import customtkinter as ctk
from tkinter import ttk
import tkinter as tk
import time
import threading
import os
import webbrowser
from tkinter import messagebox, filedialog
from PIL import ImageGrab, Image, ImageTk
import io
from model.utils import resource_path

# Tab SEO tách riêng file để giảm kích thước gui_view.py
try:
    from view.tabschuanseo import setup_seo_tab as _setup_seo_tab
except ImportError:
    try:
        from tabschuanseo import setup_seo_tab as _setup_seo_tab
    except ImportError:
        _setup_seo_tab = None

# Giả lập AppData nếu chưa có module model
class AppData:
    """Helper to store current input values."""
    def __init__(self):
        self.site_url = ""
        self.username = ""
        self.password = ""
        self.title = ""
        self.video_url = ""
        self.image_url = ""
        self.content_image = ""  # New: Image to insert in middle of content
        self.content = ""
        self.theme = "supercar"  # Default theme

class GUIView(ctk.CTk):
    def __init__(self, controller, initial_config=None):
        super().__init__()
        self.controller = controller
        self.initial_config = initial_config or {}
        
        # --- Image Cache (LRU-style) to avoid repeated disk reads ---
        self._image_cache = {}  # path -> CTkImage/PhotoImage
        self._image_cache_max = 100  # max cached images
        
        # --- Cấu hình màu sắc (GOLD THEME 🏆) ---
        self.colors = {
            'primary': '#F59E0B',       # Amber 500 (Gold-Yellow) - Main Brand Color
            'primary_hover': '#B45309', # Amber 700 (Dark Gold) - Hover State
            'success': '#10B981',       # Emerald 500
            'success_hover': '#059669', 
            'warning': '#F59E0B',       # Same as primary for consistency
            'danger': '#EF4444',        # Red 500
            'bg_light': '#FFFBEB',      # Amber 50 (Subtle Gold Tint Background)
            'bg_card': '#FFFFFF',       # White Cards
            'bg_dark': '#FEF3C7',       # Amber 100 (Light Gold Accent)
            'text_primary': '#111827',  # Gray 900
            'text_secondary': '#4B5563', # Gray 600
            'border': '#FDE68A'         # Amber 200 (Soft Gold Border)
        }
        
        # --- Cấu hình cửa sổ ---
        self.title("LVC MEDIA TOOL")
        self.geometry("1200x800")
        self.minsize(1000, 700)
        
        # --- Tắt CTk auto-DPI (nguồn lag chính khi resize) ---

        try:

            ctk.deactivate_automatic_dpi_awareness()

        except Exception:

            pass

        try:

            ctk.set_widget_scaling(1.0)

            ctk.set_window_scaling(1.0)

        except Exception:

            pass



        # --- DPI awareness thuần tk (System-aware = smooth) ---

        try:

            import ctypes

            try:

                ctypes.windll.shcore.SetProcessDpiAwareness(1)

            except Exception:

                ctypes.windll.user32.SetProcessDPIAware()

        except Exception:

            pass



        # Optimize rendering

        self.update_idletasks()



        
        # Set App Icon
        # Set App Icon
        logo_path = resource_path("logo.ico")
        if os.path.exists(logo_path):
            try:
                self.iconbitmap(logo_path)
            except: pass
        
        # Theme - LIGHT MODE mặc định
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        
        try:
            from version import get_version
            self.app_version = get_version()
        except ImportError:
            self.app_version = "2.0.3"

        # Biến dữ liệu
        self.login_frame = None
        self.main_frame = None
        self.post_queue = [] 
        
        # Initialize Update Checker
        try:
            from model.update_checker import UpdateChecker
            update_url = self.initial_config.get('update_url', "https://raw.githubusercontent.com/nguyenduyducds/WebSiteWpr/master/version.json");
            self.updater = UpdateChecker(self.app_version, update_url)
            self.after(3000, self.check_startup_update) # Check after 3s
        except Exception as e:
            print(f"Update init failed: {e}")
            
        self.content_pool = []
        self.batch_data = None
        self.published_links = []
        self.published_posts = []  # Store {title, link} pairs for copy function
        self.current_link = ""

        # Khởi tạo màn hình
        self.create_login_screen()

        # Handle X button
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def check_startup_update(self):
        """Check for updates silently on startup"""
        if hasattr(self, 'updater'):
            self.updater.check_for_updates(callback=self.on_update_found)
            
    def on_update_found(self, has_update, new_version, download_url):
        if has_update:
            msg = f"🔥 NGUYỄN DUY ĐỨC THÔNG BÁO:\n\nĐã có bản cập nhật v{new_version} !\n(Bản hiện tại: v{self.app_version})\n\n👉 Bấm OK để cập nhật ngay cho 'ngon' nhé!"
            if messagebox.askyesno("Nguyễn Duy Đức Thông Báo Update", msg):
                if download_url and download_url.endswith(".exe"):
                    # Start auto-update flow
                    self.perform_auto_update(download_url)
                elif download_url:
                    # Fallback for non-exe checks (e.g. zip)
                    webbrowser.open(download_url)
                else:
                    messagebox.showinfo("Thông tin", f"Vui lòng liên hệ Admin để nhận bản cập nhật v{new_version}.")

    def perform_auto_update(self, url):
        """Show progress window and download"""
        # Create Popup
        self.update_window = ctk.CTkToplevel(self)
        self.update_window.title("Đang cập nhật...")
        self.update_window.geometry("400x150")
        self.update_window.attributes("-topmost", True)
        
        ctk.CTkLabel(self.update_window, text="Đang tải bản cập nhật mới...", font=("Segoe UI", 14)).pack(pady=20)
        
        self.progress_bar = ctk.CTkProgressBar(self.update_window, width=300)
        self.progress_bar.pack(pady=10)
        self.progress_bar.set(0)
        
        self.lbl_progress = ctk.CTkLabel(self.update_window, text="0%")
        self.lbl_progress.pack(pady=5)
        
        # Start download
        self.updater.download_and_install(
            url, 
            progress_callback=self.on_update_progress,
            completion_callback=self.on_update_complete
        )

    def on_update_progress(self, current, total):
        try:
            percent = current / total
            self.progress_bar.set(percent)
            self.lbl_progress.configure(text=f"{int(percent*100)}%")
            self.update_window.update()
        except: pass

    def on_update_complete(self, success, result):
        if success:
            bat_path = result
            # Run batch file and exit
            try:
                import subprocess, sys
                subprocess.Popen([bat_path], shell=True)
                self.quit()
                sys.exit()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể khởi động lại: {e}")
        else:
            self.update_window.destroy()
            messagebox.showerror("Lỗi Cập Nhật", f"Tải xuống thất bại:\n{result}") 

    def on_closing(self):
        """Handle window closing event - Force cleanup"""
        try:
            # 1. Run Kill Chrome logic (Silent)
            import subprocess
            import os
            
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            bat_file = os.path.join(project_root, "kill_chrome.bat")
            
            if os.path.exists(bat_file):
                try:
                    # Run bat file silently (Wait for it)
                    subprocess.run([bat_file], shell=True, capture_output=True, timeout=5)
                    print("[EXIT] Run kill_chrome.bat success")
                except:
                    pass
            else:
                # Fallback PowerShell
                try:
                    target_path_escaped = project_root.replace("\\", "\\\\").replace("'", "''")
                    ps_cmd = f"Get-WmiObject Win32_Process | Where-Object {{ ($_.Name -eq 'chrome.exe' -or $_.Name -eq 'chromedriver.exe') -and $_.ExecutablePath -like '*{target_path_escaped}*' }} | ForEach-Object {{ Stop-Process -Id $_.ProcessId -Force }}"
                    subprocess.run(["powershell", "-NoProfile", "-Command", ps_cmd], capture_output=True, timeout=5)
                    print("[EXIT] Run Safe PowerShell success")
                except:
                    pass

        except Exception as e:
            print(f"[EXIT] Cleanup error: {e}")
            
        # 2. Destroy and Exit
        self.destroy()
        import sys
        sys.exit(0) 

    # =========================================================================
    # PHẦN 1: LOGIN SCREEN (MODERN REDESIGN ✨)
    # =========================================================================
    def _load_image_cached(self, path, size=(80, 80)):
        """Load image with caching to avoid repeated disk reads"""
        cache_key = f"{path}_{size}"
        if cache_key in self._image_cache:
            return self._image_cache[cache_key]
        
        try:
            from PIL import Image
            pil_img = Image.open(path)
            pil_img.thumbnail(size)
            result = pil_img
            
            # Evict oldest if cache full
            if len(self._image_cache) >= self._image_cache_max:
                oldest = next(iter(self._image_cache))
                del self._image_cache[oldest]
            
            self._image_cache[cache_key] = result
            return result
        except:
            return None

    def create_login_screen(self):
        # Modern Background Colors (Light Gray / Deep Dark)
        self.login_frame = ctk.CTkFrame(self, corner_radius=0, fg_color=("#F3F4F6", "#0F172A"))
        self.login_frame.pack(fill="both", expand=True)

        # --- Fake Shadow Layer (Depth Effect) ---
        shadow_box = ctk.CTkFrame(
            self.login_frame,
            corner_radius=30,
            fg_color=("#E5E7EB", "#020617"), # Subtle shadow color
            width=506, 
            height=686
        )
        shadow_box.place(relx=0.5, rely=0.5, anchor="center")

        # --- Main Card ---
        center_box = ctk.CTkFrame(
            self.login_frame, 
            corner_radius=30, 
            fg_color=("#FFFFFF", "#1E293B"), # Pure White / Slate 800
            border_width=0,
            width=500,
            height=750  # Increased from 680 to 750
        )
        center_box.place(relx=0.5, rely=0.5, anchor="center")
        center_box.pack_propagate(False) # Keep size fixed

        # Header Section
        header_frame = ctk.CTkFrame(center_box, corner_radius=20, fg_color="transparent")
        header_frame.pack(pady=(30, 15), padx=40, fill="x")  # Reduced padding
        
        # Logo Logic
        try:
            from PIL import Image
            logo_path = resource_path("logo.ico")
            if os.path.exists(logo_path):
                logo_pil = Image.open(logo_path)
                logo_img = ctk.CTkImage(light_image=logo_pil, dark_image=logo_pil, size=(70, 70))  # Reduced from 90x90
                ctk.CTkLabel(header_frame, text="", image=logo_img).pack(pady=(0, 5))
            else:
                ctk.CTkLabel(header_frame, text="🚀", font=("Segoe UI Emoji", 50)).pack(pady=(0, 5))  # Reduced from 60
        except:
             ctk.CTkLabel(header_frame, text="🚀", font=("Segoe UI Emoji", 50)).pack(pady=(0, 5))
        
        ctk.CTkLabel(
            header_frame, 
            text="LVC Tool Web", 
            font=("Segoe UI", 22, "bold"),  # Reduced from 26
            text_color=self.colors['primary']
        ).pack(pady=(5, 0))
        
        ctk.CTkLabel(
            header_frame, 
            text="Hệ thống tự động hóa WordPress chuyên nghiệp", 
            font=("Segoe UI", 11),  # Reduced from 13
            text_color=("#6B7280", "#9CA3AF") # Gray-500
        ).pack(pady=(3, 0))

        # Inputs Section
        input_frame = ctk.CTkFrame(center_box, fg_color="transparent")
        input_frame.pack(pady=5, padx=40, fill="x")  # Reduced padding from 10 to 5
        
        # Account Selector (NEW)
        from model.wp_account_manager import WPAccountManager
        self.account_manager = WPAccountManager()
        
        account_selector_frame = ctk.CTkFrame(input_frame, fg_color="transparent")
        account_selector_frame.pack(fill="x", pady=(0, 10))  # Reduced from 15 to 10
        
        ctk.CTkLabel(
            account_selector_frame,
            text="TÀI KHOẢN ĐÃ LƯU",
            font=("Segoe UI", 11, "bold"),
            text_color=("#6B7280", "#9CA3AF")
        ).pack(anchor="w", pady=(0, 5))
        
        # Dropdown + Buttons
        selector_row = ctk.CTkFrame(account_selector_frame, fg_color="transparent")
        selector_row.pack(fill="x")
        
        # Get account list
        accounts = self.account_manager.get_all_accounts()
        account_names = ["-- Chọn tài khoản --"] + [f"{acc['name']}" for acc in accounts]
        
        self.account_selector_var = ctk.StringVar(value=account_names[0])
        self.account_selector = ctk.CTkOptionMenu(
            selector_row,
            values=account_names,
            variable=self.account_selector_var,
            height=40,
            font=("Segoe UI", 12),
            fg_color=self.colors['bg_dark'],
            button_color=self.colors['border'],
            button_hover_color=self.colors['primary'],
            text_color=self.colors['text_primary'],
            dropdown_fg_color=self.colors['bg_card'],
            dropdown_hover_color=self.colors['bg_dark'],
            command=self.on_account_selected
        )
        self.account_selector.pack(side="left", fill="x", expand=True, padx=(0, 5))
        
        # Manage button
        ctk.CTkButton(
            selector_row,
            text="⚙️",
            width=40,
            height=40,
            font=("Segoe UI", 16),
            fg_color=self.colors['bg_dark'],
            hover_color=self.colors['primary'],
            command=self.open_account_manager,
            corner_radius=8
        ).pack(side="right")
        
        # Separator
        separator = ctk.CTkFrame(input_frame, height=2, fg_color=self.colors['border'])
        separator.pack(fill="x", pady=10)  # Reduced from 15 to 10
        
        self.entry_site = self.create_modern_input(
            input_frame, 
            "Site URL", 
            "https://yoursite.com/wp-admin", 
            self.initial_config.get("site_url", "")
        )
        
        self.entry_user = self.create_modern_input(
            input_frame, 
            "Username", 
            "admin", 
            self.initial_config.get("username", "")
        )
        
        self.entry_pass = self.create_modern_input(
            input_frame, 
            "Password", 
            "••••••••", 
            self.initial_config.get("password", ""), 
            show="*"
        )

        # Options Section
        checkbox_frame = ctk.CTkFrame(input_frame, fg_color="transparent")
        checkbox_frame.pack(pady=(5, 10), fill="x")  # Reduced padding
        
        self.chk_headless = ctk.CTkCheckBox(
            checkbox_frame, 
            text="Chạy ẩn (Headless - Nhanh hơn)", 
            font=("Segoe UI", 12),
            fg_color=self.colors['primary'],
            checkbox_height=20,
            checkbox_width=20,
            corner_radius=6,
            border_width=2,
            hover_color=self.colors['primary_hover']
        )
        self.chk_headless.pack(anchor="w")
        if self.initial_config.get("headless", True):
             self.chk_headless.select()

        # Login Button
        self.btn_login = ctk.CTkButton(
            input_frame, 
            text="🚀 ĐĂNG NHẬP NGAY", 
            height=50,  # Reduced from 55
            font=("Segoe UI", 13, "bold"),  # Reduced from 14
            fg_color=self.colors['primary'],
            hover_color=self.colors['primary_hover'],
            corner_radius=25,  # Reduced from 28
            command=self.on_login_click
        )
        self.btn_login.pack(pady=(5, 10), fill="x")  # Reduced padding

        # Status Label
        self.lbl_status = ctk.CTkLabel(
            input_frame, 
            text="", 
            font=("Segoe UI", 10),  # Reduced from 11
            text_color=self.colors['danger']
        )
        self.lbl_status.pack(pady=(0, 5))

    def create_modern_input(self, parent, label_text, placeholder, initial_value="", show=None):
        # Container
        input_container = ctk.CTkFrame(parent, fg_color="transparent")
        input_container.pack(fill="x", pady=8)
        
        # Label (Uppercase, Smaller, Gray)
        ctk.CTkLabel(
            input_container, 
            text=label_text.upper(), 
            font=("Segoe UI", 11, "bold"), 
            anchor="w",
            text_color=("#6B7280", "#9CA3AF") # Gray-500
        ).pack(fill="x", pady=(0, 5))
        
        # Input Field (Light Gray Background)
        entry = ctk.CTkEntry(
            input_container, 
            placeholder_text=placeholder, 
            height=48, # Taller for comfort
            font=("Segoe UI", 13),
            corner_radius=12,
            border_width=1,
            border_color=("#E5E7EB", "#374151"), # Gray-200
            fg_color=("#F9FAFB", "#111827"), # Gray-50 / Gray-900 background
            text_color=("black", "white"),
            placeholder_text_color=("#9CA3AF", "#6B7280"),
            show=show
        )
        entry.pack(fill="x")
        
        if initial_value: 
            try:
                self.after(10, lambda e=entry, v=initial_value: self._safe_insert(e, v))
            except Exception as e:
                print(f"[GUI] Error inserting initial value: {e}")
        return entry
    
    def create_form_input(self, parent, label_text, placeholder, initial_value="", show=None, height=45):
        """Create a modern form input with better styling"""
        # Container
        input_container = ctk.CTkFrame(parent, fg_color="transparent")
        input_container.pack(fill="x", pady=10)
        
        # Label
        ctk.CTkLabel(
            input_container, 
            text=label_text, 
            font=("Segoe UI", 13, "bold"), 
            anchor="w",
            text_color=self.colors['text_primary']
        ).pack(fill="x", pady=(0, 8))
        
        # Entry
        entry = ctk.CTkEntry(
            input_container, 
            placeholder_text=placeholder, 
            height=height,
            font=("Segoe UI", 12),
            corner_radius=12,
            border_width=2,
            border_color=self.colors['border'],
            fg_color=self.colors['bg_light'],
            text_color=self.colors['text_primary'],
            placeholder_text_color=self.colors['text_secondary'],
            show=show
        )
        entry.pack(fill="x")
        
        if initial_value:
            self.after(10, lambda e=entry, v=initial_value: self._safe_insert(e, v))
        
        return entry
    
    def create_image_input_section(self, parent):
        """Create image input section with paste support"""
        section = ctk.CTkFrame(parent, fg_color="transparent")
        section.pack(fill="x", pady=10)
        
        # Label
        ctk.CTkLabel(
            section, 
            text="🖼️ Ảnh Thumbnail", 
            font=("Segoe UI", 13, "bold"), 
            anchor="w",
            text_color=self.colors['text_primary']
        ).pack(fill="x", pady=(0, 8))
        
        # Hint
        hint_frame = ctk.CTkFrame(section, fg_color="transparent")
        hint_frame.pack(fill="x", pady=(0, 8))
        
        ctk.CTkLabel(
            hint_frame, 
            text="💡 Tip: Chụp màn hình rồi nhấn Ctrl+V để paste ảnh trực tiếp!", 
            font=("Segoe UI", 11), 
            text_color=self.colors['success'], 
            anchor="w"
        ).pack(side="left")
    
    def _safe_insert(self, entry, value):
        """Safely insert value into entry without blocking"""
        try:
            entry.delete(0, "end")
            entry.insert(0, value)
        except Exception as e:
            print(f"[GUI] Error in safe insert: {e}")
    
    def _handle_title_paste(self, event=None):
        """Handle title paste asynchronously to prevent freeze with Unicode"""
        try:
            # Get clipboard content
            clipboard_text = self.clipboard_get()
            
            # Clear current content
            self.entry_title.delete(0, "end")
            
            # Insert asynchronously to prevent blocking
            self.after(10, lambda: self._safe_insert(self.entry_title, clipboard_text))
            
            # Prevent default paste behavior
            return "break"
        except Exception as e:
            print(f"[GUI] Error handling title paste: {e}")
            # Allow default behavior if our handler fails
            return None

    def on_theme_changed(self, choice):
        """Update theme description when user selects a theme"""
        descriptions = {
            "⚪ Không dùng theme (Raw HTML)": "💡 Không áp dụng CSS, chỉ post HTML thuần như cũ",
            "🏎️ Supercar News (Premium)": "💡 Premium automotive design with luxury styling (English content)",
            "📰 Breaking News": "💡 Modern news layout with breaking badge and trending style",
            "📝 Classic Blog": "💡 Clean and simple blog layout for general content",
            "✨ Minimal Clean": "💡 Ultra simple and elegant design with serif fonts",
            "💻 Tech Modern": "💡 Developer-friendly tech style with code support",
            "📖 Magazine": "💡 Editorial magazine style with drop cap",
            "💼 Business Pro": "💡 Professional business layout for corporate content",
            "🌸 Lifestyle": "💡 Warm and friendly blog style for lifestyle content",
            "🌙 Dark Mode": "💡 Modern dark theme for night reading"
        }
        
        desc = descriptions.get(choice, "")
        if hasattr(self, 'theme_desc_label'):
            self.theme_desc_label.configure(text=desc)
        
        print(f"[GUI] Theme changed to: {choice}")
    
    def get_selected_theme_id(self):
        """Convert theme display name to theme ID"""
        theme_map = {
            "⚪ Không dùng theme (Raw HTML)": "none",
            "🏎️ Supercar News (Premium)": "supercar",
            "📰 Breaking News": "news",
            "📝 Classic Blog": "default",
            "✨ Minimal Clean": "minimal",
            "💻 Tech Modern": "tech",
            "📖 Magazine": "magazine",
            "💼 Business Pro": "business",
            "🌸 Lifestyle": "lifestyle",
            "🌙 Dark Mode": "dark"
        }
        
        selected = self.theme_var.get() if hasattr(self, 'theme_var') else "⚪ Không dùng theme (Raw HTML)"
        return theme_map.get(selected, "none")  # Default to "none" instead of "supercar"

    def on_account_selected(self, selected_name):
        """Khi chọn tài khoản từ dropdown"""
        if selected_name == "-- Chọn tài khoản --":
            return
        
        # Find account by name
        accounts = self.account_manager.get_all_accounts()
        for i, acc in enumerate(accounts):
            if acc['name'] == selected_name:
                # Get decrypted account
                account = self.account_manager.get_account(i)
                if account:
                    # Fill form
                    self.entry_site.delete(0, tk.END)
                    self.entry_site.insert(0, account['site_url'])
                    
                    self.entry_user.delete(0, tk.END)
                    self.entry_user.insert(0, account['username'])
                    
                    self.entry_pass.delete(0, tk.END)
                    self.entry_pass.insert(0, account['password'])
                    
                    # Set headless
                    if account.get('is_headless', True):
                        self.chk_headless.select()
                    else:
                        self.chk_headless.deselect()
                    
                    # Update last used
                    self.account_manager.update_last_used(i)
                    
                    self.lbl_status.configure(
                        text=f"✅ Đã load tài khoản: {account['name']}", 
                        text_color=self.colors['success']
                    )
                break
    
    def open_account_manager(self):
        """Mở cửa sổ quản lý tài khoản"""
        # Create popup window
        popup = ctk.CTkToplevel(self)
        popup.title("Quản Lý Tài Khoản WordPress")
        popup.geometry("900x600")
        popup.transient(self)
        popup.grab_set()
        
        # Main frame
        main = ctk.CTkFrame(popup)
        main.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        ctk.CTkLabel(
            main,
            text="👤 Quản Lý Tài Khoản WordPress",
            font=("Segoe UI", 18, "bold")
        ).pack(pady=(0, 20))
        
        # Split layout
        content = ctk.CTkFrame(main)
        content.pack(fill="both", expand=True)
        
        # LEFT: List
        left = ctk.CTkFrame(content)
        left.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        ctk.CTkLabel(left, text="📋 Danh Sách", font=("Segoe UI", 13, "bold")).pack(pady=10)
        
        listbox = tk.Listbox(
            left,
            font=("Segoe UI", 11),
            bg="#2b2b2b",
            fg="white",
            selectbackground="#1f6aa5",
            height=15
        )
        listbox.pack(fill="both", expand=True, padx=10, pady=5)
        
        # RIGHT: Form
        right = ctk.CTkFrame(content)
        right.pack(side="right", fill="both", expand=True)
        
        ctk.CTkLabel(right, text="➕ Thêm/Sửa", font=("Segoe UI", 13, "bold")).pack(pady=10)
        
        form = ctk.CTkFrame(right)
        form.pack(fill="both", expand=True, padx=10)
        
        ctk.CTkLabel(form, text="Tên:", font=("Segoe UI", 11)).pack(anchor="w", pady=(5, 0))
        name_entry = ctk.CTkEntry(form, placeholder_text="Site chính...")
        name_entry.pack(fill="x", pady=5)
        
        ctk.CTkLabel(form, text="URL:", font=("Segoe UI", 11)).pack(anchor="w", pady=(5, 0))
        url_entry = ctk.CTkEntry(form, placeholder_text="https://...")
        url_entry.pack(fill="x", pady=5)
        
        ctk.CTkLabel(form, text="Username:", font=("Segoe UI", 11)).pack(anchor="w", pady=(5, 0))
        user_entry = ctk.CTkEntry(form, placeholder_text="admin")
        user_entry.pack(fill="x", pady=5)
        
        ctk.CTkLabel(form, text="Password:", font=("Segoe UI", 11)).pack(anchor="w", pady=(5, 0))
        pass_entry = ctk.CTkEntry(form, placeholder_text="••••••", show="•")
        pass_entry.pack(fill="x", pady=5)
        
        headless_var = tk.BooleanVar(value=True)
        ctk.CTkCheckBox(form, text="Headless", variable=headless_var).pack(anchor="w", pady=10)
        
        # Functions
        def refresh_list():
            listbox.delete(0, tk.END)
            for i, acc in enumerate(self.account_manager.get_all_accounts()):
                listbox.insert(tk.END, f"[{i+1}] {acc['name']} - {acc['username']}")
        
        def on_select(event):
            sel = listbox.curselection()
            if sel:
                acc = self.account_manager.get_account(sel[0])
                if acc:
                    name_entry.delete(0, tk.END)
                    name_entry.insert(0, acc['name'])
                    url_entry.delete(0, tk.END)
                    url_entry.insert(0, acc['site_url'])
                    user_entry.delete(0, tk.END)
                    user_entry.insert(0, acc['username'])
                    pass_entry.delete(0, tk.END)
                    pass_entry.insert(0, acc['password'])
                    headless_var.set(acc.get('is_headless', True))
        
        def save_new():
            name = name_entry.get().strip()
            url = url_entry.get().strip()
            user = user_entry.get().strip()
            pwd = pass_entry.get().strip()
            if not all([name, url, user, pwd]):
                messagebox.showwarning("Thiếu thông tin", "Vui lòng điền đầy đủ!")
                return
            success, msg = self.account_manager.add_account(name, url, user, pwd, headless_var.get())
            if success:
                messagebox.showinfo("Thành công", msg)
                refresh_list()
                # Refresh dropdown in login screen
                self.refresh_account_dropdown()
            else:
                messagebox.showerror("Lỗi", msg)
        
        def delete_acc():
            sel = listbox.curselection()
            if not sel:
                messagebox.showwarning("Chưa chọn", "Chọn tài khoản cần xóa!")
                return
            if messagebox.askyesno("Xác nhận", "Xóa tài khoản này?"):
                success, msg = self.account_manager.delete_account(sel[0])
                if success:
                    messagebox.showinfo("Thành công", msg)
                    refresh_list()
                    self.refresh_account_dropdown()
        
        listbox.bind('<<ListboxSelect>>', on_select)
        
        # Buttons
        btn_frame = ctk.CTkFrame(form)
        btn_frame.pack(fill="x", pady=20)
        
        ctk.CTkButton(btn_frame, text="💾 Lưu", command=save_new, fg_color="#28a745", width=100).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="🗑️ Xóa", command=delete_acc, fg_color="#dc3545", width=100).pack(side="left", padx=5)
        
        # Bottom buttons
        bottom = ctk.CTkFrame(main)
        bottom.pack(fill="x", pady=(10, 0))
        
        ctk.CTkButton(
            bottom,
            text="✅ Đóng",
            command=popup.destroy,
            width=120
        ).pack(side="right")
        
        refresh_list()
    
    def refresh_account_dropdown(self):
        """Refresh dropdown sau khi thêm/xóa tài khoản"""
        if hasattr(self, 'account_selector'):
            accounts = self.account_manager.get_all_accounts()
            account_names = ["-- Chọn tài khoản --"] + [f"{acc['name']}" for acc in accounts]
            self.account_selector.configure(values=account_names)
    
    def on_login_click(self):
        site = self.entry_site.get().strip()
        user = self.entry_user.get().strip()
        pwd = self.entry_pass.get().strip()
        if not all([site, user, pwd]):
            self.lbl_status.configure(text="⚠️ Vui lòng nhập đủ thông tin!")
            return
        
        self.btn_login.configure(state="disabled", text="🔄 Đang kết nối...")
        # Gọi controller (Giả lập)
        if self.controller:
            self.controller.handle_login(site, user, pwd, headless=bool(self.chk_headless.get()))
        else:
            self.after(1000, self.login_success) # Test only

    def login_success(self):
        self.lbl_status.configure(text="✅ Đăng nhập thành công!", text_color=self.colors['success'])
        self.after(500, self.switch_to_main_screen)

    def login_failed(self, message):
        self.lbl_status.configure(text=f"❌ Lỗi: {message}", text_color=self.colors['danger'])
        self.btn_login.configure(state="normal", text="🚀 ĐĂNG NHẬP")

    def switch_to_main_screen(self):
        self.login_frame.destroy()
        self.create_main_screen()

    # =========================================================================
    # PHẦN 2: MAIN SCREEN (Layout đã được làm sạch)
    # =========================================================================
    def create_main_screen(self):
        # Container chính
        self.main_frame = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # 1. Header (Thanh trên cùng)
        self.create_header()

        # 2. Nội dung Tab (Trung tâm)
        self.create_tabs()

        # 3. Status Bar (Thanh trạng thái dưới cùng)
        self.create_status_bar()

    def create_header(self):
        header = ctk.CTkFrame(
            self.main_frame, 
            height=80, 
            corner_radius=16, 
            fg_color=self.colors['bg_card'],
            border_width=2,
            border_color=self.colors['border']
        )
        header.pack(fill="x", pady=(0, 15))
        
        # Logo & Title with better spacing
        title_frame = ctk.CTkFrame(header, fg_color="transparent")
        title_frame.pack(side="left", padx=25, pady=15)
        
        ctk.CTkLabel(
            title_frame, 
            text="LVC Tool Web", 
            font=("Segoe UI", 22, "bold"), 
            text_color=self.colors['primary']
        ).pack(anchor="w")
        
        ctk.CTkLabel(
            title_frame, 
            text="Chức năng đăng bài lên website WordPress", 
            font=("Segoe UI", 12), 
            text_color=self.colors['text_secondary']
        ).pack(anchor="w", pady=(2, 0))
        
        # Version info
        version_text = f"v{self.app_version}"
        
        ctk.CTkLabel(
            title_frame,
            text=version_text,
            font=("Segoe UI", 10, "bold"),
            text_color=self.colors['primary'],
            fg_color=self.colors['bg_dark'],
            corner_radius=6,
            padx=10,
            pady=4
        ).pack(anchor="w", pady=(5, 0))

        # User Info & Logout with better styling
        user_frame = ctk.CTkFrame(header, fg_color="transparent")
        user_frame.pack(side="right", padx=25, pady=15)
        
        username = self.controller.username if hasattr(self.controller, 'username') else "Admin"
        
        # User badge
        user_badge = ctk.CTkFrame(
            user_frame, 
            fg_color=self.colors['bg_dark'],
            corner_radius=10,
            border_width=1,
            border_color=self.colors['border']
        )
        user_badge.pack(side="left", padx=(0, 15))
        
        ctk.CTkLabel(
            user_badge, 
            text=f"👤 {username}", 
            font=("Segoe UI", 13, "bold"),
            text_color=self.colors['text_primary']
        ).pack(padx=15, pady=8)
        
        # Logout button
        ctk.CTkButton(
            user_frame, 
            text="� Đăng Xuất", 
            width=100, 
            height=36, 
            font=("Segoe UI", 12, "bold"),
            fg_color=self.colors['danger'], 
            hover_color="#dc2626",
            corner_radius=10,
            command=self.logout
        ).pack(side="left")
        
        # About/Version button in header
        ctk.CTkButton(
            user_frame,
            text="ℹ️ Thông Tin",
            width=110,
            height=36,
            font=("Segoe UI", 11),
            fg_color=self.colors['bg_dark'],
            hover_color=self.colors['border'],
            text_color=self.colors['text_primary'],
            corner_radius=10,
            command=self.show_about_dialog
        ).pack(side="left", padx=(0, 10))

    def create_tabs(self):
        # Container bọc nav + content
        tab_container = tk.Frame(self.main_frame, bg=self.colors["bg_light"])
        tab_container.pack(fill="both", expand=True)

        # ── NAV BAR — tk.Frame thuần cho smooth resize ──

        nav_bar = tk.Frame(

            tab_container,

            bg=self.colors['bg_dark'],

            height=62,

        )

        nav_bar.pack(fill="x", pady=(0, 6))

        nav_bar.pack_propagate(False)

        nav_inner = tk.Frame(nav_bar, bg=self.colors['bg_dark'])
        nav_inner.pack(fill="both", expand=True, padx=10, pady=8)

        # ── CONTENT AREA ─────────────────────────────────────────────
        content_area = ctk.CTkFrame(
            tab_container,
            fg_color=self.colors['bg_card'],
            corner_radius=16,
            border_width=2,
            border_color=self.colors['border']
        )
        content_area.pack(fill="both", expand=True)

        # Tab definitions: (button label, attribute name)
        tab_defs = [
            ("📝  Đăng Bài",              "tab_post"),
            ("📱  Scan Link",              "tab_fb_import"),
            ("📦  Hàng Chờ",              "tab_batch"),
            ("☁️  Upload",                "tab_upload"),
            ("🎥  Vimeo",                  "tab_vimeo"),
            ("🖼️  Ảnh",                   "tab_images"),
            ("🤖  AI Thumbnail",           "tab_thumbnail_ai"),
            ("🔍  Web Chuẩn SEO",          "tab_seo"),
            ("🗒️  Logs",                  "tab_data"),
            ("⚙️  Cài Đặt",               "tab_settings"),
        ]

        # Tạo content frames và gán vào self
        frames = {}
        for _label, attr in tab_defs:
            frame = ctk.CTkFrame(content_area, fg_color="transparent")
            frames[attr] = frame
            setattr(self, attr, frame)

        # State
        self._active_nav_btn = [None]

        def switch_tab(attr, btn):
            # Ẩn tất cả
            for f in frames.values():
                f.pack_forget()
            # Hiện frame được chọn
            frames[attr].pack(fill="both", expand=True)
            # Reset style tất cả buttons
            for b in nav_btns:
                b.configure(
                    bg=self.colors['bg_dark'],
                    fg=self.colors['text_secondary'],
                    font=("Segoe UI", 11)
                )
            # Active style
            btn.configure(
                bg=self.colors['primary'],
                fg="white",
                font=("Segoe UI", 11, "bold")
            )
            self._active_nav_btn[0] = btn

        nav_btns = []
        for label, attr in tab_defs:
            btn = tk.Button(
                nav_inner,
                text=label,
                font=("Segoe UI", 11),
                bg=self.colors['bg_dark'],
                fg=self.colors['text_secondary'],
                activebackground=self.colors['primary'],
                activeforeground="white",
                relief="flat",
                bd=0,
                padx=12,
                pady=4,
                cursor="hand2",
                highlightthickness=0,
            )
            btn.configure(command=lambda a=attr, b=btn: switch_tab(a, b))
            btn.pack(side="left", padx=3, pady=2)
            # Hover effects
            btn.bind("<Enter>", lambda e, b=btn: b.configure(
                bg=self.colors['border'], fg=self.colors['text_primary']
            ) if b.cget('bg') != self.colors['primary'] else None)
            btn.bind("<Leave>", lambda e, b=btn: b.configure(
                bg=self.colors['bg_dark'], fg=self.colors['text_secondary']
            ) if b.cget('bg') != self.colors['primary'] else None)
            nav_btns.append(btn)

        # Hiện tab đầu tiên mặc định
        first_attr = tab_defs[0][1]
        frames[first_attr].pack(fill="both", expand=True)
        nav_btns[0].configure(
            bg=self.colors['primary'],
            fg="white",
            font=("Segoe UI", 11, "bold")
        )
        self._active_nav_btn[0] = nav_btns[0]

        # Hàm tiện ích để switch tab từ code khác
        def _switch_by_attr(attr):
            for i, (_, a) in enumerate(tab_defs):
                if a == attr:
                    switch_tab(a, nav_btns[i])
                    break
        self.switch_tab = _switch_by_attr

        # Xây dựng nội dung từng Tab
        self.create_fb_import_tab_content()
        self.create_post_tab_content()
        self.create_batch_tab_content()
        self.create_upload_tab_content()
        self.create_vimeo_tab_content()
        self.create_images_tab_content()
        self.create_thumbnail_ai_tab_content()
        self._build_seo_tab()
        self.create_data_tab_content()
        self.create_settings_tab_content()

    def _build_seo_tab(self):
        """Khởi tạo tab SEO từ file tabschuanseo.py (tách riêng cho gọn)."""
        if _setup_seo_tab is not None:
            _setup_seo_tab(self, self.tab_seo)
        else:
            ctk.CTkLabel(
                self.tab_seo,
                text="⚠️ Không tải được module tabschuanseo.py",
                font=("Segoe UI", 13),
                text_color="#EF4444"
            ).pack(pady=40)


    def create_status_bar(self):
        self.status_frame = ctk.CTkFrame(
            self.main_frame, 
            height=40, 
            corner_radius=12,
            fg_color=self.colors['bg_card'],
            border_width=2,
            border_color=self.colors['border']
        )
        self.status_frame.pack(fill="x", pady=(15, 0))
        
        self.status_label = ctk.CTkLabel(
            self.status_frame, 
            text="✅ Sẵn sàng", 
            font=("Segoe UI", 12, "bold"), 
            text_color=self.colors['success']
        )
        self.status_label.pack(side="left", padx=20, pady=8)
        
        # Kill Chrome button
        ctk.CTkButton(
            self.status_frame,
            text="🔪 Kill Chrome",
            width=110,
            height=28,
            font=("Segoe UI", 10, "bold"),
            fg_color="#dc2626",
            hover_color="#b91c1c",
            text_color="white",
            corner_radius=8,
            command=self.kill_chrome_processes
        ).pack(side="left", padx=10, pady=8)
        
        # Hello Kitty Loading Indicator (Small GIF in Status Bar) 🐱💕
        self.kitty_loading_frame = ctk.CTkFrame(self.status_frame, fg_color="transparent")
        
        # 1. GIF Label
        self.kitty_gif_label = ctk.CTkLabel(self.kitty_loading_frame, text="")
        self.kitty_gif_label.pack(side="left", padx=5)
        
        # Load small pink kitty for loading indicator
        try:
            from model.animated_gif import AnimatedGIF
            gif_path = resource_path("animaition/Hello Kitty Pink GIF.gif")
            if os.path.exists(gif_path):
                # Keep reference (important!)
                self.loading_gif = AnimatedGIF(self.kitty_gif_label, gif_path, size=(40, 40))
                self.loading_gif.play()
            else:
                 self.kitty_gif_label.configure(text="🐱", font=("Segoe UI Emoji", 24))
        except: 
             self.kitty_gif_label.configure(text="🐱", font=("Segoe UI Emoji", 24))
        
        # 2. Text Label
        self.kitty_status_label = ctk.CTkLabel(
            self.kitty_loading_frame,
            text="Đang scan... 💕",
            font=("Segoe UI", 12, "bold"),
            text_color=("#FF69B4", "#FF1493")
        )
        self.kitty_status_label.pack(side="left")
        
        # Hide initially
        self.kitty_loading_frame.pack_forget()
        
        # ===== CUTE GIF DECORATIONS (Top Right) 🎀 =====
        decoration_frame = ctk.CTkFrame(self.status_frame, fg_color="transparent")
        decoration_frame.pack(side="right", padx=5, pady=8)
        
        # 1. Coffee Kitty (Left) ☕
        coffee_label = ctk.CTkLabel(decoration_frame, text="")
        coffee_label.pack(side="left", padx=2)
        try:
            from model.animated_gif import AnimatedGIF
            coffee_path = resource_path("animaition/hello kitty coffee GIF.gif")
            if os.path.exists(coffee_path):
                self.coffee_gif = AnimatedGIF(coffee_label, coffee_path, size=(32, 32))
                self.coffee_gif.play()
            else:
                coffee_label.configure(text="☕", font=("Segoe UI Emoji", 20))
        except: coffee_label.configure(text="☕", font=("Segoe UI Emoji", 20))
        
        # 2. Pink Kitty (Middle) 💕
        pink_label = ctk.CTkLabel(decoration_frame, text="")
        pink_label.pack(side="left", padx=2)
        try:
            from model.animated_gif import AnimatedGIF
            pink_path = resource_path("animaition/Hello Kitty Pink GIF.gif")
            if os.path.exists(pink_path):
                self.pink_gif = AnimatedGIF(pink_label, pink_path, size=(32, 32))
                self.pink_gif.play()
            else:
                pink_label.configure(text="💕", font=("Segoe UI Emoji", 20))
        except: pink_label.configure(text="💕", font=("Segoe UI Emoji", 20))
        
        # 3. OMG Kitty (Right) 😱
        omg_label = ctk.CTkLabel(decoration_frame, text="")
        omg_label.pack(side="left", padx=2)
        try:
            from model.animated_gif import AnimatedGIF
            omg_path = resource_path("animaition/Hello Kitty Omg GIF by Feliks Tomasz Konczakowski.gif")
            if os.path.exists(omg_path):
                self.omg_gif = AnimatedGIF(omg_label, omg_path, size=(32, 32))
                self.omg_gif.play()
            else:
                omg_label.configure(text="✨", font=("Segoe UI Emoji", 20))
        except: omg_label.configure(text="✨", font=("Segoe UI Emoji", 20))
        
        # Connection indicator
        connection_frame = ctk.CTkFrame(self.status_frame, fg_color="transparent")
        connection_frame.pack(side="right", padx=20, pady=8)
        
        ctk.CTkLabel(
            connection_frame, 
            text="🟢", 
            font=("Segoe UI", 14)
        ).pack(side="left", padx=(0, 5))
        
        ctk.CTkLabel(
            connection_frame, 
            text="Connected", 
            font=("Segoe UI", 11), 
            text_color=self.colors['text_secondary']
        ).pack(side="left")

    # =========================================================================
    # TAB 1: ĐĂNG BÀI LẺ (Post Form) - WEB LAYOUT STYLE
    # =========================================================================
    def create_post_tab_content(self):
        container = ctk.CTkScrollableFrame(
            self.tab_post, 
            fg_color="transparent"
        )
        container.pack(fill="both", expand=True, padx=0, pady=0)

        # Header Title
        title_frame = ctk.CTkFrame(container, fg_color="transparent")
        title_frame.pack(fill="x", padx=20, pady=(15, 10))
        ctk.CTkLabel(title_frame, text="📝 Soạn Thảo Bài Viết", font=("Segoe UI", 20, "bold")).pack(side="left")
        ctk.CTkLabel(title_frame, text=" | Editor Mode", font=("Segoe UI", 12), text_color="gray").pack(side="left", padx=5, pady=(5,0))

        # === GRID LAYOUT CONTAINER ===
        grid = ctk.CTkFrame(container, fg_color="transparent")
        grid.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Configure grid weights
        grid.columnconfigure(0, weight=2) # Left column (Main) ~66%
        grid.columnconfigure(1, weight=1) # Right column (Sidebar) ~33%

        # ================= LEFT COLUMN: MAIN CONTENT =================
        left_col = ctk.CTkFrame(grid, fg_color="transparent")
        left_col.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # 1. Main Info Card
        main_card = ctk.CTkFrame(left_col, fg_color=self.colors['bg_card'], corner_radius=12, border_width=1, border_color=self.colors['border'])
        main_card.pack(fill="both", expand=True, pady=(0, 15))

        # Title Input
        input_container = ctk.CTkFrame(main_card, fg_color="transparent")
        input_container.pack(fill="x", padx=20, pady=20)
        
        self.entry_title = self.create_form_input(
            input_container, 
            "📌 Tiêu đề bài viết", 
            "Nhập tiêu đề hấp dẫn vào đây...",
            height=50
        )
        # Bind paste handlers
        self.entry_title.bind('<Control-v>', self._handle_title_paste)
        
        # Link Video (More space)
        vid_container = ctk.CTkFrame(main_card, fg_color="transparent")
        vid_container.pack(fill="x", padx=20, pady=(0, 20))
        
        ctk.CTkLabel(vid_container, text="🎬 Link Video / Embed Code", font=("Segoe UI", 13, "bold"), text_color=self.colors['text_primary']).pack(anchor="w", pady=(0, 5))
        ctk.CTkLabel(vid_container, text="Hỗ trợ nhiều dòng (Facebook, Youtube, Vimeo...)", font=("Segoe UI", 11), text_color="gray").pack(anchor="w", pady=(0, 5))

        self.entry_video = ctk.CTkTextbox(
            vid_container,
            height=220,
            font=("Consolas", 12),
            corner_radius=10,
            border_width=2,
            border_color=self.colors['border'],
            fg_color=self.colors['bg_light'],
            text_color=self.colors['text_primary'],
        )
        self.entry_video.pack(fill="x")
        
        # Content Editor
        content_container = ctk.CTkFrame(main_card, fg_color="transparent")
        content_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        ctk.CTkLabel(content_container, text="📝 Nội dung chi tiết (HTML)", font=("Segoe UI", 13, "bold"), text_color=self.colors['text_primary']).pack(anchor="w", pady=(0, 5))
        
        self.textbox_content = ctk.CTkTextbox(
            content_container, 
            height=250, # Taller for writing
            font=("Segoe UI", 12),
            corner_radius=10,
            border_width=2,
            border_color=self.colors['border']
        )
        self.textbox_content.pack(fill="both", expand=True)

        # Bulk Import moved to 'Facebook Scan' tab


        # ================= RIGHT COLUMN: SIDEBAR =================
        right_col = ctk.CTkFrame(grid, fg_color="transparent")
        right_col.grid(row=0, column=1, sticky="nsew", padx=(0, 0))
        
        # 1. Action Panel (Sticky Top)
        action_card = ctk.CTkFrame(right_col, fg_color=self.colors['bg_card'], corner_radius=12, border_width=1, border_color=self.colors['border'])
        action_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(action_card, text="🚀 Tác Vụ", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=15, pady=15)
        
        self.btn_post = ctk.CTkButton(
            action_card, 
            text="ĐĂNG BÀI NGAY", 
            height=45,
            font=("Segoe UI", 13, "bold"),
            fg_color=self.colors['primary'], 
            hover_color=self.colors['primary_hover'],
            command=self.on_post_click
        )
        self.btn_post.pack(fill="x", padx=15, pady=(0, 10))
        
        ctk.CTkButton(
            action_card, 
            text="➕ Thêm vào Queue", 
            height=35,
            font=("Segoe UI", 12, "bold"),
            fg_color=self.colors['warning'], 
            hover_color="#d97706", 
            text_color="white",
            command=self.add_to_queue
        ).pack(fill="x", padx=15, pady=(0, 15))

        # 2. Settings Card
        settings_card = ctk.CTkFrame(right_col, fg_color=self.colors['bg_card'], corner_radius=12, border_width=1, border_color=self.colors['border'])
        settings_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(settings_card, text="⚙️ Cấu Hình", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=15, pady=(15, 10))
        
        # Theme
        ctk.CTkLabel(settings_card, text="Giao diện (Theme):", font=("Segoe UI", 12)).pack(anchor="w", padx=15, pady=(0, 5))
        
        self.theme_var = ctk.StringVar(value="⚪ Không dùng theme (Raw HTML)")
        theme_dropdown = ctk.CTkOptionMenu(
            settings_card,
            values=[ "⚪ Không dùng theme (Raw HTML)", "🏎️ Supercar News (Premium)", "📰 Breaking News", "📝 Classic Blog", "✨ Minimal Clean", "💻 Tech Modern", "📖 Magazine", "💼 Business Pro", "🌸 Lifestyle", "🌙 Dark Mode" ],
            variable=self.theme_var,
            height=35,
            font=("Segoe UI", 11),
            fg_color=self.colors['bg_dark'],
            button_color=self.colors['border'],
            button_hover_color=self.colors['primary'],
            text_color=self.colors['text_primary'],
            dropdown_fg_color=self.colors['bg_card'],
            dropdown_hover_color=self.colors['bg_dark'],
            command=self.on_theme_changed
        )
        theme_dropdown.pack(fill="x", padx=15, pady=(0, 5))
        
        self.theme_desc_label = ctk.CTkLabel(settings_card, text="💡 Basic HTML mode", font=("Segoe UI", 10), text_color="gray", anchor="w", wraplength=250)
        self.theme_desc_label.pack(fill="x", padx=15, pady=(0, 15))

        # 3. Media Card
        media_card = ctk.CTkFrame(right_col, fg_color=self.colors['bg_card'], corner_radius=12, border_width=1, border_color=self.colors['border'])
        media_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(media_card, text="🖼️ Media & Ảnh", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=15, pady=(15, 10))
        
        # Thumbnail
        ctk.CTkLabel(media_card, text="Ảnh Đại Diện (Thumbnail):", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15, pady=(0, 5))
        
        thumb_row = ctk.CTkFrame(media_card, fg_color="transparent")
        thumb_row.pack(fill="x", padx=15, pady=(0, 15)) 
        
        self.entry_image = ctk.CTkEntry(thumb_row, placeholder_text="URL hoặc Paste...", height=35)
        self.entry_image.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.entry_image.bind('<Control-v>', self.paste_image_from_clipboard)
        
        ctk.CTkButton(thumb_row, text="📂", width=40, height=35, command=self.browse_thumbnail).pack(side="right")
        
        # Content Images
        ctk.CTkLabel(media_card, text="Ảnh Content (Chèn giữa):", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15, pady=(0, 5))
        
        self.chk_auto_fetch_images = ctk.CTkCheckBox(media_card, text="Tự động lấy ảnh API", font=("Segoe UI", 11), checkbox_height=18, checkbox_width=18)
        self.chk_auto_fetch_images.pack(anchor="w", padx=15, pady=(0, 8))
        self.chk_auto_fetch_images.select()  # Bật mặc định
        
        # Helper to create mini image row
        def create_mini_img_input(parent, placeholder, bind_cmd, browse_cmd):
            r = ctk.CTkFrame(parent, fg_color="transparent")
            r.pack(fill="x", padx=15, pady=(0, 8))
            e = ctk.CTkEntry(r, placeholder_text=placeholder, height=30, font=("Segoe UI", 11))
            e.pack(side="left", fill="x", expand=True, padx=(0, 5))
            e.bind('<Control-v>', bind_cmd)
            ctk.CTkButton(r, text="📂", width=35, height=30, command=browse_cmd).pack(side="right")
            return e

        self.entry_content_image = create_mini_img_input(media_card, "Ảnh 1...", lambda e: self.paste_image_from_clipboard_to_field(e, self.entry_content_image, "content1"), self.browse_content_image)
        self.entry_content_image2 = create_mini_img_input(media_card, "Ảnh 2...", lambda e: self.paste_image_from_clipboard_to_field(e, self.entry_content_image2, "content2"), self.browse_content_image2)
        self.entry_content_image3 = create_mini_img_input(media_card, "Ảnh 3...", lambda e: self.paste_image_from_clipboard_to_field(e, self.entry_content_image3, "content3"), self.browse_content_image3)
        
        ctk.CTkLabel(media_card, text="", height=5).pack() # Spacer

        # Results moved to 'Facebook Scan' tab


        # Result Area


    def create_fb_import_tab_content(self):
        # Full Layout: Horizontal Split (Left Input, Right Results)
        grid = ctk.CTkFrame(self.tab_fb_import, fg_color="transparent")
        grid.pack(fill="both", expand=True, padx=5, pady=5)
        
        # LEFT COLUMN (Input) - 40%
        left_col = ctk.CTkFrame(grid, width=350, corner_radius=12, fg_color=self.colors['bg_card'], border_width=1, border_color=self.colors['border'])
        left_col.pack(side="left", fill="both", expand=False, padx=(0, 5))
        left_col.pack_propagate(False) # Force width
        
        # Header
        ctk.CTkLabel(left_col, text="Nhập link", font=("Segoe UI", 16, "bold"), text_color=self.colors['primary']).pack(pady=(15, 5))
        ctk.CTkLabel(left_col, text="Nhập link để phân tích", font=("Segoe UI", 12), text_color="gray").pack(pady=(0, 15))
        
        # NEW: Excel Import/Export Buttons
        excel_frame = ctk.CTkFrame(left_col, fg_color="transparent")
        excel_frame.pack(fill="x", padx=15, pady=(0, 10))
        
        self.btn_create_template = ctk.CTkButton(
            excel_frame, 
            text="📋 Tạo Template Excel",
            height=40,
            fg_color="#10b981",
            hover_color="#059669",
            font=("Segoe UI", 11, "bold"),
            command=self.create_excel_template
        )
        self.btn_create_template.pack(fill="x", pady=(0, 5))
        
        self.btn_import_excel = ctk.CTkButton(
            excel_frame,
            text="📥 Import File (Excel/Word)",
            height=40,
            fg_color="#3b82f6",
            hover_color="#1d4ed8",
            font=("Segoe UI", 11, "bold"),
            command=self.import_links_file
        )
        self.btn_import_excel.pack(fill="x")
        
        # Input Area (Textbox)
        self.textbox_fb_links = ctk.CTkTextbox(left_col, font=("Consolas", 11))
        self.textbox_fb_links.pack(fill="both", expand=True, padx=15, pady=5)
        self.textbox_fb_links.insert("1.0", "# Paste danh sách link vào đây (mỗi dòng 1 link)...\n")
        self.textbox_fb_links.bind("<KeyRelease>", self.update_link_count_display)
        
        # Actions
        actions = ctk.CTkFrame(left_col, fg_color="transparent")
        actions.pack(fill="x", padx=15, pady=15)
        
        self.lbl_link_count = ctk.CTkLabel(actions, text="(0 link)", text_color=self.colors['success'], font=("Segoe UI", 12, "bold"))
        self.lbl_link_count.pack(anchor="w")
        
        self.chk_auto_title_fb = ctk.CTkCheckBox(actions, text="Auto-Title (Đặt tiêu đề tự động)", font=("Segoe UI", 12))
        self.chk_auto_title_fb.pack(anchor="w", pady=5)
        self.chk_auto_title_fb.select()
        
        # NEW: Headless mode option
        self.chk_headless_scan = ctk.CTkCheckBox(actions, text="Chạy ẩn", font=("Segoe UI", 12))
        self.chk_headless_scan.pack(anchor="w", pady=5)
        self.chk_headless_scan.select()  # Default ON
        
        # NEW: Disable API image fetching
        self.chk_disable_image_api = ctk.CTkCheckBox(actions, text="Tắt Ảnh Oto ", font=("Segoe UI", 12))
        self.chk_disable_image_api.pack(anchor="w", pady=5)
        
        # Hint
        ctk.CTkLabel(actions, text="💡 Nếu máy chậm, hãy TẮT Headless", font=("Segoe UI", 10), text_color="gray").pack(anchor="w", pady=(0, 10))
        
        self.btn_import_fb = ctk.CTkButton(actions, text="🚀 PHÂN TÍCH & LẤY EMBED", 
                                          height=50, 
                                          fg_color=self.colors['warning'], 
                                          hover_color="#d97706",
                                          font=("Segoe UI", 13, "bold"),
                                          command=self.import_fb_bulk)
        self.btn_import_fb.pack(fill="x", pady=(10, 0))


        # RIGHT COLUMN (Results) - Remaining space
        right_col = ctk.CTkFrame(grid, fg_color=self.colors['bg_card'], corner_radius=12, border_width=1, border_color=self.colors['border'])
        right_col.pack(side="right", fill="both", expand=True, padx=(5, 0))
        
        # Header
        head = ctk.CTkFrame(right_col, fg_color="transparent")
        head.pack(fill="x", padx=15, pady=15)
        ctk.CTkLabel(head, text="📊 Kết quả Phân Tích (Chờ xử lý)", font=("Segoe UI", 16, "bold")).pack(side="left")
        
        # Controls
        ctrl = ctk.CTkFrame(head, fg_color="transparent")
        ctrl.pack(side="right")
        
        self.chk_select_all = ctk.CTkCheckBox(ctrl, text="Chọn Tất Cả", font=("Segoe UI", 12), command=self.toggle_select_all_results)
        self.chk_select_all.pack(side="left", padx=10)
        
        # NEW: Export Excel Button
        ctk.CTkButton(ctrl, text="📥 XUẤT EXCEL", 
                     width=140, height=35,
                     font=("Segoe UI", 11, "bold"),
                     fg_color="#10b981",
                     hover_color="#059669",
                     command=self.export_scan_results_to_excel).pack(side="left", padx=5)
        
        ctk.CTkButton(ctrl, text="⬇ THÊM VÀO HÀNG CHỜ", 
                     width=160, height=35,
                     font=("Segoe UI", 12, "bold"),
                     fg_color=self.colors['success'], 
                     hover_color="#059669",
                     command=self.add_selected_results_to_queue).pack(side="left")
        
        # List
        self.results_scroll = ctk.CTkScrollableFrame(right_col, fg_color="transparent") 
        self.results_scroll.pack(fill="both", expand=True, padx=5, pady=5)


    # =========================================================================
    # TAB 2: BATCH & QUEUE (Hợp nhất Logic Hàng chờ vào đây)
    # =========================================================================
    def create_batch_tab_content(self):
        # Chia làm 2 cột: Trái (Công cụ Import) - Phải (Danh sách Hàng Chờ)
        grid = ctk.CTkFrame(self.tab_batch, fg_color="transparent")
        grid.pack(fill="both", expand=True, padx=5, pady=5)
        
        # --- LEFT COLUMN: TOOLS ---
        left_col = ctk.CTkFrame(grid, width=350)
        left_col.pack(side="left", fill="both", padx=5, expand=False)
        
        ctk.CTkLabel(left_col, text="📂 Nhập Dữ Liệu", font=("Segoe UI", 14, "bold")).pack(pady=10)
        
        # CSV Section
        ctk.CTkLabel(left_col, text="Nhập từ CSV:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10)
        self.entry_csv = ctk.CTkEntry(left_col, placeholder_text="Chưa chọn file CSV...")
        self.entry_csv.pack(fill="x", padx=10, pady=5)
        
        btn_row_csv = ctk.CTkFrame(left_col, fg_color="transparent")
        btn_row_csv.pack(fill="x", padx=10)
        ctk.CTkButton(btn_row_csv, text="Chọn File", width=80, command=self.browse_csv).pack(side="left", padx=2)
        ctk.CTkButton(btn_row_csv, text="Tải Mẫu", width=80, fg_color="gray", command=self.create_example_csv).pack(side="right", padx=2)

        # TXT Folder Section
        ctk.CTkLabel(left_col, text="Nhập từ Folder TXT:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10, pady=(20, 0))
        ctk.CTkButton(left_col, text="📂 Chọn Folder TXT", command=self.import_txt_folder).pack(fill="x", padx=10, pady=5)

        # Facebook Tools Section
        ctk.CTkLabel(left_col, text="Công cụ Facebook:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10, pady=(20, 0))
        ctk.CTkLabel(left_col, text="💡 Nhập nhiều link FB trong tab 'Facebook Scan'", 
                    font=("Segoe UI", 10), 
                    text_color="gray",
                    wraplength=300).pack(anchor="w", padx=10, pady=(5, 0))
        ctk.CTkButton(left_col, text="📱 Mở Tab Facebook Scan", 
                     fg_color="#1877F2", 
                     command=lambda: self.tabview.set("📱 Facebook Scan")).pack(fill="x", padx=10, pady=5)

        # --- RIGHT COLUMN: QUEUE LIST ---
        right_col = ctk.CTkFrame(grid)
        right_col.pack(side="right", fill="both", expand=True, padx=5)

        # Queue Header
        q_header = ctk.CTkFrame(right_col, height=40, fg_color="transparent")
        q_header.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(q_header, text="📋 Danh sách Hàng Chờ", font=("Segoe UI", 14, "bold")).pack(side="left")
        self.lbl_queue_count = ctk.CTkLabel(q_header, text="(0 bài)", text_color="gray")
        self.lbl_queue_count.pack(side="left", padx=5)
        
        # Copy All Links button
        ctk.CTkButton(
            q_header, 
            text="📋 Copy Tất Cả Link", 
            width=130, 
            fg_color=self.colors['primary'],
            hover_color=self.colors['primary_hover'],
            command=self.copy_all_links_with_titles
        ).pack(side="right", padx=5)
        
        ctk.CTkButton(q_header, text="🗑️ Xóa All", width=80, fg_color=self.colors['danger'], command=self.clear_queue).pack(side="right")

        # Listbox replacement with Scrollable Frame for Cards
        self.queue_scroll = ctk.CTkScrollableFrame(right_col, label_text="Danh sách Video đang chờ")
        self.queue_scroll.pack(fill="both", expand=True, padx=10, pady=5)
        # self.queue_listbox = ctk.CTkTextbox(right_col, font=("Consolas", 12)) # Removed
        # self.queue_listbox.pack(fill="both", expand=True, padx=10, pady=5)
        # self.queue_listbox.configure(state="disabled")

        # Action Button (Run Batch)
        self.btn_batch_post = ctk.CTkButton(right_col, text="▶️ CHẠY AUTO POST (HÀNG CHỜ)", height=50, 
                                            fg_color=self.colors['success'], 
                                            font=("Segoe UI", 13, "bold"),
                                            state="disabled",
                                            command=self.on_queue_post_click)
        self.btn_batch_post.pack(fill="x", padx=10, pady=10)

    # =========================================================================
    # TAB 3: UPLOAD VIDEO
    # =========================================================================
    def create_upload_tab_content(self):
        container = ctk.CTkFrame(self.tab_upload, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Upload Options
        options_frame = ctk.CTkFrame(container)
        options_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(options_frame, text="Tùy chọn Upload:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10, pady=5)
        
        self.chk_headless_upload = ctk.CTkCheckBox(options_frame, text="🚀 Chạy ẩn (Nhanh hơn)", font=("Segoe UI", 11))
        self.chk_headless_upload.pack(anchor="w", padx=20, pady=2)
        self.chk_headless_upload.select()  # Default to headless for speed
        
        self.chk_auto_add_queue = ctk.CTkCheckBox(options_frame, text="📝 Tự động thêm vào hàng chờ đăng bài", font=("Segoe UI", 11))
        self.chk_auto_add_queue.pack(anchor="w", padx=20, pady=2)
        self.chk_auto_add_queue.select()  # Default to auto-add
        
        # Input Area
        input_frame = ctk.CTkFrame(container)
        input_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(input_frame, text="Chọn Video Upload:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=10, pady=5)
        
        row = ctk.CTkFrame(input_frame, fg_color="transparent")
        row.pack(fill="x", padx=10, pady=5)
        
        self.entry_upload_path = ctk.CTkEntry(row, placeholder_text="Đường dẫn file video...")
        self.entry_upload_path.pack(side="left", fill="x", expand=True)
        ctk.CTkButton(row, text="📂 Chọn File", width=100, command=self.browse_video_upload).pack(side="left", padx=5)
        
        self.btn_upload = ctk.CTkButton(input_frame, text="⬆️ Bắt đầu Upload", fg_color=self.colors['primary'], command=self.on_upload_click)
        self.btn_upload.pack(fill="x", padx=10, pady=10)

        # Upload Log/Result Area
        ctk.CTkLabel(container, text="Kết quả Upload (Link nhúng sẽ hiện ở đây):", font=("Segoe UI", 12, "bold")).pack(anchor="w", pady=(10,0))
        self.txt_upload_list = ctk.CTkTextbox(container, font=("Consolas", 11))
        self.txt_upload_list.pack(fill="both", expand=True, pady=5)

    # =========================================================================
    # TAB 4: VIMEO TOOLS
    # =========================================================================
    def create_vimeo_tab_content(self):
        frm = ctk.CTkFrame(self.tab_vimeo)
        frm.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(frm, text="🎥 Công cụ Vimeo", font=("Segoe UI", 16, "bold")).pack(pady=10)
        
        # Cookie Login Test Section
        login_frame = ctk.CTkFrame(frm)
        login_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(login_frame, text="🔐 Test Cookie Login", font=("Segoe UI", 14, "bold")).pack(pady=5)
        
        self.chk_headless_test = ctk.CTkCheckBox(login_frame, text="🚀 Chạy ẩn (Nhanh hơn)", font=("Segoe UI", 11))
        self.chk_headless_test.pack(anchor="w", padx=20, pady=2)
        self.chk_headless_test.select()  # Default to headless
        
        self.btn_test_cookie = ctk.CTkButton(login_frame, text="🍪 Test Cookie Login", 
                                           height=40, fg_color=self.colors['success'], 
                                           command=self.test_vimeo_cookie_login)
        self.btn_test_cookie.pack(fill="x", padx=20, pady=5)
        
        # Account Creation Section
        reg_frame = ctk.CTkFrame(frm)
        reg_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(reg_frame, text="📝 Tạo Tài Khoản Vimeo Hàng Loạt", font=("Segoe UI", 14, "bold")).pack(pady=5)
        
        self.entry_vm_count = self.create_modern_input(reg_frame, "Số lượng tài khoản cần tạo:", "10")
        
        # Vimeo Options
        self.chk_headless_vimeo = ctk.CTkCheckBox(reg_frame, text="🚀 Chạy ẩn (Nhanh hơn)", font=("Segoe UI", 11))
        self.chk_headless_vimeo.pack(anchor="w", padx=30, pady=5)
        self.chk_headless_vimeo.select()  # Default to headless
        
        self.btn_vm_reg = ctk.CTkButton(reg_frame, text="🚀 Bắt đầu Tạo", height=50, 
                                      fg_color=self.colors['warning'], 
                                      command=self.on_vimeo_reg_click)
        self.btn_vm_reg.pack(fill="x", padx=30, pady=20)

        # Account List Section
        list_frame = ctk.CTkFrame(frm)
        list_frame.pack(fill="both", expand=True, pady=10)
        
        header_frame = ctk.CTkFrame(list_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(header_frame, text="📋 Danh sách Tài Khoản", font=("Segoe UI", 14, "bold")).pack(side="left")
        ctk.CTkButton(header_frame, text="🔄 Refresh", width=80, height=25, command=self.refresh_account_list).pack(side="right")
        
        # Scrollable list
        self.acc_scroll = ctk.CTkScrollableFrame(list_frame, label_text="Click 'Login' để mở trình duyệt")
        self.acc_scroll.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Initial load
        self.after(500, self.refresh_account_list)

        # Log Section
        ctk.CTkLabel(frm, text="📜 Nhật ký chạy:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=30, pady=(10,0))
        self.txt_vimeo_log = ctk.CTkTextbox(frm, font=("Consolas", 10), height=300)
        self.txt_vimeo_log.pack(fill="both", expand=True, padx=30, pady=10)
        self.txt_vimeo_log.configure(state="disabled")

    def refresh_account_list(self):
        """Read vimeo_accounts.txt and populate the scrollable list"""
        # Clear existing
        try:
            for widget in self.acc_scroll.winfo_children():
                widget.destroy()
        except:
            pass # Widget might not be ready
            
        try:
            if os.path.exists("vimeo_accounts.txt"):
                with open("vimeo_accounts.txt", "r", encoding="utf-8") as f:
                    lines = [l.strip() for l in f.readlines() if l.strip()]
                
                if not lines:
                    ctk.CTkLabel(self.acc_scroll, text="File tài khoản trống").pack(pady=10)
                    return

                for idx, line in enumerate(lines):
                    if "|" in line:
                        parts = line.split("|")
                        if len(parts) >= 2:
                            email = parts[0].strip()
                            pwd = parts[1].strip()
                            
                            row = ctk.CTkFrame(self.acc_scroll)
                            row.pack(fill="x", pady=2, padx=2)
                            
                            # Login Button (Pack FIRST to ensure visibility)
                            btn = ctk.CTkButton(row, text="Login 🚀", width=80, height=24,
                                              fg_color=self.colors['primary'],
                                              command=lambda e=email, p=pwd: self.start_vimeo_session(e, p))
                            btn.pack(side="right", padx=(5, 10), pady=2)
                            
                            # Copy Password
                            btn_pass = ctk.CTkButton(row, text="Copy MK", width=60, height=24,
                                                   fg_color="#4b5563", # Gray
                                                   font=("Segoe UI", 11))
                            btn_pass.configure(command=lambda t=pwd, b=btn_pass: self.copy_account_info(t, b))
                            btn_pass.pack(side="right", padx=2, pady=2)
                            
                            # Copy Check (Copy TK)
                            btn_user = ctk.CTkButton(row, text="Copy TK", width=60, height=24,
                                                   fg_color="#4b5563", # Gray 
                                                   font=("Segoe UI", 11))
                            btn_user.configure(command=lambda t=email, b=btn_user: self.copy_account_info(t, b))
                            btn_user.pack(side="right", padx=2, pady=2)

                            # Info (Pack SECOND to take remaining space)
                            info_text = f"{idx+1}. {email} | {pwd}"
                            ctk.CTkLabel(row, text=info_text, font=("Consolas", 12), anchor="w").pack(side="left", padx=10, fill="x", expand=True)
            else:
                ctk.CTkLabel(self.acc_scroll, text="Không tìm thấy file vimeo_accounts.txt\nTạo file này với định dạng: email|password").pack(pady=20)
                
        except Exception as e:
            print(f"Error loading accounts: {e}")

    def copy_account_info(self, text, btn_widget):
        """Helper to copy text and animate button"""
        try:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.update() # Required to finalize clipboard
            
            orig_text = btn_widget.cget("text")
            orig_color = btn_widget.cget("fg_color")
            
            btn_widget.configure(text="✅", fg_color=self.colors['success'])
            
            def restore():
                try:
                    btn_widget.configure(text=orig_text, fg_color=orig_color)
                except: pass # Widget might be destroyed
                
            self.after(1000, restore)
        except Exception as e:
            print(f"Copy error: {e}")

    def start_vimeo_session(self, email, pwd):
        """Launch browser and auto-login"""
        import threading
        from model.vimeo_helper import VimeoHelper
        
        if hasattr(self, 'txt_vimeo_log'):
            self.txt_vimeo_log.configure(state="normal")
            self.txt_vimeo_log.insert("end", f"\n[SYSTEM] 🚀 Đang mở trình duyệt cho: {email}...\n")
            self.txt_vimeo_log.see("end")
            self.txt_vimeo_log.configure(state="disabled")
            
        def _run():
            helper = VimeoHelper()
            try:
                helper.login_interactive(email, pwd)
            except Exception as e:
                print(f"Login error: {e}")
                
        # Run in thread so GUI doesn't freeze
        threading.Thread(target=_run, daemon=True).start()

    # =========================================================================
    # TAB 4.5: IMAGE MANAGEMENT (Quản lý ảnh API car)
    # =========================================================================
    def create_images_tab_content(self):
        """Tab to manage car API images - save and delete"""
        container = ctk.CTkScrollableFrame(self.tab_images, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Header
        header_frame = ctk.CTkFrame(container, fg_color="transparent")
        header_frame.pack(fill="x", pady=(0, 10))
        
        ctk.CTkLabel(header_frame, text="Ảnh Đã Up Lên Web", font=("Segoe UI", 18, "bold")).pack(side="left")
        ctk.CTkButton(header_frame, text="🔄 Làm Mới", width=100, command=self.refresh_image_lists).pack(side="right", padx=5)
        
        # Description
        ctk.CTkLabel(container, 
                    text="💡 Ảnh xe sẽ tự động lưu vào thư viện sau khi upload lên WordPress thành công",
                    font=("Segoe UI", 11),
                    text_color="gray",
                    wraplength=800).pack(anchor="w", pady=(0, 20))
        
        # ===== SECTION: SAVED LIBRARY (Thư viện ảnh xe từ API) =====
        saved_section = ctk.CTkFrame(container)
        saved_section.pack(fill="both", expand=True)
        
        ctk.CTkLabel(saved_section, text=" Thư Viện Ảnh ", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=20, pady=(10, 5))
        ctk.CTkLabel(saved_section, 
                    text="Ảnh xe đã upload lên WordPress. Tự động lưu sau mỗi lần đăng bài thành công.",
                    font=("Segoe UI", 10),
                    text_color="gray").pack(anchor="w", padx=20, pady=(0, 10))
        
        # Saved list with scrollbar
        saved_list_frame = ctk.CTkFrame(saved_section)
        saved_list_frame.pack(fill="both", expand=True, padx=20, pady=(0, 15))
        
        self.saved_listbox = tk.Listbox(
            saved_list_frame,
            font=("Consolas", 10),
            bg="#2b2b2b",
            fg="white",
            selectbackground="#10b981",
            height=8
        )
        self.saved_listbox.pack(side="left", fill="both", expand=True)
        
        saved_scrollbar = ctk.CTkScrollbar(saved_list_frame, command=self.saved_listbox.yview)
        saved_scrollbar.pack(side="right", fill="y")
        self.saved_listbox.configure(yscrollcommand=saved_scrollbar.set)
        
        # Saved actions
        saved_actions = ctk.CTkFrame(saved_section, fg_color="transparent")
        saved_actions.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkButton(saved_actions, text="📋 Copy Đường Dẫn", 
                     fg_color=self.colors['primary'], 
                     width=150,
                     command=self.copy_saved_image_path).pack(side="left", padx=5)
        
        ctk.CTkButton(saved_actions, text="🗑️ Xoá Khỏi Thư Viện", 
                     fg_color=self.colors['danger'], 
                     width=150,
                     command=self.delete_selected_saved).pack(side="left", padx=5)
        
        ctk.CTkButton(saved_actions, text="👁️ Xem Ảnh", 
                     fg_color="gray", 
                     width=100,
                     command=self.view_selected_saved).pack(side="left", padx=5)
        
        ctk.CTkButton(saved_actions, text="📂 Mở Thư Mục", 
                     fg_color="gray", 
                     width=120,
                     command=lambda: self.open_folder("saved_car_images")).pack(side="left", padx=5)
        
        self.lbl_saved_count = ctk.CTkLabel(saved_actions, text="(0 ảnh)", text_color="gray")
        self.lbl_saved_count.pack(side="right", padx=10)
        
        # Load initial data
        self.refresh_image_lists()
    
    # =========================================================================
    # TAB 4.6: AI THUMBNAIL SETTINGS (Cài đặt AI cho Thumbnail)
    # =========================================================================
    def create_thumbnail_ai_tab_content(self):
        """Tab for AI Thumbnail customization"""
        from model.thumbnail_config import ThumbnailConfig
        
        # Load config
        self.thumb_config = ThumbnailConfig()
        
        container = ctk.CTkScrollableFrame(self.tab_thumbnail_ai, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Header
        header = ctk.CTkFrame(container, fg_color="transparent")
        header.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(header, text="🤖 Cài Đặt AI Thumbnail", 
                    font=("Segoe UI", 22, "bold")).pack(anchor="w")
        ctk.CTkLabel(header, 
                    text="Tùy chỉnh cách AI xử lý ảnh thumbnail để tối ưu cho Facebook OG Image",
                    font=("Segoe UI", 12),
                    text_color="gray").pack(anchor="w", pady=(5, 0))
        
        # === SECTION 1: OUTPUT RESOLUTION ===
        res_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                               corner_radius=12, border_width=1, border_color=self.colors['border'])
        res_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(res_card, text="📐 Độ Phân Giải Đầu Ra", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        self.resolution_var = ctk.StringVar(value=self.thumb_config.get("output_resolution", "720p"))
        
        res_options = ctk.CTkFrame(res_card, fg_color="transparent")
        res_options.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkRadioButton(res_options, text="720p (1280x720) - Khuyến nghị ✅", 
                          variable=self.resolution_var, value="720p",
                          font=("Segoe UI", 12)).pack(anchor="w", pady=5)
        ctk.CTkLabel(res_options, text="   → Đủ cho Facebook, không quá xử lý, tự nhiên nhất",
                    font=("Segoe UI", 10), text_color="gray").pack(anchor="w")
        
        ctk.CTkRadioButton(res_options, text="1080p (1920x1080) - Chất lượng cao", 
                          variable=self.resolution_var, value="1080p",
                          font=("Segoe UI", 12)).pack(anchor="w", pady=5)
        ctk.CTkLabel(res_options, text="   → Rất nét nhưng có thể làm ảnh trông giả nếu nguồn kém",
                    font=("Segoe UI", 10), text_color="gray").pack(anchor="w")
        
        # === SECTION 2: AI UPSCALE ===
        ai_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                              corner_radius=12, border_width=1, border_color=self.colors['border'])
        ai_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(ai_card, text="🚀 AI Upscaling (Real-ESRGAN)", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        self.ai_upscale_var = ctk.BooleanVar(value=self.thumb_config.get("use_ai_upscale", False))
        
        ctk.CTkCheckBox(ai_card, text="Bật AI Upscale cho ảnh nhỏ", 
                       variable=self.ai_upscale_var,
                       font=("Segoe UI", 13, "bold"),
                       command=self.on_ai_upscale_toggle).pack(anchor="w", padx=20, pady=(0, 10))
        
        ctk.CTkLabel(ai_card, 
                    text="⚠️ Chỉ dùng khi ảnh gốc quá nhỏ (< 360p). AI sẽ tăng độ phân giải mà không làm vỡ ảnh.",
                    font=("Segoe UI", 11), text_color="orange", wraplength=700).pack(anchor="w", padx=20, pady=(0, 10))
        
        # AI Threshold slider
        threshold_frame = ctk.CTkFrame(ai_card, fg_color="transparent")
        threshold_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkLabel(threshold_frame, text="Ngưỡng kích hoạt AI (pixel):", 
                    font=("Segoe UI", 12)).pack(anchor="w")
        
        self.ai_threshold_var = ctk.IntVar(value=self.thumb_config.get("ai_upscale_threshold", 360))
        self.ai_threshold_slider = ctk.CTkSlider(threshold_frame, from_=240, to=720, 
                                                 variable=self.ai_threshold_var,
                                                 command=self.on_threshold_change)
        self.ai_threshold_slider.pack(fill="x", pady=5)
        
        self.ai_threshold_label = ctk.CTkLabel(threshold_frame, text=f"Dùng AI khi chiều cao < {self.ai_threshold_var.get()}px",
                                              font=("Segoe UI", 10), text_color="gray")
        self.ai_threshold_label.pack(anchor="w")
        
        # === SECTION 3: FACE RESTORATION ===
        face_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                                corner_radius=12, border_width=1, border_color=self.colors['border'])
        face_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(face_card, text="👤 Phục Hồi Khuôn Mặt (GFPGAN)", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        self.face_restore_var = ctk.BooleanVar(value=self.thumb_config.get("use_face_restoration", False))
        
        ctk.CTkCheckBox(face_card, text="Tự động làm nét khuôn mặt trong bodycam", 
                       variable=self.face_restore_var,
                       font=("Segoe UI", 13, "bold")).pack(anchor="w", padx=20, pady=(0, 10))
        
        ctk.CTkLabel(face_card, 
                    text="💡 Hữu ích cho video bodycam có người. AI sẽ phục hồi chi tiết khuôn mặt bị mờ.",
                    font=("Segoe UI", 11), text_color="gray", wraplength=700).pack(anchor="w", padx=20, pady=(0, 15))
        
        # === SECTION 4: SHARPENING ===
        sharp_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                                 corner_radius=12, border_width=1, border_color=self.colors['border'])
        sharp_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(sharp_card, text="✨ Làm Nét (Sharpening)", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        sharp_frame = ctk.CTkFrame(sharp_card, fg_color="transparent")
        sharp_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkLabel(sharp_frame, text="Cường độ làm nét:", 
                    font=("Segoe UI", 12)).pack(anchor="w")
        
        self.sharpen_var = ctk.DoubleVar(value=self.thumb_config.get("sharpen_strength", 0.5))
        self.sharpen_slider = ctk.CTkSlider(sharp_frame, from_=0.0, to=1.0, 
                                           variable=self.sharpen_var,
                                           command=self.on_sharpen_change)
        self.sharpen_slider.pack(fill="x", pady=5)
        
        self.sharpen_label = ctk.CTkLabel(sharp_frame, text=self._get_sharpen_desc(0.5),
                                         font=("Segoe UI", 10), text_color="gray")
        self.sharpen_label.pack(anchor="w")
        
        # === SECTION 5: COLOR CORRECTION ===
        color_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                                 corner_radius=12, border_width=1, border_color=self.colors['border'])
        color_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(color_card, text="🎨 Hiệu Chỉnh Màu Sắc", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        self.color_correct_var = ctk.BooleanVar(value=self.thumb_config.get("color_correction", True))
        
        ctk.CTkCheckBox(color_card, text="Tự động khử màu cam/vàng (bodycam)", 
                       variable=self.color_correct_var,
                       font=("Segoe UI", 13, "bold")).pack(anchor="w", padx=20, pady=(0, 10))
        
        ctk.CTkLabel(color_card, 
                    text="💡 Bodycam thường có màu cam/vàng. Bật tùy chọn này để ảnh trông tự nhiên hơn.",
                    font=("Segoe UI", 11), text_color="gray", wraplength=700).pack(anchor="w", padx=20, pady=(0, 15))
        
        # === SECTION 6: OUTPUT QUALITY ===
        quality_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                                   corner_radius=12, border_width=1, border_color=self.colors['border'])
        quality_card.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(quality_card, text="💾 Chất Lượng Lưu File", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        quality_frame = ctk.CTkFrame(quality_card, fg_color="transparent")
        quality_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        ctk.CTkLabel(quality_frame, text="JPEG Quality (%):", 
                    font=("Segoe UI", 12)).pack(anchor="w")
        
        self.quality_var = ctk.IntVar(value=self.thumb_config.get("output_quality", 92))
        self.quality_slider = ctk.CTkSlider(quality_frame, from_=70, to=100, 
                                           variable=self.quality_var,
                                           command=self.on_quality_change)
        self.quality_slider.pack(fill="x", pady=5)
        
        self.quality_label = ctk.CTkLabel(quality_frame, text=f"Quality: {self.quality_var.get()}% (Sweet spot: 92%)",
                                         font=("Segoe UI", 10), text_color="gray")
        self.quality_label.pack(anchor="w")
        
        # === PREVIEW SECTION ===
        preview_card = ctk.CTkFrame(container, fg_color=self.colors['bg_card'], 
                                   corner_radius=12, border_width=1, border_color=self.colors['border'])
        preview_card.pack(fill="both", expand=True, pady=(0, 15))
        
        ctk.CTkLabel(preview_card, text="🖼️ Xem Trước & Test", 
                    font=("Segoe UI", 16, "bold")).pack(anchor="w", padx=20, pady=(15, 10))
        
        ctk.CTkLabel(preview_card, 
                    text="Upload ảnh mẫu để test các cài đặt AI trước khi áp dụng",
                    font=("Segoe UI", 11), text_color="gray").pack(anchor="w", padx=20, pady=(0, 10))
        
        # Upload button
        upload_frame = ctk.CTkFrame(preview_card, fg_color="transparent")
        upload_frame.pack(fill="x", padx=20, pady=(0, 10))
        
        ctk.CTkButton(upload_frame, text="📂 Chọn Ảnh Test", 
                     width=150,
                     height=35,
                     font=("Segoe UI", 12, "bold"),
                     fg_color=self.colors['primary'],
                     command=self.upload_test_thumbnail).pack(side="left", padx=5)
        
        self.test_image_label = ctk.CTkLabel(upload_frame, text="Chưa có ảnh", 
                                            font=("Segoe UI", 11), text_color="gray")
        self.test_image_label.pack(side="left", padx=10)
        
        # Before/After comparison
        comparison_frame = ctk.CTkFrame(preview_card, fg_color=self.colors['bg_dark'])
        comparison_frame.pack(fill="both", expand=True, padx=20, pady=(0, 15))
        
        # Before image
        before_container = ctk.CTkFrame(comparison_frame, fg_color="transparent")
        before_container.pack(side="left", fill="both", expand=True, padx=5, pady=10)
        
        ctk.CTkLabel(before_container, text="📷 Ảnh Gốc", 
                    font=("Segoe UI", 13, "bold")).pack(pady=(0, 5))
        
        self.before_image_frame = ctk.CTkFrame(before_container, 
                                              width=300, height=200,
                                              fg_color="gray20")
        self.before_image_frame.pack(fill="both", expand=True)
        
        self.before_image_label = ctk.CTkLabel(self.before_image_frame, 
                                              text="Chưa có ảnh\n\nBấm 'Chọn Ảnh Test' để upload",
                                              font=("Segoe UI", 11),
                                              text_color="gray")
        self.before_image_label.pack(expand=True)
        
        # After image
        after_container = ctk.CTkFrame(comparison_frame, fg_color="transparent")
        after_container.pack(side="right", fill="both", expand=True, padx=5, pady=10)
        
        ctk.CTkLabel(after_container, text="✨ Sau Khi Xử Lý", 
                    font=("Segoe UI", 13, "bold")).pack(pady=(0, 5))
        
        self.after_image_frame = ctk.CTkFrame(after_container, 
                                             width=300, height=200,
                                             fg_color="gray20")
        self.after_image_frame.pack(fill="both", expand=True)
        
        self.after_image_label = ctk.CTkLabel(self.after_image_frame, 
                                             text="Chưa xử lý\n\nBấm 'Áp Dụng Hiệu Ứng' để test",
                                             font=("Segoe UI", 11),
                                             text_color="gray")
        self.after_image_label.pack(expand=True)
        
        # Processing info
        self.processing_info_label = ctk.CTkLabel(preview_card, 
                                                 text="",
                                                 font=("Segoe UI", 10),
                                                 text_color=self.colors['success'])
        self.processing_info_label.pack(padx=20, pady=(0, 10))
        
        # === ACTION BUTTONS ===
        actions = ctk.CTkFrame(container, fg_color="transparent")
        actions.pack(fill="x", pady=(20, 0))
        
        ctk.CTkButton(actions, text="✨ Áp Dụng Hiệu Ứng", 
                     height=45,
                     font=("Segoe UI", 14, "bold"),
                     fg_color="#8b5cf6",  # Purple
                     hover_color="#7c3aed",
                     command=self.apply_thumbnail_effects).pack(side="left", padx=5)
        
        ctk.CTkButton(actions, text="💾 Lưu Cài Đặt", 
                     height=45,
                     font=("Segoe UI", 14, "bold"),
                     fg_color=self.colors['success'],
                     hover_color=self.colors['success_hover'],
                     command=self.save_thumbnail_config).pack(side="left", padx=5)
        
        ctk.CTkButton(actions, text="🔄 Khôi Phục Mặc Định", 
                     height=45,
                     font=("Segoe UI", 14, "bold"),
                     fg_color=self.colors['warning'],
                     command=self.reset_thumbnail_config).pack(side="left", padx=5)
        
        ctk.CTkButton(actions, text="💾 Lưu Ảnh Đã Xử Lý", 
                     height=45,
                     font=("Segoe UI", 14, "bold"),
                     fg_color="#10b981",
                     command=self.save_processed_thumbnail).pack(side="left", padx=5)
        
        # Store test image path
        self.test_thumbnail_path = None
        self.processed_thumbnail_path = None
    
    def _get_sharpen_desc(self, value):
        """Get description for sharpen strength"""
        if value < 0.2:
            return "Rất nhẹ (Gần như không làm nét)"
        elif value < 0.4:
            return "Nhẹ (Tự nhiên, khuyến nghị)"
        elif value < 0.6:
            return "Trung bình (Cân bằng)"
        elif value < 0.8:
            return "Mạnh (Rất nét, có thể hơi giả)"
        else:
            return "Rất mạnh (Cực nét, có thể quá xử lý)"
    
    def on_ai_upscale_toggle(self):
        """Handle AI upscale checkbox toggle"""
        enabled = self.ai_upscale_var.get()
        self.ai_threshold_slider.configure(state="normal" if enabled else "disabled")
    
    def on_threshold_change(self, value):
        """Update threshold label"""
        self.ai_threshold_label.configure(text=f"Dùng AI khi chiều cao < {int(value)}px")
    
    def on_sharpen_change(self, value):
        """Update sharpen label"""
        self.sharpen_label.configure(text=self._get_sharpen_desc(value))
    
    def on_quality_change(self, value):
        """Update quality label"""
        self.quality_label.configure(text=f"Quality: {int(value)}% (Sweet spot: 92%)")
    
    def save_thumbnail_config(self):
        """Save thumbnail AI configuration"""
        try:
            self.thumb_config.set("output_resolution", self.resolution_var.get())
            self.thumb_config.set("use_ai_upscale", self.ai_upscale_var.get())
            self.thumb_config.set("ai_upscale_threshold", self.ai_threshold_var.get())
            self.thumb_config.set("use_face_restoration", self.face_restore_var.get())
            self.thumb_config.set("sharpen_strength", self.sharpen_var.get())
            self.thumb_config.set("color_correction", self.color_correct_var.get())
            self.thumb_config.set("output_quality", self.quality_var.get())
            
            if self.thumb_config.save():
                messagebox.showinfo("Thành Công", "✅ Đã lưu cài đặt AI Thumbnail!")
                self.log("💾 Đã lưu cài đặt AI Thumbnail")
            else:
                messagebox.showerror("Lỗi", "❌ Không thể lưu cài đặt!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"❌ Lỗi lưu cài đặt: {e}")
    
    def reset_thumbnail_config(self):
        """Reset to default settings"""
        if messagebox.askyesno("Xác Nhận", "Khôi phục về cài đặt mặc định?"):
            self.thumb_config.reset_to_defaults()
            
            # Update UI
            self.resolution_var.set(self.thumb_config.get("output_resolution"))
            self.ai_upscale_var.set(self.thumb_config.get("use_ai_upscale"))
            self.ai_threshold_var.set(self.thumb_config.get("ai_upscale_threshold"))
            self.face_restore_var.set(self.thumb_config.get("use_face_restoration"))
            self.sharpen_var.set(self.thumb_config.get("sharpen_strength"))
            self.color_correct_var.set(self.thumb_config.get("color_correction"))
            self.quality_var.set(self.thumb_config.get("output_quality"))
            
            messagebox.showinfo("Thành Công", "✅ Đã khôi phục cài đặt mặc định!")
            self.log("🔄 Đã khôi phục cài đặt AI Thumbnail về mặc định")
    
    def upload_test_thumbnail(self):
        """Upload a test thumbnail image"""
        file_path = filedialog.askopenfilename(
            title="Chọn ảnh test",
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.webp"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                self.test_thumbnail_path = file_path
                filename = os.path.basename(file_path)
                self.test_image_label.configure(text=f"✅ {filename}")
                
                # Display before image
                self._display_image(file_path, self.before_image_label, max_size=(400, 300))
                
                # Clear after image
                self.after_image_label.configure(image=None, text="Chưa xử lý\n\nBấm 'Áp Dụng Hiệu Ứng' để test")
                self.processed_thumbnail_path = None
                self.processing_info_label.configure(text="")
                
                self.log(f"📂 Đã load ảnh test: {filename}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể load ảnh: {e}")
    
    def _display_image(self, image_path, label_widget, max_size=(400, 300)):
        """Display image in a label with max size"""
        try:
            from PIL import Image, ImageTk
            
            img = Image.open(image_path)
            
            # Calculate resize ratio
            width_ratio = max_size[0] / img.width
            height_ratio = max_size[1] / img.height
            ratio = min(width_ratio, height_ratio, 1.0)  # Don't upscale
            
            new_width = int(img.width * ratio)
            new_height = int(img.height * ratio)
            
            img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(img_resized)
            
            label_widget.configure(image=photo, text="")
            label_widget.image = photo  # Keep reference
            
        except Exception as e:
            print(f"Error displaying image: {e}")
            label_widget.configure(text=f"Lỗi hiển thị:\n{str(e)}")
    
    def apply_thumbnail_effects(self):
        """Apply AI effects to test thumbnail with current settings"""
        if not self.test_thumbnail_path:
            messagebox.showwarning("Chưa có ảnh", "Vui lòng chọn ảnh test trước!")
            return
        
        try:
            import threading
            
            def process():
                try:
                    self.after(0, lambda: self.processing_info_label.configure(
                        text="⏳ Đang xử lý... Vui lòng đợi", text_color="orange"))
                    
                    from model.facebook_thumbnail_optimizer import FacebookThumbnailOptimizerUltra
                    
                    # Create optimizer
                    optimizer = FacebookThumbnailOptimizerUltra()
                    
                    # Override settings from UI
                    optimizer.OUTPUT_WIDTH = 1920 if self.resolution_var.get() == "1080p" else 1280
                    optimizer.OUTPUT_HEIGHT = 1080 if self.resolution_var.get() == "1080p" else 720
                    
                    # Process image
                    import time
                    start_time = time.time()
                    
                    output_filename = f"test_preview_{int(time.time())}.jpg"
                    result_path = optimizer.optimize_for_facebook(
                        self.test_thumbnail_path,
                        output_filename=output_filename,
                        enhance=True
                    )
                    
                    processing_time = time.time() - start_time
                    
                    if result_path and os.path.exists(result_path):
                        self.processed_thumbnail_path = result_path
                        
                        # Display after image
                        self.after(0, lambda: self._display_image(result_path, self.after_image_label, max_size=(400, 300)))
                        
                        # Show info
                        file_size = os.path.getsize(result_path) / 1024
                        info_text = f"✅ Xử lý xong trong {processing_time:.1f}s | Kích thước: {file_size:.0f} KB"
                        self.after(0, lambda: self.processing_info_label.configure(
                            text=info_text, text_color=self.colors['success']))
                        
                        self.after(0, lambda: self.log(f"✨ Đã xử lý ảnh test thành công"))
                    else:
                        self.after(0, lambda: messagebox.showerror("Lỗi", "Không thể xử lý ảnh!"))
                        self.after(0, lambda: self.processing_info_label.configure(
                            text="❌ Lỗi xử lý", text_color="red"))
                        
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    self.after(0, lambda: messagebox.showerror("Lỗi", f"Lỗi xử lý ảnh: {e}"))
                    self.after(0, lambda: self.processing_info_label.configure(
                        text=f"❌ Lỗi: {str(e)}", text_color="red"))
            
            # Run in thread
            threading.Thread(target=process, daemon=True).start()
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể xử lý: {e}")
    
    def save_processed_thumbnail(self):
        """Save the processed thumbnail to a custom location"""
        if not self.processed_thumbnail_path or not os.path.exists(self.processed_thumbnail_path):
            messagebox.showwarning("Chưa có ảnh", "Vui lòng xử lý ảnh test trước!")
            return
        
        try:
            save_path = filedialog.asksaveasfilename(
                title="Lưu ảnh đã xử lý",
                defaultextension=".jpg",
                filetypes=[("JPEG Image", "*.jpg"), ("All files", "*.*")],
                initialfile=f"thumbnail_processed_{int(time.time())}.jpg"
            )
            
            if save_path:
                import shutil
                shutil.copy2(self.processed_thumbnail_path, save_path)
                messagebox.showinfo("Thành Công", f"✅ Đã lưu ảnh:\n{save_path}")
                self.log(f"💾 Đã lưu ảnh đã xử lý: {os.path.basename(save_path)}")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu ảnh: {e}")
    
    
    def refresh_image_lists(self):
        """Refresh saved car images from API"""
        try:
            from model.image_api import ImageAPI
            image_api = ImageAPI()
            
            # Get saved images
            saved_images = image_api.get_saved_images()
            self.saved_listbox.delete(0, tk.END)
            for img_path in saved_images:
                filename = os.path.basename(img_path)
                size_kb = os.path.getsize(img_path) // 1024
                self.saved_listbox.insert(tk.END, f"{filename} ({size_kb} KB)")
            
            self.lbl_saved_count.configure(text=f"({len(saved_images)} ảnh)")
            
            self.log(f"🔄 Đã làm mới thư viện: {len(saved_images)} ảnh xe từ API")
            
        except Exception as e:
            self.log(f"❌ Lỗi làm mới danh sách ảnh: {e}")
    
    
    def delete_selected_saved(self):
        """Delete selected saved image"""
        try:
            selection = self.saved_listbox.curselection()
            if not selection:
                messagebox.showwarning("Chưa chọn ảnh", "Vui lòng chọn ảnh cần xoá!")
                return
            
            # Confirm deletion
            if not messagebox.askyesno("Xác nhận", "Bạn có chắc muốn xoá ảnh này khỏi thư viện?"):
                return
            
            from model.image_api import ImageAPI
            image_api = ImageAPI()
            
            # Get selected image path
            saved_images = image_api.get_saved_images()
            selected_idx = selection[0]
            
            if selected_idx >= len(saved_images):
                messagebox.showerror("Lỗi", "Không tìm thấy ảnh!")
                return
            
            image_path = saved_images[selected_idx]
            
            # Delete image
            if image_api.delete_image(image_path):
                self.log(f"🗑️ Đã xoá ảnh khỏi thư viện: {os.path.basename(image_path)}")
                messagebox.showinfo("Thành công", "Đã xoá ảnh khỏi thư viện!")
                self.refresh_image_lists()
            else:
                messagebox.showerror("Lỗi", "Không thể xoá ảnh!")
                
        except Exception as e:
            self.log(f"❌ Lỗi xoá ảnh: {e}")
            messagebox.showerror("Lỗi", f"Lỗi xoá ảnh: {e}")
    
    
    def view_selected_saved(self):
        """Open selected saved image in default image viewer"""
        try:
            selection = self.saved_listbox.curselection()
            if not selection:
                messagebox.showwarning("Chưa chọn ảnh", "Vui lòng chọn ảnh cần xem!")
                return
            
            from model.image_api import ImageAPI
            image_api = ImageAPI()
            
            saved_images = image_api.get_saved_images()
            selected_idx = selection[0]
            
            if selected_idx >= len(saved_images):
                return
            
            image_path = saved_images[selected_idx]
            
            # Open with default viewer
            import subprocess
            subprocess.Popen(['start', '', image_path], shell=True)
            
        except Exception as e:
            self.log(f"❌ Lỗi mở ảnh: {e}")
    
    def copy_saved_image_path(self):
        """Copy selected saved image path to clipboard"""
        try:
            selection = self.saved_listbox.curselection()
            if not selection:
                messagebox.showwarning("Chưa chọn ảnh", "Vui lòng chọn ảnh!")
                return
            
            from model.image_api import ImageAPI
            image_api = ImageAPI()
            
            saved_images = image_api.get_saved_images()
            selected_idx = selection[0]
            
            if selected_idx >= len(saved_images):
                return
            
            image_path = saved_images[selected_idx]
            
            # Copy to clipboard
            self.clipboard_clear()
            self.clipboard_append(os.path.abspath(image_path))
            
            self.log(f"📋 Đã copy đường dẫn: {image_path}")
            messagebox.showinfo("Thành công", f"Đã copy đường dẫn!\n{os.path.abspath(image_path)}")
            
        except Exception as e:
            self.log(f"❌ Lỗi copy đường dẫn: {e}")
    
    def open_folder(self, folder_name):
        """Open folder in file explorer"""
        try:
            if not os.path.exists(folder_name):
                os.makedirs(folder_name)
            
            import subprocess
            subprocess.Popen(['explorer', os.path.abspath(folder_name)])
            
        except Exception as e:
            self.log(f"❌ Lỗi mở thư mục: {e}")

    # =========================================================================
    # TAB 5: DATA & LOGS (Thay thế cho Bottom UI cũ)
    # =========================================================================
    def create_data_tab_content(self):
        tabview_logs = ctk.CTkTabview(self.tab_data)
        tabview_logs.pack(fill="both", expand=True, padx=5, pady=5)
        
        tab_sys_log = tabview_logs.add("🖥️ System Log")
        tab_history = tabview_logs.add("✅ Lịch sử Link")
        
        # System Log
        self.textbox_log = ctk.CTkTextbox(tab_sys_log, font=("Consolas", 11))
        self.textbox_log.pack(fill="both", expand=True)
        self.textbox_log.configure(state="disabled")
        
        # History
        hist_toolbar = ctk.CTkFrame(tab_history, height=40)
        hist_toolbar.pack(fill="x", pady=5)
        
        # Left side - Import
        ctk.CTkButton(hist_toolbar, text="📥 Import File", command=self.import_csv_for_export, fg_color="#8b5cf6").pack(side="left", padx=10)
        
        # Right side - Export & Actions
        ctk.CTkButton(hist_toolbar, text="📋 Copy Tất Cả Link", command=self.copy_history_links).pack(side="right", padx=10)
        # Smart Export Button - exports in same format as imported
        ctk.CTkButton(hist_toolbar, text="📄 Xuất File", command=self.smart_export_history, fg_color="#3b82f6").pack(side="right", padx=10)
        # Export Excel Button (Green)
        ctk.CTkButton(hist_toolbar, text="📊 Xuất Excel", command=self.export_history_to_excel, fg_color="#28a745").pack(side="right", padx=10)
        # Open Documents Folder Button (Orange)
        ctk.CTkButton(hist_toolbar, text="📂 Mở Documents", command=self.open_documents_folder, fg_color="#f59e0b").pack(side="right", padx=10)
        ctk.CTkButton(hist_toolbar, text="🗑️ Xóa Lịch Sử", command=self.clear_history, fg_color="red").pack(side="right", padx=10)
        
        # Use tkinter Text widget instead of CTkTextbox for better link handling
        import tkinter as tk
        history_frame = ctk.CTkFrame(tab_history)
        history_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        self.history_textbox = tk.Text(
            history_frame, 
            font=("Segoe UI", 11),
            bg="#2b2b2b",
            fg="white",
            insertbackground="white",
            wrap="word",
            cursor="arrow"
        )
        self.history_textbox.pack(fill="both", expand=True)
        self.history_textbox.configure(state="disabled")
        
        # Configure tag for clickable links
        self.history_textbox.tag_config("link", foreground="#3b82f6", underline=True)
        self.history_textbox.tag_bind("link", "<Button-1>", self._on_link_click)
        self.history_textbox.tag_bind("link", "<Enter>", lambda e: self.history_textbox.config(cursor="hand2"))
        self.history_textbox.tag_bind("link", "<Leave>", lambda e: self.history_textbox.config(cursor="arrow"))

    # Method removed (using the one defined later which correctly clears self.published_links)

    # =========================================================================
    # =========================================================================
    # TAB 6: SETTINGS
    # =========================================================================
    def create_settings_tab_content(self):
        frm = ctk.CTkFrame(self.tab_settings)
        frm.pack(fill="both", expand=True, padx=20, pady=20)
        
        # --- UI Settings ---
        ctk.CTkLabel(frm, text="Cài đặt giao diện", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=20, pady=10)
        ctk.CTkOptionMenu(frm, values=["Dark", "Light", "System"], command=lambda v: ctk.set_appearance_mode(v)).pack(padx=20, anchor="w")
        
        # --- Facebook Cookie Settings ---
        ctk.CTkLabel(frm, text="Facebook Cookies (Fix lỗi Reel/Video)", font=("Segoe UI", 14, "bold")).pack(anchor="w", padx=20, pady=(30, 10))
        
        cookie_frame = ctk.CTkFrame(frm, fg_color="transparent")
        cookie_frame.pack(anchor="w", padx=20, fill="x")
        
        # Status Label
        self.lbl_fb_cookie_status = ctk.CTkLabel(cookie_frame, text="Checking...", font=("Segoe UI", 12))
        self.lbl_fb_cookie_status.pack(side="left", padx=(0, 15))
        
        # Import Button
        ctk.CTkButton(
            cookie_frame, 
            text="📂 Chọn File Cookie (.txt)", 
            command=self.browse_fb_cookie,
            fg_color="#3b82f6",
            width=150
        ).pack(side="left")
        
        # Paste Button (NEW)
        ctk.CTkButton(
            cookie_frame, 
            text="📋 Paste từ Clipboard", 
            command=self.paste_fb_cookie_from_clipboard,
            fg_color="#8b5cf6", # Purple
            width=150
        ).pack(side="left", padx=10)
        
        # Help Text
        ctk.CTkLabel(
            frm, 
            text="💡 Hướng dẫn: Dùng extension 'Get cookies.txt LOCALLY' -> Copy -> Bấm nút Paste ở trên.",
            text_color="gray",
            font=("Segoe UI", 11)
        ).pack(anchor="w", padx=20, pady=5)
        
        # Initial check
        self.check_fb_cookie_status()
        
        ctk.CTkLabel(frm, text="Thông tin phiên bản: v2.0 Refactored", text_color="gray").pack(side="bottom", pady=20)

    def check_fb_cookie_status(self):
        if hasattr(self, 'lbl_fb_cookie_status'):
             if os.path.exists("facebook_cookies.txt"):
                 self.lbl_fb_cookie_status.configure(text="✅ Đã có cookie", text_color="#4ade80") # Green
             else:
                 self.lbl_fb_cookie_status.configure(text="❌ Chưa có cookie", text_color="#ef4444") # Red

    def browse_fb_cookie(self):
        filename = filedialog.askopenfilename(title="Chọn file cookie Facebook", filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if filename:
            try:
                import shutil
                shutil.copy2(filename, "facebook_cookies.txt")
                self.log(f"✅ Đã import cookie từ: {os.path.basename(filename)}")
                self.check_fb_cookie_status()
                messagebox.showinfo("Thành công", "Đã nạp Facebook Cookie thành công!")
            except Exception as e:
                self.log(f"❌ Lỗi copy cookie: {e}")
                
    def paste_fb_cookie_from_clipboard(self):
        try:
            content = self.clipboard_get()
            if not content or len(content) < 10:
                messagebox.showwarning("Clipboard Trống", "Vui lòng copy nội dung cookie trước!")
                return
            
            # Simple check for netscape format
            if ".facebook.com" not in content and "c_user" not in content:
                res = messagebox.askyesno("Cảnh báo", "Nội dung không giống cookie Facebook. Bạn có chắc muốn lưu?")
                if not res: return
                
            with open("facebook_cookies.txt", "w", encoding="utf-8") as f:
                f.write(content)
                
            self.check_fb_cookie_status()
            self.log("✅ Đã paste cookie từ clipboard.")
            messagebox.showinfo("Thành công", "Đã lưu nội dung cookie!")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lấy nội dung clipboard: {e}")

    # =========================================================================
    # LOGIC HELPERS & EVENT HANDLERS
    # =========================================================================

    def log(self, message):
        """Unified logging method - Optimized with Batching & UI Buffer 🚀"""
        # 1. Store in buffer for periodic update (Batching)
        if not hasattr(self, '_log_buffer'):
            self._log_buffer = []
            
        # Optional: Filter out spam logs if needed (e.g. "Link OK - Status 200")
        msg_str = str(message)
        timestamp = time.strftime("[%H:%M:%S] ")
        
        # Only buffer if not redundant (simple dedup check could be added here)
        self._log_buffer.append(timestamp + msg_str + "\n")
        
        # 2. Schedule flush if not already scheduled
        if not hasattr(self, '_log_flush_scheduled') or not self._log_flush_scheduled:
            self._log_flush_scheduled = True
            self.after(500, self._flush_logs)  # Update UI every 500ms
        
        # 3. Print to console for real-time debugging (optional)
        try:
             print(msg_str)
        except: pass

    def _flush_logs(self):
        """Flush log buffer to UI in one go"""
        self._log_flush_scheduled = False
        if not hasattr(self, '_log_buffer') or not self._log_buffer:
            return
            
        try:
             # Get messages and clear buffer
             # Use a local copy to avoid race conditions if thread appends while we read
             current_buffer = list(self._log_buffer)
             self._log_buffer = []  # Clear immediately
             
             messages_to_write = "".join(current_buffer)
             
             # Update Log Tab (Efficiently)
             if hasattr(self, 'textbox_log'):
                 self.textbox_log.configure(state="normal")
                 
                 # Optimization: Truncate if too long (keep last 50KB)
                 current_len = len(self.textbox_log.get("1.0", "end"))
                 if current_len > 50000:
                      self.textbox_log.delete("1.0", "200.0")
                 
                 self.textbox_log.insert("end", messages_to_write)
                 self.textbox_log.see("end")
                 self.textbox_log.configure(state="disabled")
             
             # Update Status Bar (Last message only)
             # DISABLE OLD STATUS BAR UPDATE (User Request) - Because we have new Kitty Loading
             # if hasattr(self, 'status_label') and current_buffer:
             #     last_msg = current_buffer[-1].strip()
             #     # Clean timestamp for status bar
             #     if "] " in last_msg: 
             #          last_msg = last_msg.split('] ')[-1]
             #     self.status_label.configure(text=last_msg[:60])
                 
        except Exception as e:
             print(f"Log Flush Error: {e}")

    def browse_csv(self):
        path = filedialog.askopenfilename(filetypes=[("CSV", "*.csv")])
        if path:
            self.entry_csv.delete(0, "end")
            self.entry_csv.insert(0, path)
            # Load data logic here
            try:
                from model.batch_helper import BatchPostData
                self.batch_data = BatchPostData(path)
                self.log(f"Đã load {self.batch_data.total_posts()} bài từ CSV.")
                self.btn_batch_post.configure(state="normal")
            except:
                self.log("Lỗi load CSV (Giả lập: File OK)")
                self.btn_batch_post.configure(state="normal")

    def create_example_csv(self):
        path = filedialog.asksaveasfilename(defaultextension=".csv", initialfile="mau_bai_viet.csv")
        if path:
            with open(path, "w", encoding="utf-8-sig") as f:
                f.write("title,video_url,image_url,content\nBai 1,http://video,,Noi dung 1")
            self.log(f"Đã tạo file mẫu: {path}")

    def import_txt_folder(self):
        """Import content from TXT files to be used as body content for video posts"""
        folder = filedialog.askdirectory(title="Chọn folder chứa file TXT (nội dung body)")
        if not folder: 
            return
        
        try:
            # Get all .txt files
            files = [f for f in os.listdir(folder) if f.endswith(".txt")]
            
            if not files:
                self.log("❌ Không tìm thấy file .txt nào trong folder!")
                return
            
            # Store content in a pool for later use
            if not hasattr(self, 'content_pool'):
                self.content_pool = []
            
            imported_count = 0
            for filename in files:
                try:
                    file_path = os.path.join(folder, filename)
                    
                    # Read file content
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                    
                    # Skip empty files
                    if not content:
                        self.log(f"⚠️ Bỏ qua file trống: {filename}")
                        continue
                    
                    # Add to content pool WITH file path for later deletion
                    self.content_pool.append({
                        'filename': filename,
                        'filepath': file_path,  # Store full path for deletion
                        'content': content
                    })
                    
                    imported_count += 1
                    self.log(f"✅ Đã đọc: {filename} ({len(content)} ký tự)")
                    
                except Exception as e:
                    self.log(f"❌ Lỗi đọc file {filename}: {e}")
                    continue
            
            self.log(f"🎉 Đã load {imported_count} file nội dung vào bộ nhớ.")
            self.log(f"💡 Tip: Upload video để tự động ghép với nội dung này!")
            
        except Exception as e:
            self.log(f"❌ Lỗi import folder: {e}")

    def add_result_card(self, idx, title, embed_code, image_url, analysis_text, url="", video_id=""):
        """Add a result card to the scrollable list"""
        try:
            # Card Frame
            card = ctk.CTkFrame(self.results_scroll, fg_color="#2b2b2b", corner_radius=8, border_width=1, border_color="#404040")
            card.pack(fill="x", pady=5, padx=5)
            
            # STORE DATA ON CARD FOR BULK ACTIONS (including URL and video_id for export)
            card.data = {
                'url': url,
                'title': title,
                'embed_code': embed_code,
                'thumbnail': image_url,
                'video_id': video_id,
                'status': 'scanned'
            }
            
            # Checkbox (Left)
            card.checkbox = ctk.CTkCheckBox(card, text="", width=24, checkbox_width=20, checkbox_height=20)
            card.checkbox.pack(side="left", padx=(10, 0))
            # Auto-select if face detected (Optional smart feature? maybe just leave unchecked)
            # card.checkbox.select() 
            
            # --- Layout: Image Left, Info Right ---
            # Image Placeholder/Canvas
            img_frame = ctk.CTkFrame(card, width=60, height=60, fg_color="#1a1a1a")
            img_frame.pack(side="left", padx=5, pady=2)
            img_frame.pack_propagate(False) # Fixed size
            
            # Load Image (cached)
            if image_url and os.path.exists(image_url):
                try:
                    pil_img = self._load_image_cached(image_url, (60, 60))
                    if pil_img:
                        ctk_img = ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=(60, 60))
                        ctk.CTkLabel(img_frame, text="", image=ctk_img).pack(expand=True)
                    else:
                        ctk.CTkLabel(img_frame, text="No Image").pack(expand=True)
                except:
                    ctk.CTkLabel(img_frame, text="No Image").pack(expand=True)
            else:
                ctk.CTkLabel(img_frame, text="No Image").pack(expand=True)

            # Info Frame
            info_frame = ctk.CTkFrame(card, fg_color="transparent")
            info_frame.pack(side="left", fill="both", expand=True, padx=5, pady=2)
            
            # Title
            # Title Row: Number label + Editable title entry
            title_row = ctk.CTkFrame(info_frame, fg_color="transparent")
            title_row.pack(fill="x")
            ctk.CTkLabel(title_row, text=f"{idx}.", font=("Segoe UI", 11, "bold"), text_color="white", width=25).pack(side="left")
            card.title_entry = ctk.CTkEntry(title_row, font=("Segoe UI", 11, "bold"), text_color="white", fg_color="#374151", border_width=1, border_color="#555555", height=28, placeholder_text="Nhap tieu de...")
            card.title_entry.pack(side="left", fill="x", expand=True, padx=(3, 0))
            if title:
                card.title_entry.insert(0, title)
            
            # AI Analysis Result
            ai_color = "#9ca3af" # default light gray for neutral/error
            if "Phát hiện" in analysis_text or "Face Detected" in analysis_text: ai_color = "#4ade80" # Screen Green (Brighter)
            elif "No Image" in analysis_text or "Failed" in analysis_text: ai_color = "#ef4444" # Red for obvious error
            
            ctk.CTkLabel(info_frame, text=f"🧠 AI: {analysis_text}", font=("Segoe UI", 10), text_color=ai_color, anchor="w").pack(fill="x")
            
            # Embed Code (Entry for easy copying)
            embed_entry = ctk.CTkEntry(info_frame, font=("Consolas", 10), height=22, text_color="white", fg_color="#374151")
            embed_entry.pack(fill="x", pady=2)
            embed_entry.insert(0, embed_code)
            
            # Add to Queue Button
            ctk.CTkButton(info_frame, text="⬇ Queue", 
                         width=80, height=22, 
                         fg_color="#3b82f6", font=("Segoe UI", 10),
                         command=lambda c=card: self.add_single_result_to_queue(c)).pack(anchor="e", pady=2)
            
        except Exception as e:
            print(f"Error adding card: {e}")

    def toggle_select_all_results(self):
        """Select or deselect all cards"""
        state = self.chk_select_all.get()
        for widget in self.results_scroll.winfo_children():
            if hasattr(widget, 'checkbox'):
                if state: widget.checkbox.select()
                else: widget.checkbox.deselect()

    def add_single_result_to_queue(self, card, update_ui=True):
        """Add a single card's data to the queue"""
        try:
            if hasattr(card, 'data'):
                data = card.data
            else:
                # Fallback if card is not the widget storing data directly
                # This depends on how the card was created
                return False

            # Create post data
            # Use manually edited title if available, otherwise use auto-detected title
            manual_title = ""
            if hasattr(card, 'title_entry'):
                manual_title = card.title_entry.get().strip()
            auto_title = data.get('title', 'Video')
            final_title = manual_title if manual_title else auto_title
            
            # FIXED: Use 'thumbnail' key (matches add_result_card)
            thumbnail_path = data.get('thumbnail', '') or data.get('image_url', '')
            
            post_data = {
                'title': final_title,
                'video_url': data.get('embed_code', ''),
                'image_url': thumbnail_path,  # Use thumbnail path
                'content': '',
                'needs_body_content': True,
                'theme': self.get_selected_theme_id()
            }
            # Add to list
            self.post_queue.append(post_data)
            
            # Update UI only if requested
            if update_ui:
                self.update_queue_display()
                self.log(f"✔ Da them '{final_title[:30]}...' vao hang cho.")
            return True
        except Exception as e:
            self.log(f"❌ Lỗi thêm bài: {e}")
            return False

    def add_selected_results_to_queue(self):
        """Add all checked cards to queue - Optimized for bulk adding"""
        added_count = 0
        cards_to_add = []
        
        # 1. Collect all selected cards first (fast)
        for widget in self.results_scroll.winfo_children():
            # Check for custom 'checkbox' attribute we added
            if hasattr(widget, 'checkbox') and widget.checkbox.get() == 1:
                cards_to_add.append(widget)
        
        if not cards_to_add:
            self.log("⚠️ Chưa chọn video nào!")
            return

        # 2. Add to queue data structure (fast, no UI updates)
        for card in cards_to_add:
            if self.add_single_result_to_queue(card, update_ui=False):
                added_count += 1
        
        # 3. Update UI ONCE at the end (prevents freezing)
        if added_count > 0:
            self.update_queue_display()
            self.log(f"🎉 Đã thêm {added_count} video đã chọn vào hàng chờ!")

    def update_link_count_display(self, event=None):
        """Update the link count label based on textbox content"""
        try:
            if not hasattr(self, 'textbox_fb_links') or not hasattr(self, 'lbl_link_count'):
                return
                
            text = self.textbox_fb_links.get("1.0", "end").strip()
            if not text:
                count = 0
            else:
                lines = text.split('\n')
                # Count non-empty non-comment lines
                count = sum(1 for line in lines if line.strip() and not line.strip().startswith('#'))
            
            self.lbl_link_count.configure(text=f"({count} link)")
        except:
            pass
    
    def create_excel_template(self):
        """Create empty Excel template for user to fill with links"""
        try:
            from model.excel_helper import ExcelHelper
            
            # Ask user where to save
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile="batch_videos_template.xlsx"
            )
            
            if not file_path:
                return
            
            self.log("📋 Đang tạo template Excel...")
            helper = ExcelHelper()
            result = helper.create_template_excel(file_path)
            
            if result:
                self.log(f"✅ Template Excel đã tạo: {os.path.basename(file_path)}")
                self.log("💡 Hướng dẫn:")
                self.log("   1. Mở file Excel")
                self.log("   2. Paste link video vào cột 'Watch full rescues'")
                self.log("   3. Lưu file")
                self.log("   4. Nhấn 'Import Excel' để scan tất cả link")
                
                # Ask if user wants to open file
                if messagebox.askyesno("Thành công", f"Template đã tạo!\n\nBạn muốn mở file ngay không?"):
                    os.startfile(file_path)
            else:
                messagebox.showerror("Lỗi", "Không thể tạo template Excel.\nHãy cài đặt: pip install openpyxl")
                
        except Exception as e:
            self.log(f"❌ Lỗi: {e}")
            messagebox.showerror("Lỗi", f"Lỗi tạo template: {e}")
    
    def import_excel_file(self):
        """Import video links from Excel file"""
        try:
            from model.excel_helper import ExcelHelper
            
            # Ask user to select file
            file_path = filedialog.askopenfilename(
                # Put "All files" first so dialog doesn't look empty
                # when the current folder has no .xlsx files.
                filetypes=[("All files", "*.*"), ("Excel files", "*.xlsx")],
                title="Chọn file Excel chứa danh sách link"
            )
            
            if not file_path:
                return
            
            self.log(f"📥 Đang import Excel: {os.path.basename(file_path)}")
            
            helper = ExcelHelper()
            videos, error = helper.read_excel(file_path)
            
            if error:
                messagebox.showerror("Lỗi", f"Lỗi đọc Excel: {error}")
                self.log(f"❌ {error}")
                return
            
            if not videos:
                messagebox.showwarning("Trống", "File Excel không có link nào")
                return
            
            # Clear textbox and add links
            self.textbox_fb_links.delete("1.0", "end")
            
            for video in videos:
                self.textbox_fb_links.insert("end", f"{video['url']}\n")
            
            self.log(f"✅ Đã import {len(videos)} link từ Excel")
            self.log("💡 Nhấn '🚀 PHÂN TÍCH & LẤY EMBED' để bắt đầu scan")
            
            # Update link count
            self.update_link_count_display()
            
        except Exception as e:
            self.log(f"❌ Lỗi: {e}")
            messagebox.showerror("Lỗi", f"Lỗi import Excel: {e}")

    def import_links_file(self):
        """Import links from Excel (.xlsx), Word (.docx), or CSV file"""
        try:
            from pathlib import Path
            from model.excel_helper import ExcelHelper
            from model.word_helper import WordHelper
            import csv

            file_path = filedialog.askopenfilename(
                filetypes=[
                    ("All files", "*.*"),
                    ("Excel files", "*.xlsx"),
                    ("CSV files", "*.csv"),
                    ("Word files", "*.docx"),
                ],
                title="Chọn file Excel/Word/CSV chứa danh sách link"
            )

            if not file_path:
                return

            ext = Path(file_path).suffix.lower()

            if ext == ".xlsx":
                self.log(f"📥 Đang import Excel: {os.path.basename(file_path)}")
                videos, error = ExcelHelper.read_excel(file_path)
                if error:
                    messagebox.showerror("Lỗi", f"Lỗi đọc Excel: {error}")
                    self.log(f"❌ {error}")
                    return
            elif ext == ".csv":
                self.log(f"📥 Đang import CSV: {os.path.basename(file_path)}")
                try:
                    videos = []
                    with open(file_path, 'r', encoding='utf-8-sig') as f:
                        reader = csv.reader(f)
                        header = next(reader, None)  # Skip header
                        
                        # Detect which column has links
                        link_col = 0  # Default to first column
                        if header:
                            # Check if header has "Link" or "URL" keyword
                            for i, col_name in enumerate(header):
                                if col_name and ('link' in col_name.lower() or 'url' in col_name.lower() or 'bài viết' in col_name.lower()):
                                    link_col = i
                                    break
                        
                        for row in reader:
                            if row and len(row) > link_col:
                                url = row[link_col].strip()
                                # Validate URL
                                if url and ('http' in url or 'vimeo' in url or 'facebook' in url or 'youtube' in url):
                                    videos.append({'url': url, 'title': 'Auto-generated'})
                    
                    if not videos:
                        messagebox.showwarning("Trống", "File CSV không có link nào")
                        return
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Lỗi đọc CSV: {e}")
                    self.log(f"❌ Lỗi đọc CSV: {e}")
                    return
            elif ext == ".docx":
                self.log(f"� Đang import Word: {os.path.basename(file_path)}")
                videos, error = WordHelper.read_docx(file_path)
                if error:
                    messagebox.showerror("Lỗi", f"Lỗi đọc Word: {error}")
                    self.log(f"❌ {error}")
                    return
            else:
                messagebox.showwarning("Không hỗ trợ", f"Chỉ hỗ trợ .xlsx, .csv và .docx\nFile bạn chọn: {ext}")
                return

            if not videos:
                messagebox.showwarning("Trống", "File không có link nào")
                return

            # Clear textbox and add links
            self.textbox_fb_links.delete("1.0", "end")
            for video in videos:
                self.textbox_fb_links.insert("end", f"{video['url']}\n")

            self.log(f"✅ Đã import {len(videos)} link từ file")
            self.log("💡 Nhấn '🚀 PHÂN TÍCH & LẤY EMBED' để bắt đầu scan")
            self.update_link_count_display()

        except Exception as e:
            self.log(f"❌ Lỗi: {e}")
            messagebox.showerror("Lỗi", f"Lỗi import file: {e}")

    def export_scan_results_to_excel(self):
        """Export scan results to Excel file"""
        try:
            from model.excel_helper import ExcelHelper
            
            # Collect all results from cards
            videos = []
            for widget in self.results_scroll.winfo_children():
                if hasattr(widget, 'data'):
                    videos.append(widget.data)
            
            if not videos:
                messagebox.showwarning("Trống", "Không có kết quả để export")
                return
            
            # Ask user where to save
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile="batch_videos_results.xlsx"
            )
            
            if not file_path:
                return
            
            self.log(f"📥 Đang export {len(videos)} video sang Excel...")
            
            helper = ExcelHelper()
            result = helper.export_results(videos, file_path)
            
            if result:
                self.log(f"✅ Đã export Excel: {os.path.basename(file_path)}")
                
                if messagebox.askyesno("Thành công", f"Đã export {len(videos)} video!\n\nBạn muốn mở file ngay không?"):
                    os.startfile(file_path)
            else:
                messagebox.showerror("Lỗi", "Không thể export Excel.\nHãy cài đặt: pip install openpyxl")
                
        except Exception as e:
            self.log(f"❌ Lỗi: {e}")
            messagebox.showerror("Lỗi", f"Lỗi export: {e}")

    def import_fb_bulk(self):
        """Import multiple video links, generate embeds, and SHOW RESULT in List"""
        try:
            # Get text from textbox
            fb_text = self.textbox_fb_links.get("1.0", "end").strip()
            
            if not fb_text:
                self.log("❌ Chưa nhập link Video nào!")
                return
            
            # Split by lines and filter out comments and empty lines
            lines = fb_text.split('\n')
            video_links = []  # List of tuples: (url, manual_title or None)
            
            # Support 2 formats:
            # Format 1 (links only): Each line is a URL
            # Format 2 (title + link): Line 1 = title, Line 2 = URL, repeat
            # Auto-detect: if a non-URL line is followed by a URL line, it's a title
            
            supported_domains = ['facebook.com', 'fb.watch', 'vimeo.com', 'youtube.com', 'youtu.be', 'tiktok.com']
            pending_title = None
            
            for line in lines:
                line = line.strip()
                if not line or line.startswith('#'): continue
                
                is_url = any(domain in line for domain in supported_domains)
                
                if is_url:
                    # This is a URL - check if previous line was a title
                    video_links.append((line, pending_title))
                    if pending_title:
                        print(f"[SCAN] Cặp: '{pending_title}' -> {line[:50]}")
                    pending_title = None  # Reset
                else:
                    # Not a URL - treat as title for the next URL
                    pending_title = line
            
            if not video_links:
                self.log("❌ Không tìm thấy link Video hợp lệ!")
                return
            
            # Count how many have manual titles
            manual_count = sum(1 for _, t in video_links if t)
            if manual_count > 0:
                self.log(f"📝 Phát hiện {manual_count}/{len(video_links)} link có tiêu đề thủ công")
            
            self.log(f"📱 Bắt đầu xử lý {len(video_links)} link Video (AI Analysis)...")
            self.btn_import_fb.configure(state="disabled", text="⏳ Đang phân tích...")
            
            # Clear previous results in scrollable frame
            for widget in self.results_scroll.winfo_children():
                widget.destroy()
            
            # ===== SHOW SMALL HELLO KITTY LOADING IN STATUS BAR =====
            # User request: Small pretty Hello Kitty near Kill Chrome button
            self.kitty_loading_frame.pack(side="left", padx=10, pady=8)
            self.kitty_status_label.configure(text=f"Đang xử lý {len(video_links)} videos... 💕")
            
            # Simple status updater
            def update_progress(current, total):
                try:
                    progress = int((current / total) * 100) if total > 0 else 0
                    messages = [
                        f"Đang phân tích {current}/{total} ({progress}%)... 💕",
                        f"Đang lấy dữ liệu {current}/{total} ({progress}%)... ✨",
                        f"Đang xử lý {current}/{total} ({progress}%)... 🌸"
                    ]
                    self.kitty_status_label.configure(text=messages[current % len(messages)])
                except: pass
            
            # Process in thread

            import threading
            import requests 
            from bs4 import BeautifulSoup
            import cv2
            import numpy as np
            import urllib.request

            def _get_meta(url):
                title = None
                img_url = None
                try:
                    headers = {'User-Agent': 'Mozilla/5.0'}
                    resp = requests.get(url, headers=headers, timeout=5)
                    if resp.status_code == 200:
                        soup = BeautifulSoup(resp.text, 'html.parser')
                        # Title
                        og_title = soup.find("meta", property="og:title")
                        if og_title: title = og_title.get("content")
                        if not title and soup.title: title = soup.title.string
                        # Image
                        og_image = soup.find("meta", property="og:image")
                        if og_image: img_url = og_image.get("content")
                except: pass
                return title, img_url
            
            def _analyze_face(img_path):
                """Simplified analysis - Skip OpenCV to avoid build issues and speed up processing"""
                try:
                    # Just check if file exists
                    if not os.path.exists(img_path):
                        return "No Image"
                    
                    # Get file size for basic validation
                    file_size = os.path.getsize(img_path)
                    if file_size < 1024:  # Less than 1KB
                        return "Invalid"
                    
                    return "Ready ✓"
                    
                except Exception as e:
                    return f"Error: {str(e)[:15]}"

            def _process_links():
                processed_count = 0
                shared_driver = None
                
                # Check for FB/TikTok links to init One Shared Driver (Complex platforms)
                has_complex = any(('facebook.com' in url or 'fb.watch' in url or 'tiktok.com' in url) for url, _ in video_links)
                
                if has_complex:
                    try:
                        self.log("🚀 Đang khởi động trình duyệt ẩn danh (Shared) để quét Facebook...")
                        import undetected_chromedriver as uc
                        options = uc.ChromeOptions()
                        
                        # Check headless option from GUI
                        use_headless = bool(self.chk_headless_scan.get()) if hasattr(self, 'chk_headless_scan') else True
                        
                        if use_headless:
                            self.log("   🔇 Chế độ: Headless (Ẩn)")
                        else:
                            self.log("   👁️ Chế độ: Visible (Hiện - Chậm hơn nhưng ổn định hơn)") 
                        
                        # Use Mobile User Agent
                        options.add_argument('--user-agent=Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1')
                        
                        options.add_argument("--disable-gpu")
                        options.add_argument("--mute-audio")
                        options.add_argument("--window-size=375,812") # Mobile dimensions
                        
                        chrome_path = os.path.join(os.getcwd(), "chrome_portable", "chrome.exe")
                        if os.path.exists(chrome_path): options.binary_location = chrome_path
                        
                        options.add_argument("--no-sandbox")
                        options.add_argument("--disable-dev-shm-usage")
                        
                        shared_driver = uc.Chrome(options=options, version_main=144, headless=use_headless)
                        shared_driver.set_page_load_timeout(30)
                    except Exception as e:
                        print(f"Failed to init shared driver: {e}")

                try:
                    for idx, (link, manual_title) in enumerate(video_links):
                        try:
                            title = manual_title  # Use manual title if provided (format 2)
                            img_remote = None
                            video_url_final = link
                            platform = "Video"

                            # 1. FACEBOOK
                            if 'facebook.com' in link or 'fb.watch' in link:
                                platform = "Facebook"
                                self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] FB: Đang lấy dữ liệu..."))
                                
                                # Only auto-detect title if no manual title provided
                                if not manual_title:
                                    # Use shared driver!
                                    title = self.get_facebook_title(link, driver=shared_driver)
                                else:
                                    # Manual title provided - still fetch image via driver
                                    self.get_facebook_title(link, driver=shared_driver)  # For image only
                                
                                # Get Image (Prefer cached from get_facebook_title, fallback to _get_meta)
                                img_remote = None
                                if hasattr(self, '_last_fb_image') and self._last_fb_image:
                                    img_remote = self._last_fb_image
                                else:
                                    _, img_remote = _get_meta(link)
                                
                                # Convert to embed (Iframe)
                                video_url_final = self.create_facebook_embed(link, use_sdk=False)
                            
                            # 2. YOUTUBE
                            elif 'youtube.com' in link or 'youtu.be' in link:
                                platform = "YouTube"
                                self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] YT: Đang lấy dữ liệu..."))
                                
                                # Get ID
                                import re
                                vid_id = None
                                if 'youtu.be' in link:
                                    match = re.search(r'youtu\.be/([\w-]+)', link)
                                    if match: vid_id = match.group(1)
                                else:
                                    match = re.search(r'v=([\w-]+)', link)
                                    if match: vid_id = match.group(1)
                                
                                if vid_id:
                                    # Embed Code
                                    video_url_final = f'<iframe width="560" height="315" src="https://www.youtube.com/embed/{vid_id}" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen></iframe>'
                                    
                                    # Thumbnail (High Quality)
                                    img_remote = f"https://img.youtube.com/vi/{vid_id}/maxresdefault.jpg"
                                    
                                    # Title via oEmbed
                                    try:
                                        oembed_url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={vid_id}&format=json"
                                        res = requests.get(oembed_url, timeout=3)
                                        if res.status_code == 200:
                                            title = res.json().get('title')
                                    except: pass
                                
                                if not title:
                                    title, _ = _get_meta(link) # Fallback

                            # 3. VIMEO
                            elif 'vimeo.com' in link:
                                platform = "Vimeo"
                                self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] Vimeo: Đang lấy dữ liệu..."))
                                
                                # Convert to Player URL / Embed
                                # Default to keeping it as is, but we will try to make it an embed if possible
                                video_url_final = link 
                                
                                # Extract ID
                                import re
                                vid_id = None
                                match = re.search(r'vimeo\.com/(\d+)', link)
                                if match:
                                    vid_id = match.group(1)
                                    # Use helper function for proper aspect ratio
                                    video_url_final = self.create_vimeo_embed(vid_id, title or "Vimeo Video")
                                    
                                    # Fetch oEmbed - dùng width=1280 để lấy thumbnail chất lượng cao
                                    try:
                                        oembed_url = f"https://vimeo.com/api/oembed.json?url=https://vimeo.com/{vid_id}&width=1280"
                                        res = requests.get(oembed_url, timeout=3)
                                        if res.status_code == 200:
                                            data = res.json()
                                            if not title: title = data.get('title')
                                            # oEmbed trả thumbnail_url kích thước đúng với width=1280
                                            img_remote = data.get('thumbnail_url')
                                    except: pass

                                    # Fallback: thử lấy qua Vimeo API /videos/{id} (không cần auth)
                                    if not img_remote:
                                        try:
                                            api_url = f"https://api.vimeo.com/videos/{vid_id}?fields=pictures"
                                            res2 = requests.get(api_url, timeout=3)
                                            if res2.status_code == 200:
                                                sizes = res2.json().get('pictures', {}).get('sizes', [])
                                                if sizes:
                                                    sizes.sort(key=lambda x: x.get('width', 0), reverse=True)
                                                    img_remote = sizes[0].get('link') or sizes[0].get('link_with_play_button')
                                        except: pass

                                if not title:
                                    title, img_remote_meta = _get_meta(link)
                                    if not img_remote: img_remote = img_remote_meta

                            # 4. TIKTOK (New)
                            elif 'tiktok.com' in link:
                                platform = "TikTok"
                                self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] TikTok: Đang lấy dữ liệu (API)..."))
                                
                                # 1. Try TikTok oEmbed API (Fastest & Most Reliable)
                                try:
                                    oembed_url = f"https://www.tiktok.com/oembed?url={link}"
                                    res = requests.get(oembed_url, timeout=5)
                                    if res.status_code == 200:
                                        data = res.json()
                                        title = data.get('title')
                                        img_remote = data.get('thumbnail_url')
                                        if title: print(f"[TikTok] oEmbed success: {title[:30]}...")
                                except Exception as e:
                                    print(f"[TikTok] oEmbed error: {e}")

                                # 2. Try yt-dlp (Fallback)
                                if (not title or not img_remote):
                                    self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] TikTok: Thử cách 2 (Deep Scan)..."))
                                    try:
                                        import yt_dlp
                                        ydl_opts = {
                                            'quiet': True, 
                                            'no_warnings': True,
                                            'extract_flat': True,
                                            'impersonate': 'chrome',
                                        }
                                        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                                            info = ydl.extract_info(link, download=False)
                                            if info:
                                                if not title: title = info.get('title')
                                                # Get best thumbnail
                                                if not img_remote:
                                                    thumbs = info.get('thumbnails', [])
                                                    if thumbs: img_remote = thumbs[-1].get('url')
                                                    else: img_remote = info.get('thumbnail')
                                    except: pass
                                
                                # 3. Fallback to Shared Driver (Last Resort)
                                if (not title or not img_remote) and shared_driver:
                                    try:
                                        shared_driver.get(link)
                                        time.sleep(3)
                                        if not title:
                                            t = shared_driver.title
                                            if t: title = t.replace("TikTok", "").replace("|", "").strip()
                                        if not img_remote:
                                            try:
                                                ie = shared_driver.find_element(uc.By.XPATH, "//meta[@property='og:image']")
                                                if ie: img_remote = ie.get_attribute("content")
                                            except: pass
                                    except: pass
                                
                                video_url_final = link # Let wp_model handle the embed block generation

                            # 5. OTHERS
                            else:
                                self.after(0, lambda i=idx: self.log(f"   🔍 [{i+1}] Đang lấy dữ liệu..."))
                                title, img_remote = _get_meta(link)
                                video_url_final = link

                            if not title: title = f"Video {idx+1}"
                            # Log if manual title was used
                            if manual_title:
                                print(f"[SCAN] Using manual title: '{manual_title}'")

                            # Download Image for Analysis
                            local_img_path = None
                            analysis_result = "No Image"
                            
                            if img_remote:
                                try:
                                    timestamp = time.strftime("%Y%m%d_%H%M%S")
                                    if not os.path.exists("temp_analysis"): os.makedirs("temp_analysis")
                                    local_img_path = f"temp_analysis/thumb_{timestamp}_{idx}.jpg"
                                    urllib.request.urlretrieve(img_remote, local_img_path)
                                    
                                    # Analyze
                                    analysis_result = _analyze_face(local_img_path)
                                except:
                                    analysis_result = "Download Failed"

                            # Add card to UI
                            # Extract video_id from link (for Vimeo: /video/123456789)
                            video_id = ""
                            if "vimeo.com" in link:
                                try:
                                    video_id = link.split("/")[-1].split("?")[0]
                                except:
                                    pass
                            
                            self.after(0, lambda i=idx+1, t=title, e=video_url_final, img=local_img_path, a=analysis_result, u=link, vid=video_id: 
                                       self.add_result_card(i, t, e, img, a, u, vid))
                            
                            processed_count += 1
                            
                            # Update progress
                            self.after(0, lambda c=processed_count, t=len(video_links): update_progress(c, t))
                            
                        except Exception as e:
                            print(e)
                            continue
                finally:
                    # Clean up shared driver
                    if shared_driver:
                        try:
                            shared_driver.quit()
                        except: pass
                
                # Stop emoji animation
                self._scanning_active = False
                
                # Hide status bar loading
                self.after(0, lambda: self.kitty_loading_frame.pack_forget())

                
                self.after(0, lambda: self.log(f"🏁 Hoàn tất phân tích {processed_count} link."))
                self.after(0, lambda: self.btn_import_fb.configure(state="normal", text="🚀 Phân Tích & Lấy Embed"))


            # Start thread
            threading.Thread(target=_process_links, daemon=True).start()

        except Exception as e:
            self.log(f"❌ Lỗi xử lý: {e}")
            self.btn_import_fb.configure(state="normal", text="🚀 Phân Tích & Lấy Embed")

    def get_facebook_title(self, fb_url, driver=None):
        """Get title from Facebook video page - Fast API method (bypass yt-dlp)"""
        fetched_title = None
        self._last_fb_image = None # Reset image cache
        
        # --- Method 1: Facebook Fast Fetcher (BYPASS YT-DLP) ---
        try:
            from model.facebook_fast_fetcher import FacebookFastFetcher
            
            # Initialize Facebook fast fetcher (singleton pattern)
            if not hasattr(self, '_fb_fast_fetcher'):
                self._fb_fast_fetcher = FacebookFastFetcher()
                print("[FB-FAST] ⚙️ Initialized Facebook Fast Fetcher (bypass yt-dlp)")
            
            # Get video info (skip yt-dlp by default)
            result = self._fb_fast_fetcher.get_video_info(fb_url, skip_ytdlp=True)
            
            if result['success'] and result['title']:
                fetched_title = result['title']
                self._last_fb_image = result['thumbnail']
                print(f"[FB] ✅ Fast Fetcher ({result['method']}): {fetched_title[:50]}...")
                return fetched_title
            else:
                print(f"[FB] ⚠️ Fast Fetcher failed: {result.get('error', 'Unknown')}, trying yt-dlp...")
                
        except ImportError:
            print("[FB] FacebookFastFetcher not available, trying yt-dlp...")
        except Exception as e:
            print(f"[FB] Fast Fetcher error: {e}, trying yt-dlp...")
        
        # --- Method 2: Enhanced yt-dlp (FALLBACK ONLY) ---
        try:
            from model.enhanced_ytdlp import EnhancedYTDLP
            
            # Initialize enhanced yt-dlp with optimized settings (singleton pattern)
            if not hasattr(self, '_enhanced_ytdlp'):
                self._enhanced_ytdlp = EnhancedYTDLP(
                    cookies_file="facebook_cookies.txt",
                    max_workers=5,        # Parallel workers
                    request_delay=0.5,    # 0.5s delay between requests
                    timeout=30            # 30s timeout (better for rate limiting)
                )
                print("[YTDLP+] ⚙️ Initialized with optimized settings for bulk processing")
            
            # Get video info with retry & cache
            result = self._enhanced_ytdlp.get_video_info(fb_url, use_cache=True, max_retries=3)
            
            if result['success'] and result['title']:
                fetched_title = result['title']
                self._last_fb_image = result['thumbnail']
                print(f"[FB] ✅ Enhanced yt-dlp: {fetched_title[:50]}...")
                return fetched_title
            else:
                print(f"[FB] ⚠️ Enhanced yt-dlp failed: {result.get('error', 'Unknown')}")
                
        except ImportError:
            print("[FB] Enhanced yt-dlp not available, trying standard yt-dlp...")
        except Exception as e:
            print(f"[FB] Enhanced yt-dlp error: {e}, trying fallback...")
        
        # --- Method 1.5: Standard yt-dlp (Fallback) ---
        try:
            import yt_dlp
            
            ydl_opts = {
                'quiet': True,
                'no_warnings': True,
                'simulate': True,
                'extract_flat': False,
                'noplaylist': True,
                'ignoreerrors': True,
                'socket_timeout': 10,
            }
            
            cookie_path = "facebook_cookies.txt"
            if os.path.exists(cookie_path):
                ydl_opts['cookiefile'] = cookie_path
                print(f"[FB] 🍪 Đã tìm thấy facebook_cookies.txt -> Đang sử dụng cookie!")
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(fb_url, download=False)
                
                if info:
                    title = info.get('title')
                    if title:
                        fetched_title = title.strip()
                        self._last_fb_image = info.get('thumbnail')
                        print(f"[FB] yt_dlp Lib extracted: {fetched_title[:50]}...")
                        return fetched_title

        except ImportError:
            print("[FB] Python module 'yt_dlp' lỗi import.")
        except Exception as e:
            print(f"[FB] yt-dlp chưa support video này, đang chuyển sang chạy Browser... ({e})")

        # --- Method 1.5: Try yt-dlp Subprocess (Fallback nếu không import được lib) ---
        try:
            import subprocess
            import json
            
            cmd = ['yt-dlp', '--dump-json', '--no-download', '--no-warnings', '--quiet', fb_url]
            
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=8, startupinfo=startupinfo)  # Giảm từ 15s -> 8s
            if result.returncode == 0 and result.stdout:
                data = json.loads(result.stdout)
                title = data.get('title') or data.get('fulltitle')
                if title and title.strip():
                    fetched_title = title.strip()
                    self._last_fb_image = data.get('thumbnail') 
                    print(f"[FB] yt-dlp.exe extracted: {fetched_title[:50]}...")
                    return fetched_title
        except: pass
        
        # --- Method 2: Requests (Only if no driver provided) ---
        if not driver:
            try:
                import requests
                from bs4 import BeautifulSoup
                
                headers = {
                    'User-Agent': 'facebookexternalhit/1.1 (+http://www.facebook.com/externalhit_uatext.php)',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
                }
                
                response = requests.get(fb_url, headers=headers, timeout=5, allow_redirects=True)
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    t = None
                    og_title = soup.find('meta', property='og:title')
                    if og_title: t = og_title.get('content')
                    
                    invalid_titles = ['facebook', 'watch', 'video', 'log into facebook', 'login', 'log in', 
                                     'this page isn', 'content not found', 'unavailable', 'facebook video']
                    
                    if t:
                        is_invalid = any(inv in t.lower() for inv in invalid_titles)
                        if not is_invalid:
                            fetched_title = t.strip()
                            og_image = soup.find('meta', property='og:image')
                            if og_image: self._last_fb_image = og_image.get('content')
            except: pass

        # --- Method 2: undetected_chromedriver (Fallback Mạnh Mẽ) ---
        invalid_titles = ["facebook video", "facebook", "log into facebook", "log in", "login", "watch",
                         "this page isn", "content not found", "page not found", "unavailable"]
        
        is_invalid_title = False
        if fetched_title:
             is_invalid_title = any(inv in fetched_title.lower() for inv in invalid_titles)
             
        if not fetched_title or is_invalid_title:
            print("[FB] Using undetected_chromedriver for title extraction...")
            
            # Helper for safe import
            try:
                import undetected_chromedriver as uc
            except ImportError:
                print("[FB] Could not import undetected_chromedriver")
                return "Facebook Video"

            use_driver = None
            should_quit = False
            
            # Reuse driver if possible
            if driver and isinstance(driver, uc.Chrome):
                use_driver = driver
                print("[FB] Reusing existing driver")
            else:
                try:
                    # import undetected_chromedriver as uc # Already imported above
                    options = uc.ChromeOptions()
                    
                    # Check headless option from GUI
                    use_headless = bool(self.chk_headless_scan.get()) if hasattr(self, 'chk_headless_scan') else True
                    
                    # Mobile UA
                    options.add_argument('--user-agent=Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1')
                    
                    options.add_argument("--disable-gpu")
                    options.add_argument("--mute-audio")
                    options.add_argument("--window-size=375,812") 
                    # Chống detect bot
                    options.add_argument("--disable-blink-features=AutomationControlled")
                    
                    chrome_path = os.path.join(os.getcwd(), "chrome_portable", "chrome.exe")
                    if os.path.exists(chrome_path): options.binary_location = chrome_path
                    
                    options.add_argument("--no-sandbox")
                    options.add_argument("--disable-dev-shm-usage")
                    use_driver = uc.Chrome(options=options, version_main=144, headless=use_headless)
                    use_driver.set_page_load_timeout(30)  # Giảm từ 45s -> 30s
                    should_quit = True
                    
                    # --- LOAD COOKIES INTO DRIVER ---
                    if os.path.exists(cookie_path):
                        try:
                            # Use mobile site for cookies too
                            use_driver.get("https://m.facebook.com/")
                            with open(cookie_path, 'r', encoding='utf-8') as f:
                                for line in f:
                                    if line.strip() and not line.strip().startswith('#'):
                                        parts = line.split('\t')
                                        if len(parts) >= 7:
                                            cookie = {
                                                'name': parts[5],
                                                'value': parts[6].strip(),
                                                'domain': parts[0],
                                                'path': parts[2]
                                            }
                                            try: use_driver.add_cookie(cookie)
                                            except: pass
                            print("[FB] 🍪 Đã nạp Cookies vào Browser Fallback!")
                            time.sleep(1) # Chờ cookie ăn
                        except Exception as e:
                            print(f"[FB] Lỗi nạp cookie: {e}")

                except Exception as e:
                    print(f"[FB] Failed to create driver: {e}")
                    return "Facebook Video" 

            try:
                use_driver.get(fb_url)
                
                # Scroll nhẹ
                try: use_driver.execute_script("window.scrollTo(0, 300);")
                except: pass
                
                # --- NEW: Dùng WebDriverWait thay vì sleep cứng ---
                from selenium.webdriver.common.by import By # Ensure imports
                from selenium.webdriver.support.ui import WebDriverWait
                from selenium.webdriver.support import expected_conditions as EC
                
                wait = WebDriverWait(use_driver, 5)  # Giảm từ 15s -> 5s
                
                # 1. Image Strategy (Wait for og:image)
                try:
                    # Chờ thẻ meta og:image xuất hiện
                    og_img_el = wait.until(EC.presence_of_element_located((By.XPATH, "//meta[@property='og:image']")))
                    if og_img_el:
                        img_url = og_img_el.get_attribute("content")
                        if img_url:
                            self._last_fb_image = img_url
                            print(f"[FB] Selenium found og:image: {img_url[:30]}...")
                except: 
                    # Last resort: Ảnh to nhất
                    if not self._last_fb_image:
                        try:
                            largest_img = use_driver.execute_script("""
                                let maxArea = 0; let bestSrc = null;
                                document.querySelectorAll('img').forEach(img => {
                                    let rect = img.getBoundingClientRect();
                                    let area = rect.width * rect.height;
                                    if (area > maxArea && rect.width > 200) { maxArea = area; bestSrc = img.src; }
                                });
                                return bestSrc;
                            """)
                            if largest_img: self._last_fb_image = largest_img
                        except: pass

                # 2. Title Strategy
                try:
                    # a. Thử lấy og:title (Title chuẩn của bài post)
                    og_title_el = use_driver.find_elements(By.XPATH, "//meta[@property='og:title']")
                    if og_title_el:
                        og_title = og_title_el[0].get_attribute("content")
                        if og_title and "Facebook" not in og_title:
                            fetched_title = og_title
                            print(f"[FB] Selenium found og:title: {fetched_title}")
                    
                    # b. Nếu không có, thử lấy Page Title
                except Exception as e:
                    print(f"[FB] Selenium getting og:title error: {e}")

                try:
                    # Wait for title in <title> tag
                    import time
                    time.sleep(1)  # Giảm từ 3s -> 1s
                    
                    # Try getting title from document.title
                    t = use_driver.title
                    
                    # Clean up title
                    if t:
                        if "Log in" in t or "Đăng nhập" in t:
                             self.log("⚠️ Facebook yêu cầu đăng nhập (Login Wall) - Vui lòng nhập Cookie!")
                             fetched_title = "Facebook Login Required"
                        else:
                             t = t.replace(" | Facebook", "").replace("Facebook", "").strip()
                    
                    # If empty or generic, try finding h1 or video label
                    if not t or len(t) < 5:
                         try:
                             # Try to find user status/caption
                             # Ensure uc is imported for By if not already in scope
                             if 'uc' not in locals() and 'uc' not in globals():
                                 import undetected_chromedriver as uc
                             elem = use_driver.find_element(uc.By.CSS_SELECTOR, 'div[data-ad-preview="message"]')
                             if elem: t = elem.text
                         except: pass

                    if t and len(t) > 3:
                        fetched_title = t
                        print(f"[FB] Browser extracted: {fetched_title[:50]}...")
                    else:
                        fetched_title = "Facebook Video"

                except Exception as e:
                     print(f"[FB] Browser error: {e}")
                     fetched_title = "Facebook Video"
                
            except Exception as e:
                print(f"[FB] Driver fallback error: {e}")
            finally:
                if should_quit and use_driver:
                    try: 
                        use_driver.quit()
                    except: pass


        if fetched_title:
            fetched_title = fetched_title.replace(' | Facebook', '').replace('Facebook', '').strip()
            if any(inv in fetched_title.lower() for inv in invalid_titles): fetched_title = None

        return fetched_title or "Facebook Video"
    
    def create_vimeo_embed(self, video_id, title="Vimeo Video"):
        """
        Create Vimeo embed code with automatic aspect ratio detection
        
        Args:
            video_id: Vimeo video ID
            title: Video title for accessibility
            
        Returns:
            str: HTML embed code with responsive wrapper
        """
        try:
            import requests
            
            # Get video info from Vimeo oEmbed API
            oembed_url = f"https://vimeo.com/api/oembed.json?url=https://vimeo.com/{video_id}"
            response = requests.get(oembed_url, timeout=5)
            
            if response.status_code == 200:
                data = response.json()
                width = data.get('width', 640)
                height = data.get('height', 360)
                
                # Calculate aspect ratio
                aspect_ratio = width / height if height > 0 else 16/9
                
                # Determine if vertical (9:16) or horizontal (16:9)
                is_vertical = aspect_ratio < 1  # width < height
                
                if is_vertical:
                    # Vertical video (9:16) - Use max-width and center
                    padding_percent = (height / width) * 100  # e.g., 177.78% for 9:16
                    
                    embed_code = f'''<div style="max-width:400px;margin:0 auto;">
<div style="padding:{padding_percent:.2f}% 0 0 0;position:relative;">
<iframe src="https://player.vimeo.com/video/{video_id}?badge=0&autopause=0&player_id=0&app_id=58479" 
frameborder="0" 
allow="autoplay; fullscreen; picture-in-picture; clipboard-write" 
style="position:absolute;top:0;left:0;width:100%;height:100%;" 
title="{title}"></iframe>
</div>
</div>
<script src="https://player.vimeo.com/api/player.js"></script>'''
                    
                    print(f"[VIMEO] 📱 Vertical video detected: {width}x{height} (aspect: {aspect_ratio:.2f})")
                else:
                    # Horizontal video (16:9) - Full width responsive
                    padding_percent = (height / width) * 100  # e.g., 56.25% for 16:9
                    
                    embed_code = f'''<div style="padding:{padding_percent:.2f}% 0 0 0;position:relative;">
<iframe src="https://player.vimeo.com/video/{video_id}?badge=0&autopause=0&player_id=0&app_id=58479" 
frameborder="0" 
allow="autoplay; fullscreen; picture-in-picture; clipboard-write" 
style="position:absolute;top:0;left:0;width:100%;height:100%;" 
title="{title}"></iframe>
</div>
<script src="https://player.vimeo.com/api/player.js"></script>'''
                    
                    print(f"[VIMEO] 🖥️ Horizontal video detected: {width}x{height} (aspect: {aspect_ratio:.2f})")
                
                return embed_code
            else:
                print(f"[VIMEO] ⚠️ Could not fetch video info, using default 16:9")
                
        except Exception as e:
            print(f"[VIMEO] ❌ Error detecting aspect ratio: {e}")
        
        # Fallback to standard 16:9 embed
        return f'''<div style="padding:56.25% 0 0 0;position:relative;">
<iframe src="https://player.vimeo.com/video/{video_id}?badge=0&autopause=0&player_id=0&app_id=58479" 
frameborder="0" 
allow="autoplay; fullscreen; picture-in-picture; clipboard-write" 
style="position:absolute;top:0;left:0;width:100%;height:100%;" 
title="{title}"></iframe>
</div>
<script src="https://player.vimeo.com/api/player.js"></script>'''


    def create_facebook_embed(self, fb_url, use_sdk=False):
        """Convert Facebook URL to embed code (267x591 - 9:16 ratio for portrait video)
        
        Args:
            fb_url: Facebook video URL
            use_sdk: If True, use Facebook SDK method (bypass security). If False, use iframe method.
        """
        import re
        from urllib.parse import quote
        
        # Clean URL
        clean_url = fb_url.strip()
        clean_url = clean_url.replace("m.facebook.com", "www.facebook.com")
        clean_url = clean_url.replace("web.facebook.com", "www.facebook.com")
        
        # Decide method: Use SDK if requested, otherwise Iframe
        if use_sdk:
            # Method 2: Facebook SDK (Javascript based)
            embed_code = (
                f'<div id="fb-root"></div>\n'
                f'<script async defer crossorigin="anonymous" src="https://connect.facebook.net/en_US/sdk.js#xfbml=1&version=v18.0" nonce="Aut0Wpr"></script>\n'
                f'<div class="fb-video" data-href="{clean_url}" data-width="267" data-show-text="false"></div>'
            )
        else:
            # Method 1: Iframe (HTML only, no JS required)
            # URL encode for Facebook embed
            encoded_url = quote(clean_url, safe='')
            
            # Dimensions for vertical video (9:16 approx)
            w = "360"
            h = "640"
            
            # Use a robust iframe URL compatible with Reels
            embed_code = (
                f'<iframe src="https://www.facebook.com/plugins/video.php?height={h}&href={encoded_url}&show_text=false&width={w}&t=0" '
                f'width="{w}" height="{h}" '
                f'style="border:none;overflow:hidden" '
                f'scrolling="no" frameborder="0" allowfullscreen="true" '
                f'allow="autoplay; clipboard-write; encrypted-media; picture-in-picture; web-share"></iframe>'
            )
        
        return embed_code

    def add_to_queue(self):
        # Grab raw text from Textbox (Bulk support)
        raw_video_text = self.entry_video.get("1.0", "end").strip()
        data = self.get_post_data() # Gets other fields (Title, Image, etc)
        
        # Check if empty
        if not data.title and not raw_video_text:
            messagebox.showwarning("Thiếu thông tin", "Nhập ít nhất tiêu đề hoặc video URL!")
            return

        # SPLIT LINES for Bulk Processing
        video_lines = [line.strip() for line in raw_video_text.split('\n') if line.strip()]
        
        if len(video_lines) > 1:
            # BULK MODE
            added_count = 0
            base_title = data.title
            
            # --- SHARED DRIVER INIT ---
            shared_driver = None
            has_fb = any(('facebook.com' in x or 'fb.watch' in x) for x in video_lines)
            if has_fb:
                try:
                    self.log("🚀 Đang khởi động trình duyệt ẩn danh (Shared) để quét Facebook...")
                    import undetected_chromedriver as uc
                    options = uc.ChromeOptions()
                    options.add_argument('--user-agent=Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1')
                    options.add_argument("--disable-gpu")
                    options.add_argument("--no-sandbox")
                    options.add_argument("--disable-dev-shm-usage")
                    options.add_argument("--mute-audio")
                    options.add_argument("--window-size=375,812")
                    
                    chrome_path = os.path.join(os.getcwd(), "chrome_portable", "chrome.exe")
                    if os.path.exists(chrome_path): options.binary_location = chrome_path
                    
                    shared_driver = uc.Chrome(options=options, version_main=144, headless=True)
                    shared_driver.set_page_load_timeout(45)
                except Exception as e:
                    print(f"Failed to init shared driver: {e}")
            # --------------------------

            self.log(f"📦 Phát hiện {len(video_lines)} dòng video. Đang thêm xử lý hàng loạt...")
            
            # SPEED OPTIMIZATION: Process in parallel
            import concurrent.futures
            
            def process_single_video(idx, vid_line, base_title, shared_driver):
                """Process a single video line (runs in parallel)"""
                try:
                    # Clone data for each post
                    import copy
                    current_post = copy.deepcopy(data)
                    current_post.video_url = vid_line
                    
                    # Handle Title
                    if not base_title: 
                         current_post.title = "" 
                    else:
                         current_post.title = f"{base_title} (Part {idx+1})"

                    # --- Facebook Logic (Mini version for bulk) ---
                    if 'facebook.com' in vid_line or 'fb.watch' in vid_line:
                        if not vid_line.startswith('<'): # URL not Embed
                             # Fetch title if missing (ONLY if no base title)
                             if not current_post.title:
                                 try:
                                     # Fast fetch attempt
                                     ft = self.get_facebook_title(vid_line, driver=shared_driver)
                                     if ft and ft != "Facebook Video":
                                         current_post.title = ft
                                 except: pass
                                 
                             # Generate Embed
                             current_post.video_url = self.create_facebook_embed(vid_line)
                    
                    # --- YouTube Logic ---
                    elif 'youtube.com' in vid_line or 'youtu.be' in vid_line:
                        if not vid_line.startswith('<'):
                             # Extract video ID
                             import re
                             video_id = None
                             if 'youtu.be' in vid_line:
                                 match = re.search(r'youtu\.be/([\w-]+)', vid_line)
                                 if match: video_id = match.group(1)
                             else:
                                 match = re.search(r'v=([\w-]+)', vid_line)
                                 if match: video_id = match.group(1)
                             
                             if video_id:
                                 width = 560
                                 height = 315
                                 current_post.video_url = f'<iframe width="{width}" height="{height}" src="https://www.youtube.com/embed/{video_id}" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen="true"></iframe>'
                                 
                                 # Optional: Fetch title via oEmbed (fast)
                                 if not current_post.title:
                                     try:
                                         import requests
                                         oembed_url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
                                         res = requests.get(oembed_url, timeout=2)
                                         if res.status_code == 200:
                                             yt_data = res.json()
                                             current_post.title = yt_data.get('title', '')
                                     except: pass

                    # --- Vimeo Logic ---
                    elif 'vimeo.com' in vid_line:
                        if not vid_line.strip().startswith('<') and 'player.vimeo.com' not in vid_line:
                             import re
                             match = re.search(r'vimeo\.com/(\d+)', vid_line)
                             if match:
                                 vid_id = match.group(1)
                                 current_post.video_url = self.create_vimeo_embed(vid_id, current_post.title or "Vimeo Video")
                                 
                                 try:
                                      import requests
                                      oembed_url = f"https://vimeo.com/api/oembed.json?url=https://vimeo.com/{vid_id}&width=1280"
                                      res = requests.get(oembed_url, timeout=2)
                                      if res.status_code == 200:
                                          data = res.json()
                                          if not current_post.title:
                                              current_post.title = data.get('title', '')
                                          # Lấy thumbnail chất lượng cao (1280px)
                                          if not current_post.image_url:
                                              current_post.image_url = data.get('thumbnail_url', '')
                                 except: pass
                    
                    # Return cleaned data
                    return {
                        'title': current_post.title,
                        'video_url': current_post.video_url,
                        'image_url': current_post.image_url,
                        'content': current_post.content,
                        'needs_body_content': getattr(current_post, 'needs_body_content', True),
                        'theme': self.get_selected_theme_id()
                    }
                    
                except Exception as e:
                    print(f"Error processing line {idx}: {e}")
                    return None
            
            # Process all videos in parallel (max 8 workers for speed)
            results = []
            with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
                futures = {
                    executor.submit(process_single_video, idx, vid_line, base_title, shared_driver): idx
                    for idx, vid_line in enumerate(video_lines)
                }
                
                for future in concurrent.futures.as_completed(futures):
                    try:
                        result = future.result()
                        if result:
                            results.append(result)
                            added_count += 1
                    except Exception as e:
                        print(f"Error in parallel processing: {e}")
            
            # Add all results to queue
            self.post_queue.extend(results)
            
            # --- SHARED DRIVER CLEANUP ---
            if shared_driver:
                try:
                    shared_driver.quit()
                    self.log("🏁 Đã đóng trình duyệt Shared.")
                except: pass
            # -----------------------------

            self.log(f"✅ Đã thêm {added_count} bài viết vào hàng chờ!")
            self.update_queue_display()
            # Clear ALL inputs after bulk add
            self.entry_video.delete("1.0", "end")
            self.entry_title.delete(0, "end")
            self.entry_image.delete(0, "end")
            if hasattr(self, 'entry_content_image'): self.entry_content_image.delete(0, "end")
            if hasattr(self, 'entry_content_image2'): self.entry_content_image2.delete(0, "end")
            if hasattr(self, 'entry_content_image3'): self.entry_content_image3.delete(0, "end")
            self.textbox_content.delete("1.0", "end")
            return

        # SINGLE MODE (Original Logic)
        # Ensure data.video_url is set correctly from the textbox
        data.video_url = raw_video_text
        
        if not data.title and not data.video_url:
             messagebox.showwarning("Thiếu thông tin", "Nhập ít nhất tiêu đề hoặc video URL!")
             return
             
        # ... Continue with original single logic below ...
        
        # --- Facebook Processing for Single Post ---
        if data.video_url and ('facebook.com' in data.video_url or 'fb.watch' in data.video_url):
            # Only process if it looks like a URL, not already an embed code
            if not data.video_url.strip().startswith('<'):
                # Check if should use oEmbed (let WordPress handle it)
                use_oembed = self.initial_config.get('facebook_use_oembed', False)
                
                if use_oembed:
                    self.log("🔍 Đang xử lý link Facebook (lấy tiêu đề, giữ URL cho WordPress oEmbed)...")
                    
                    # Only fetch title, keep URL as-is for WordPress oEmbed
                    if not data.title:
                        try:
                            self.update_idletasks() 
                            fetched_title = self.get_facebook_title(data.video_url)
                            if fetched_title and fetched_title != "Facebook Video":
                                data.title = fetched_title
                                self.entry_title.delete(0, "end")
                                self.entry_title.insert(0, fetched_title)
                        except Exception as e:
                            print(f"Lỗi lấy tiêu đề FB: {e}")
                    
                    # Keep URL as-is - WordPress will handle oEmbed
                    self.log("✅ Giữ URL Facebook cho WordPress tự xử lý (oEmbed)")
                else:
                    self.log("🔍 Đang xử lý link Facebook (lấy tiêu đề & mã nhúng)...")
                    
                    # 1. Fetch Title if missing
                    if not data.title:
                        try:
                            # Force UI update
                            self.update_idletasks() 
                            fetched_title = self.get_facebook_title(data.video_url)
                            if fetched_title and fetched_title != "Facebook Video":
                                data.title = fetched_title
                                self.entry_title.delete(0, "end")
                                self.entry_title.insert(0, fetched_title)
                        except Exception as e:
                            print(f"Lỗi lấy tiêu đề FB: {e}")

                    # 2. Generate Embed Code
                    try:
                        use_sdk = self.initial_config.get('facebook_use_sdk', False)
                        embed_code = self.create_facebook_embed(data.video_url, use_sdk=use_sdk)
                        if embed_code:
                            data.video_url = embed_code
                            # Optional: Update UI to show it's now an embed code? 
                            # self.entry_video.delete(0, "end")
                            # self.entry_video.insert(0, "[Converted to Embed Code]")
                    except Exception as e:
                        self.log(f"⚠️ Lỗi tạo embed FB: {e}")

        # --- Vimeo Processing for Single Post (If not handled by _extract_video_url) ---
        if data.video_url and 'vimeo.com' in data.video_url:
             # Check if it is a raw link that needs conversion to full embed code
             if not data.video_url.strip().startswith('<') and 'player.vimeo.com' not in data.video_url:
                  import re
                  match = re.search(r'vimeo\.com/(\d+)', data.video_url)
                  if match:
                      vid_id = match.group(1)
                      # Use new helper function for proper aspect ratio
                      data.video_url = self.create_vimeo_embed(vid_id, data.title or "Vimeo Video")
                      self.log(f"✅ Đã chuyển link Vimeo sang Embed Code (tự động phát hiện tỷ lệ)")
                      
                      # Fetch title if missing
                      if not data.title:
                          try:
                              import requests
                              oembed_url = f"https://vimeo.com/api/oembed.json?url=https://vimeo.com/{vid_id}"
                              res = requests.get(oembed_url, timeout=2)
                              if res.status_code == 200:
                                  data.title = res.json().get('title', '')
                                  self.entry_title.delete(0, "end")
                                  self.entry_title.insert(0, data.title)
                          except: pass


        # Auto-generate title if missing but has video
        if not data.title and data.video_url:
            if "youtube" in data.video_url.lower():
                data.title = f"Video YouTube - {time.strftime('%H:%M:%S')}"
            elif "vimeo" in data.video_url.lower():
                data.title = f"Video Vimeo - {time.strftime('%H:%M:%S')}"
            elif "facebook" in data.video_url.lower() or "fb.watch" in data.video_url.lower():
                data.title = f"Video Facebook - {time.strftime('%H:%M:%S')}"
            else:
                data.title = f"Video Embed - {time.strftime('%H:%M:%S')}"
        
        # Prepare queue item
        queue_item = data.__dict__
        
        # Check if content is empty -> AUTO ENABLE CONTENT MATCHING FROM POOL
        if not data.content.strip():
            queue_item['needs_body_content'] = True
            if hasattr(self, 'content_pool') and self.content_pool:
                self.log(f"   ℹ️ Bài viết trống nội dung -> Sẽ tự động ghép từ pool ({len(self.content_pool)} bài sẵn có)")
        
        self.post_queue.append(queue_item)
        self.update_queue_display()
        self.log(f"✅ Đã thêm vào hàng chờ: {data.title}")
        
        # Clear inputs
        self.entry_title.delete(0, "end")
        self.entry_video.delete("1.0", "end")
        self.entry_image.delete(0, "end")
        if hasattr(self, 'entry_content_image'): self.entry_content_image.delete(0, "end")  # Clear content image 1
        if hasattr(self, 'entry_content_image2'): self.entry_content_image2.delete(0, "end")  # Clear content image 2
        if hasattr(self, 'entry_content_image3'): self.entry_content_image3.delete(0, "end")  # Clear content image 3
        self.textbox_content.delete("1.0", "end")

    def update_queue_display(self):
        """Update Queue display using Cards with Thumbnails - Debounced + Optimized"""
        # Debounce: cancel pending update, schedule new one 500ms later
        # This prevents UI freeze when adding many items rapidly
        if hasattr(self, '_queue_update_job') and self._queue_update_job:
            try:
                self.after_cancel(self._queue_update_job)
            except:
                pass
        self._queue_update_job = self.after(500, self._do_update_queue_display)

    def _do_update_queue_display(self):
        # Update label count
        if hasattr(self, 'lbl_queue_count'):
            self.lbl_queue_count.configure(text=f"({len(self.post_queue)} bài)")
            
        # Enable/Disable Run button
        if hasattr(self, 'btn_batch_post'):
            if self.post_queue:
                self.btn_batch_post.configure(state="normal", text="🚀 CHAY AUTO")
            else:
                self.btn_batch_post.configure(state="disabled", text="▶️ CHẠY AUTO POST (HÀNG CHỜ)")
            
        # Clear existing cards
        if hasattr(self, 'queue_scroll'):
            for widget in self.queue_scroll.winfo_children():
                widget.destroy()
        else:
            return
            
        if not self.post_queue:
            ctk.CTkLabel(self.queue_scroll, text="Hàng chờ trống. Thêm bài viết để bắt đầu.", text_color="gray").pack(pady=20)
            return

        # Limit display to avoid lag
        max_display = 20
        display_queue = self.post_queue[:max_display]
        hidden_count = len(self.post_queue) - max_display

        # Persistent image refs
        if not hasattr(self, '_queue_images'):
            self._queue_images = []
        self._queue_images.clear()

        for idx, post in enumerate(display_queue):
            try:
                title = post.get('title', 'No Title')
                video_url = post.get('video_url', '')
                image_url = post.get('thumbnail', '') or post.get('image_url', '')
                
                card = ctk.CTkFrame(self.queue_scroll, fg_color="#2b2b2b", corner_radius=8, border_width=1, border_color="#404040")
                card.pack(fill="x", pady=3, padx=5)
                
                # Thumbnail (Left) - use cache
                thumb_frame = ctk.CTkFrame(card, width=70, height=50, fg_color="#1a1a1a")
                thumb_frame.pack(side="left", padx=5, pady=4)
                thumb_frame.pack_propagate(False)
                
                has_img = False
                if image_url and os.path.exists(image_url):
                    try:
                        pil_img = self._load_image_cached(image_url, (70, 50))
                        if pil_img:
                            ctk_img = ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=(70, 50))
                            self._queue_images.append(ctk_img)
                            ctk.CTkLabel(thumb_frame, text="", image=ctk_img).pack(expand=True)
                            has_img = True
                    except Exception as e:
                        print(f"Thumb error: {e}")
                
                if not has_img:
                    ctk.CTkLabel(thumb_frame, text="🖼️", font=("Arial", 18)).pack(expand=True)
                
                # Info (Center)
                info_frame = ctk.CTkFrame(card, fg_color="transparent")
                info_frame.pack(side="left", fill="both", expand=True, padx=8)
                
                post_type = "📄"
                if 'facebook' in video_url: post_type = "🟦 FB"
                elif 'youtube' in video_url: post_type = "🟥 YT"
                elif 'vimeo' in video_url: post_type = "🟦 Vimeo"
                
                ctk.CTkLabel(info_frame, text=f"{idx+1}. {post_type}", font=("Segoe UI", 9), text_color="gray", anchor="w").pack(fill="x")
                ctk.CTkLabel(info_frame, text=title[:55] + "..." if len(title) > 55 else title, font=("Segoe UI", 11, "bold"), text_color="white", anchor="w").pack(fill="x")
                
                # Actions (Right)
                action_frame = ctk.CTkFrame(card, fg_color="transparent")
                action_frame.pack(side="right", padx=8)
                
                ctk.CTkButton(action_frame, text="🗑️", width=36, height=30, fg_color="#ef4444", hover_color="#dc2626",
                              command=lambda i=idx: self.remove_from_queue(i)).pack()
                
                # Yield to UI every 5 cards to prevent freeze
                if (idx + 1) % 5 == 0:
                    self.update_idletasks()
                
            except Exception as e:
                print(f"Error displaying queue item {idx}: {e}")

        if hidden_count > 0:
            ctk.CTkLabel(self.queue_scroll, text=f"... và {hidden_count} bài viết nữa đang ẩn (vẫn sẽ được tự động đăng)", font=("Segoe UI", 12, "italic"), text_color="gray").pack(pady=15)

    def remove_from_queue(self, index):
        """Remove item from queue by index and refresh display"""
        if 0 <= index < len(self.post_queue):
            removed = self.post_queue.pop(index)
            self.log(f"🗑️ Đã xoá khỏi hàng chờ: {removed.get('title')}")
            self.update_queue_display()

    def clear_queue(self):
        self.post_queue = []
        self.update_queue_display()
        self.log("🗑️ Đã xóa toàn bộ hàng chờ.")

    def get_post_data(self):
        data = AppData()
        data.title = self.entry_title.get()
        
        # Extract video URL from iframe if user pasted full iframe code
        video_input = self.entry_video.get("1.0", "end").strip()
        data.video_url = self._extract_video_url(video_input)
        
        data.image_url = self.entry_image.get()
        data.content_image = self.entry_content_image.get()  # Content image 1
        data.content_image2 = self.entry_content_image2.get()  # Content image 2
        data.content_image3 = self.entry_content_image3.get()  # Content image 3
        data.auto_fetch_images = self.chk_auto_fetch_images.get()  # Auto-fetch flag
        data.content = self.textbox_content.get("1.0", "end")
        data.theme = self.get_selected_theme_id()  # Get selected theme
        return data
    
    def _extract_video_url(self, input_text):
        """
        Extract video URL or full embed code from various input formats:
        - Full Vimeo embed with div wrapper: Use entire code
        - Facebook video URL: Convert to fixed-size iframe (267x591)
        - Iframe code: Extract src URL
        - Direct URL: Use as-is
        - Vimeo share link: Convert to player URL
        """
        if not input_text:
            return ""
        
        import re
        from urllib.parse import quote
        
        # Case 1: Full Vimeo embed code with <div> wrapper and <script>
        # This is the BEST format - use it entirely
        if '<div style="padding:' in input_text and '<script src="https://player.vimeo.com/api/player.js">' in input_text:
            print(f"[GUI] Detected full Vimeo embed code with wrapper")
            # Return the entire embed code - WordPress will handle it
            return input_text
        
        # Case 2: Facebook video URL - Convert to responsive iframe
        if 'facebook.com' in input_text or 'fb.watch' in input_text or 'fb.com' in input_text:
            print(f"[GUI] Detected Facebook video URL")
            
            # Extract the actual Facebook URL if it's in iframe already
            if '<iframe' in input_text:
                match = re.search(r'href=([^&\s]+)', input_text)
                if match:
                    fb_url = match.group(1)
                    # URL decode if needed
                    fb_url = fb_url.replace('%3A', ':').replace('%2F', '/')
                else:
                    # Try to extract from src
                    match = re.search(r'src=["\']([^"\']+)["\']', input_text)
                    if match:
                        fb_url = match.group(1)
                    else:
                        fb_url = input_text
            else:
                fb_url = input_text.strip()
            
            # Ensure it's a full URL
            if not fb_url.startswith('http'):
                fb_url = 'https://' + fb_url
            
            # URL encode for Facebook embed
            encoded_url = quote(fb_url, safe='')
            
            # Create Facebook iframe (267x591 - 9:16 ratio for portrait video)
            # Check if should use SDK method
            use_sdk = self.initial_config.get('facebook_use_sdk', False)
            
            if use_sdk:
                # Method 2: Facebook SDK (bypass security)
                # Format: <div class="fb-video" data-href="URL" data-width="267"></div>
                fb_iframe = (
                    f'<!-- Facebook SDK Script -->'
                    f'<script>window.fbAsyncInit = function() {{'
                    f'FB.init({{appId: "YOUR_APP_ID",xfbml: true,version: "v12.0"}});'
                    f'}};'
                    f'(function(d, s, id){{'
                    f'var js, fjs = d.getElementsByTagName(s)[0];'
                    f'if (d.getElementById(id)) {{return;}}'
                    f'js = d.createElement(s); js.id = id;'
                    f'js.src = "https://connect.facebook.net/en_US/sdk.js";'
                    f'fjs.parentNode.insertBefore(js, fjs);'
                    f'}}(document, "script", "facebook-jssdk"));</script>'
                    f'<!-- Facebook Video Embed -->'
                    f'<div class="fb-video" data-href="{fb_url}" data-width="267"></div>'
                )
            else:
                # Method 1: Direct iframe (default) - CRITICAL: NO SPACES between attributes!
                fb_iframe = (
                    f'<iframe src="https://www.facebook.com/plugins/video.php?height=591&href={encoded_url}&show_text=false&width=267&t=0" '
                    f'width="267" '
                    f'height="591" '
                    f'style="border:none;overflow:hidden" '
                    f'scrolling="no" '
                    f'frameborder="0" '
                    f'allowfullscreen="true" '
                    f'allow="autoplay; clipboard-write; encrypted-media; picture-in-picture; web-share" '
                    f'allowFullScreen="true">'
                    f'</iframe>'
                )
            
            print(f"[GUI] Created Facebook iframe: {fb_iframe[:100]}...")
            return fb_iframe
        
        # Case 3: Already a clean player URL
        if 'player.vimeo.com/video/' in input_text and '<' not in input_text:
            return input_text
        
        # Case 4: Iframe code - extract src attribute
        if '<iframe' in input_text.lower():
            # Match src="..." or src='...'
            match = re.search(r'src=["\']([^"\']+)["\']', input_text)
            if match:
                url = match.group(1)
                print(f"[GUI] Extracted URL from iframe: {url}")
                return url
        
        # Case 5: Regular vimeo.com link - convert to player URL
        if 'vimeo.com/' in input_text and 'player.vimeo.com' not in input_text:
            # Extract video ID
            match = re.search(r'vimeo\.com/(\d+)', input_text)
            if match:
                video_id = match.group(1)
                player_url = f"https://player.vimeo.com/video/{video_id}"
                print(f"[GUI] Converted vimeo.com link to player URL: {player_url}")
                return player_url
        
        # Case 6: Just return as-is (might be other video platform)
        return input_text

    def on_post_click(self):
        data = self.get_post_data()
        self.log("Đang đăng bài lẻ...")
        self.controller.handle_post_request(data)

    def on_queue_post_click(self):
        if not self.post_queue: 
            self.log("❌ Hàng chờ trống!")
            return
        
        self.log(f"🚀 Bắt đầu chạy AUTO - {len(self.post_queue)} bài trong hàng chờ")
        self.btn_batch_post.configure(state="disabled", text="⏳ Đang chạy AUTO...")
        
        # Start processing queue
        self.process_next_queue_post()

    def process_next_queue_post(self):
        if not self.post_queue:
            self.log("✅ Hoàn thành AUTO! Tất cả bài đã được xử lý.")
            self.btn_batch_post.configure(state="normal", text="🚀 CHAY AUTO")
            return
        
        item = self.post_queue[0]
        self.log(f"📝 Đang xử lý bài {len(self.post_queue)} còn lại: {item.get('title', 'Không có tiêu đề')}")
        
        # Check if this post needs body content from pool
        if item.get('needs_body_content', False) and hasattr(self, 'content_pool') and self.content_pool:
            # Get content from pool NOW (when posting)
            content_item = self.content_pool.pop(0)
            body_content = content_item['content']
            content_filename = content_item['filename']
            content_filepath = content_item.get('filepath', '')  # Get filepath for deletion
            
            # DON'T combine here - let wp_model.py handle video placement
            # Just update the raw_content with body text only
            item['content'] = body_content  # Only body content, no embed code
            item['needs_body_content'] = False
            item['content_source'] = content_filename
            item['txt_filepath'] = content_filepath  # Store filepath for deletion after post
            
            self.log(f"   📄 Ghép nội dung từ: {content_filename}")
            self.log(f"   🗑️ Đã xóa khỏi pool để tránh trùng lặp")
            self.log(f"   ✅ Nội dung body: {len(body_content)} ký tự")
            
            # Show remaining
            remaining = len(self.content_pool)
            if remaining > 0:
                self.log(f"   💾 Còn {remaining} nội dung trong pool")
            else:
                self.log(f"   ℹ️ Pool đã hết nội dung")
        
        # Convert dict to AppData
        data = AppData()
        data.title = item.get('title', '')
        data.video_url = item.get('video_url', '')  # This will be used to generate embed in wp_model.py
        data.image_url = item.get('image_url', '')

        # --- FALLBACK: Use Profile Pic from 'thumbnails' if no video thumb ---
        if not data.image_url:
            try:
                import random
                thumb_dir = "thumbnails"
                if os.path.exists(thumb_dir):
                    profiles = [
                        f for f in os.listdir(thumb_dir) 
                        if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')) 
                        and not f.startswith("car_api_") # Filter out downloaded car images
                    ]
                    if profiles:
                        # Pick random profile pic
                        selected_profile = random.choice(profiles)
                        data.image_url = os.path.abspath(os.path.join(thumb_dir, selected_profile))
                        self.log(f"   🖼️ Không có ảnh video -> Lấy ảnh đại diện: {selected_profile}")
            except Exception as e:
                print(f"Error picking profile pic: {e}")
        data.content_image = item.get('content_image', '')  # Content image 1
        data.content_image2 = item.get('content_image2', '')  # Content image 2
        data.content_image3 = item.get('content_image3', '')  # Content image 3
        
        # Feature: Tắt lấy ảnh từ API
        disable_api_scan = bool(self.chk_disable_image_api.get()) if hasattr(self, 'chk_disable_image_api') else False
        disable_api_media = not bool(self.chk_auto_fetch_images.get()) if hasattr(self, 'chk_auto_fetch_images') else False
        data.auto_fetch_images = not (disable_api_scan or disable_api_media)
        
        data.content = item.get('content', '')
        
        # Get theme from item or fallback to selected theme in GUI
        # This ensures that even old queue items respect the current GUI selection if they don't have a theme
        data.theme = item.get('theme', self.get_selected_theme_id())
        print(f"[DEBUG] Batch item theme: {data.theme}")
        
        # LƯU TITLE VÀO BIẾN TẠM để dùng trong on_post_finished
        self.current_posting_title = data.title
        print(f"[DEBUG] Saved current_posting_title: '{self.current_posting_title}'")  # Terminal only
        
        # Call controller to handle the post
        if self.controller:
            self.controller.handle_post_request(data, is_batch=True)
        else:
            # Mock success for testing
            self.after(2000, lambda: self.on_post_finished(True, f"https://test.com/{data.title.replace(' ','-')}", True))

    def on_post_finished(self, success, message, is_batch=False, post_title=None, image_path=None):
        if success:
            self.log(f"✅ THÀNH CÔNG: {message}")
            
            # Use the passed post_title if available logic
            final_title = "Unknown"
            
            if post_title:
                final_title = post_title
                print(f"[DEBUG] Got clean title passed from controller: '{final_title}'")  # Terminal only
            elif is_batch:
                # Fallback (should rarely happen now)
                if self.post_queue:
                    final_title = self.post_queue[0].get('title', 'Unknown')
            else:
                # Đăng bài lẻ - lấy từ input fields
                if hasattr(self, 'entry_title'):
                    try:
                        final_title = self.entry_title.get().strip()
                    except:
                        pass
            
            if not final_title or final_title == "Unknown":
                 final_title = f"Post {len(self.published_links) + 1}"

            # --- Resolve Image Path for History if not provided ---
            if not image_path:
                if is_batch and self.post_queue:
                    # Get from current queue item
                    image_path = self.post_queue[0].get('image_url', '')
                elif not is_batch:
                    # Get from UI
                    if hasattr(self, 'entry_image'):
                        image_path = self.entry_image.get()

            # Add to history with title AND image
            try:
                print(f"[DEBUG] Adding to history: '{final_title}' -> {message}")
                # Store in list with image path
                timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
                history_item = {
                    'title': final_title,
                    'link': message,
                    'timestamp': timestamp,
                    'status': 'checking', # checking, success, error
                    'image_path': image_path
                }
                self.published_links.append(history_item)
                
                # Render immediately vào history_textbox
                def _render_history(item=history_item, idx=len(self.published_links)):
                    try:
                        if not hasattr(self, 'history_textbox'):
                            return
                        self.history_textbox.configure(state="normal")
                        
                        # --- Thumbnail ---
                        img_path = item.get('image_path', '')
                        has_image = False
                        if img_path and os.path.exists(img_path):
                            try:
                                pil_img = self._load_image_cached(img_path, (80, 80))
                                if pil_img:
                                    from PIL import ImageTk
                                    photo = ImageTk.PhotoImage(pil_img)
                                    if not hasattr(self, 'history_images'):
                                        self.history_images = []
                                    self.history_images.append(photo)
                                    self.history_textbox.image_create("end", image=photo)
                                    self.history_textbox.insert("end", "  ")
                                    has_image = True
                            except Exception as img_err:
                                print(f"[GUI] Error loading history thumb: {img_err}")
                        
                        indent = "          " if has_image else ""
                        
                        # --- Số thứ tự + icon trạng thái ---
                        self.history_textbox.insert("end", f"[{idx}] ⏳ ")
                        
                        # --- Link có thể click ---
                        link_start = self.history_textbox.index("end-1c")
                        self.history_textbox.insert("end", item['link'])
                        link_end = self.history_textbox.index("end-1c")
                        self.history_textbox.tag_add("link", link_start, link_end)
                        
                        # --- Tiêu đề & thời gian ---
                        self.history_textbox.insert("end", f"\n{indent}📝 {item['title']}\n")
                        self.history_textbox.insert("end", f"{indent}📅 {item['timestamp']}\n")
                        self.history_textbox.insert("end", "─" * 60 + "\n\n")
                        
                        self.history_textbox.see("end")
                        self.history_textbox.configure(state="disabled")
                    except Exception as render_err:
                        print(f"[ERROR] Render history failed: {render_err}")
                
                self.after(0, _render_history)
                
                # AUTO-CHECK LINK sau 5 giây (tăng từ 2s để WordPress kịp xử lý)
                self.after(5000, lambda: self.check_published_link(message, final_title))
                
            except Exception as e:
                self.log(f"❌ Lỗi khi thêm vào lịch sử: {e}")
                import traceback
                traceback.print_exc()
            
            # Auto-delete files after successful post
            if is_batch and self.post_queue:
                completed_item = self.post_queue[0]
                
                # Delete thumbnail file if it exists
                thumbnail_path = completed_item.get('image_url', '')
                if thumbnail_path and os.path.exists(thumbnail_path):
                    try:
                        os.remove(thumbnail_path)
                        self.log(f"🗑️ Đã xóa thumbnail: {os.path.basename(thumbnail_path)}")
                    except Exception as e:
                        self.log(f"⚠️ Không thể xóa thumbnail: {e}")
                
                # Delete txt content file if it exists
                txt_filepath = completed_item.get('txt_filepath', '')
                if txt_filepath and os.path.exists(txt_filepath):
                    try:
                        os.remove(txt_filepath)
                        self.log(f"🗑️ Đã xóa file nội dung: {os.path.basename(txt_filepath)}")
                    except Exception as e:
                        self.log(f"⚠️ Không thể xóa file txt: {e}")
                
                # Remove completed item from queue
                self.post_queue.pop(0)
                self.log(f"✅ Hoàn thành: {completed_item.get('title', 'Không có tiêu đề')}")
                self.update_queue_display()
                # Continue with next item after 2 seconds
                self.after(2000, self.process_next_queue_post)
        else:
            self.log(f"❌ THẤT BẠI: {message}")
            if is_batch:
                # On failure, still remove item and continue (or you can choose to stop)
                failed_item = self.post_queue.pop(0) if self.post_queue else None
                if failed_item:
                    self.log(f"❌ Bỏ qua bài lỗi: {failed_item.get('title', 'Không có tiêu đề')}")
                    
                    # Also delete files on failure to avoid accumulation
                    thumbnail_path = failed_item.get('image_url', '')
                    if thumbnail_path and os.path.exists(thumbnail_path):
                        try:
                            os.remove(thumbnail_path)
                            self.log(f"🗑️ Đã xóa thumbnail của bài lỗi: {os.path.basename(thumbnail_path)}")
                        except:
                            pass
                    
                    # Delete txt file on failure too
                    txt_filepath = failed_item.get('txt_filepath', '')
                    if txt_filepath and os.path.exists(txt_filepath):
                        try:
                            os.remove(txt_filepath)
                            self.log(f"🗑️ Đã xóa file txt của bài lỗi: {os.path.basename(txt_filepath)}")
                        except:
                            pass
                
                self.update_queue_display()
                # Continue with next item after 3 seconds
                self.after(3000, self.process_next_queue_post)

    def _on_link_click(self, event):
        """Handle click on link in history"""
        try:
            # Get the index of the click
            index = self.history_textbox.index(f"@{event.x},{event.y}")
            
            # Get the range of the clicked tag
            tag_ranges = self.history_textbox.tag_ranges("link")
            
            # Find which link was clicked
            for i in range(0, len(tag_ranges), 2):
                start = tag_ranges[i]
                end = tag_ranges[i + 1]
                
                if self.history_textbox.compare(start, "<=", index) and self.history_textbox.compare(index, "<", end):
                    # Get the URL text
                    url = self.history_textbox.get(start, end)
                    
                    # Open in browser
                    import webbrowser
                    webbrowser.open(url)
                    self.log(f"🌐 Đã mở link: {url}")
                    break
        except Exception as e:
            print(f"[ERROR] Failed to open link: {e}")
    
    def add_to_history(self, link, title="", status="pending", image_path=None):
        """Add published post to history with title, link, status AND thumbnail"""
        # This function is now deprecated/replaced by the logic in on_post_finished
        # and update_history_status for initial rendering.
        # Keeping it for now to avoid breaking other parts if any still call it directly.
        print("[WARNING] add_to_history called directly. Consider using on_post_finished logic.")
        try:
            # Store in list with image path
            self.published_links.append({
                'title': title,
                'link': link,
                'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                'status': status,
                'image_path': image_path
            })
            
            # Display in history textbox
            def update_ui():
                if hasattr(self, 'history_textbox'):
                    self.history_textbox.configure(state="normal")
                    
                    index = len(self.published_links)
                    
                    # --- Insert Image (Thumbnail) ---
                    has_image = False
                    if image_path and os.path.exists(image_path):
                        try:
                            pil_img = self._load_image_cached(image_path, (80, 80))
                            if pil_img:
                                from PIL import ImageTk
                                photo = ImageTk.PhotoImage(pil_img)
                                if not hasattr(self, 'history_images'):
                                    self.history_images = []
                                self.history_images.append(photo)
                                self.history_textbox.image_create("end", image=photo)
                                self.history_textbox.insert("end", "  ")
                                has_image = True
                        except Exception as e:
                            print(f"[GUI] Error loading history thumb: {e}")
                    
                    # --- Status Icon ---
                    status_icon = "⏳"
                    if status == "success": status_icon = "✅"
                    elif status == "error": status_icon = "❌"
                    
                    # --- Link First (NEW ORDER) ---
                    indent = "          " if has_image else ""
                    self.history_textbox.insert("end", f"[{index}] {status_icon} ")
                    
                    link_start = self.history_textbox.index("end-1c")
                    self.history_textbox.insert("end", link)
                    link_end = self.history_textbox.index("end-1c")
                    self.history_textbox.tag_add("link", link_start, link_end)
                    
                    # --- Title Second ---
                    self.history_textbox.insert("end", f"\n{indent}📝 {title}\n")
                    
                    # --- Timestamp ---
                    self.history_textbox.insert("end", f"{indent}📅 {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                    self.history_textbox.insert("end", "─" * 60 + "\n\n")
                    
                    self.history_textbox.see("end")
                    self.history_textbox.configure(state="disabled")
                else:
                    print(f"[WARNING] history_textbox not found, but saved to list: {title}")
            
            # Schedule UI update on main thread
            self.after(0, update_ui)
            
        except Exception as e:
            print(f"[ERROR] Failed to add to history: {e}")
            import traceback
            traceback.print_exc()
    
    def update_history_status(self, index, status):
        """Update status icon in history - only update the specific item, no full rebuild"""
        if index >= len(self.published_links): return
        
        self.published_links[index]['status'] = status
        
        # Just update the status icon text in-place - no full rebuild needed
        # The history textbox already has the item, just log the change
        status_icon = "✅" if status == "success" else "❌" if status == "error" else "⏳"
        self.log(f"{status_icon} Bài [{index+1}] cập nhật: {status}")
    
    def check_published_link(self, url, title):
        """
        Tự động check xem link đã đăng có accessible không
        Retry 3 lần với delay để WordPress kịp xử lý
        """
        def _check():
            try:
                import requests
                import time
                import re
                self.log(f"🔍 Đang kiểm tra link: {title}")
                
                # Find index of this link in history
                link_index = len(self.published_links)
                
                # Extract post ID from URL for alternative check
                post_id_match = re.search(r'[?&]p=(\d+)', url)
                post_id = post_id_match.group(1) if post_id_match else None
                
                # Retry up to 3 times with increasing delay
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        response = requests.get(url, timeout=10, allow_redirects=True)
                        status_code = response.status_code
                        
                        if status_code == 200:
                            self.log(f"   ✅ Link OK - Status 200")
                            self.update_history_status(link_index, "success")
                            return
                        elif status_code == 404:
                            if attempt < max_retries - 1:
                                # Retry after delay (WordPress might still be processing)
                                wait_time = (attempt + 1) * 3  # 3s, 6s, 9s
                                print(f"[LINK_CHECK] 404 on attempt {attempt+1}, retrying in {wait_time}s...")
                                time.sleep(wait_time)
                                continue
                            else:
                                # Final attempt failed - provide helpful message
                                if post_id:
                                    self.log(f"   ❌ Link lỗi 404 - Post ID {post_id} tồn tại nhưng URL không accessible")
                                    self.log(f"   💡 Có thể cần flush permalinks: WP Admin → Settings → Permalinks → Save")
                                else:
                                    self.log(f"   ❌ Link lỗi 404 (Not Found)")
                                self.update_history_status(link_index, "error")
                                return
                        elif status_code == 403:
                            self.log(f"   ❌ Link lỗi 403 (Forbidden) - Có thể bị security plugin block")
                            self.update_history_status(link_index, "error")
                            return
                        elif status_code == 500:
                            self.log(f"   ❌ Link lỗi 500 (Server Error)")
                            self.update_history_status(link_index, "error")
                            return
                        else:
                            self.log(f"   ⚠️ Link trả về status {status_code}")
                            self.update_history_status(link_index, "error")
                            return
                    except requests.exceptions.Timeout:
                        if attempt < max_retries - 1:
                            time.sleep(3)
                            continue
                        else:
                            self.log(f"   ⏱️ Timeout khi check link")
                            self.update_history_status(link_index, "error")
                            return
                    except requests.exceptions.ConnectionError:
                        if attempt < max_retries - 1:
                            time.sleep(3)
                            continue
                        else:
                            self.log(f"   🔌 Lỗi kết nối khi check link")
                            self.update_history_status(link_index, "error")
                            return
                    
            except Exception as e:
                self.log(f"   ❌ Lỗi check link: {e}")
                link_index = len(self.published_links)
                self.update_history_status(link_index, "error")
        
        # Run check in background thread
        import threading
        threading.Thread(target=_check, daemon=True).start()

    def copy_history_links(self):
        """Copy all links to clipboard"""
        self.clipboard_clear()
        
        # Create formatted text with Link on one line, Title on next line
        text_lines = []
        for item in self.published_links:
            if isinstance(item, dict):
                # Link on first line
                text_lines.append(item['link'])
                # Title on second line
                text_lines.append(item['title'])
                # Empty line between items
                text_lines.append("")
            else:
                # Old format (just string)
                text_lines.append(str(item))
                text_lines.append("")
        
        # Use Windows line ending for better compatibility
        clipboard_text = "\r\n".join(text_lines)
        self.clipboard_append(clipboard_text)
        
        # Also try using pyperclip for better compatibility
        try:
            import pyperclip
            pyperclip.copy(clipboard_text)
        except:
            pass
        
        self.log(f"📋 Đã copy {len(self.published_links)} link vào clipboard (Link xuống dòng Tiêu đề).")

    def export_history_to_word(self):
        """Export history links to CSV file (Excel compatible)"""
        try:
            if not getattr(self, "published_links", []):
                messagebox.showwarning("Canh bao", "Lich su trong! Khong co gi de xuat.")
                return

            from datetime import datetime
            import os
            import csv
            import subprocess

            # Auto-generate file path in Documents folder
            docs_folder = os.path.expanduser("~\\Documents")
            os.makedirs(docs_folder, exist_ok=True)
            
            # Use CSV format instead of Word (more reliable)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"DanhSachBaiViet_{timestamp}.csv"
            file_path = os.path.join(docs_folder, filename)

            self.log(f"Dang xuat CSV...")
            self.log(f"📁 Luu vao: {file_path}")

            # Write CSV file
            try:
                with open(file_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    # Header: LINK FIRST, TITLE SECOND
                    writer.writerow(['Link Bài Viết', 'Tiêu Đề'])
                    
                    # Data: LINK FIRST, TITLE SECOND
                    for item in self.published_links:
                        if isinstance(item, dict):
                            link = item.get('link', 'N/A')
                            title = item.get('title', 'N/A')
                            writer.writerow([link, title])
                
                self.log(f"✅ Da xuat CSV: {filename}")
                messagebox.showinfo("Thanh cong", f"✅ Da xuat {len(self.published_links)} bai viet!\n\n📁 File: {filename}\n📂 Thu muc: {docs_folder}\n\n💡 File CSV co the mo bang Excel")

                # Auto open file
                try:
                    os.startfile(file_path)
                except Exception as e:
                    self.log(f"⚠️ Khong the mo file tu dong: {e}")
                    try:
                        subprocess.Popen(['start', file_path], shell=True)
                    except:
                        pass
            except Exception as e:
                self.log(f"❌ Loi xuat CSV: {e}")
                messagebox.showerror("Loi", f"Khong the xuat file: {e}")

        except Exception as e:
            self.log(f"❌ Loi: {e}")
            messagebox.showerror("Loi", f"Loi xuat: {e}")


    def export_history_to_excel(self):
        """Export history links to Excel (CSV format) 📊"""
        try:
            if not getattr(self, "published_links", []):
                messagebox.showwarning("Cảnh báo", "Lịch sử trống! Không có gì để xuất.")
                return

            import subprocess
            
            # Auto-generate file path in Documents folder (no dialog)
            docs_folder = os.path.expanduser("~\\Documents")
            os.makedirs(docs_folder, exist_ok=True)
            
            # Use simple filename without special characters
            timestamp = time.strftime('%Y%m%d_%H%M%S')
            filename = f"LichSuLink_{timestamp}.csv"
            file_path = os.path.join(docs_folder, filename)
            
            self.log(f"Dang xuat Excel...")
            self.log(f"📁 Luu vao: {file_path}")
            
            import csv
            # Use 'utf-8-sig' for Excel compatibility with Vietnamese characters
            with open(file_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                # Header
                writer.writerow(["STT", "Thời Gian", "Tiêu Đề", "Link Bài Viết", "Trạng Thái", "Ảnh Thumbnail"])
                
                # Data
                for i, item in enumerate(self.published_links, 1):
                    # Check if item is dict or old format
                    if isinstance(item, dict):
                        time_str = item.get('timestamp', item.get('time', ''))
                        title = item.get('title', 'No Title')
                        link = item.get('link', item.get('url', ''))
                        status = item.get('status', 'N/A')
                        if status == 'success': status = "Thành công" 
                        elif status == 'error': status = "Lỗi"
                        elif status == 'checking': status = "Đang kiểm tra"
                        img = item.get('image_path', '')
                    else:
                        # Fallback for old simple string logs
                        time_str = ""
                        title = "Old Item"
                        link = str(item)
                        status = ""
                        img = ""
                    
                    writer.writerow([i, time_str, title, link, status, img])
                    
            self.log(f"✅ Đã xuất Excel: {filename}")
            messagebox.showinfo("Thành công", f"✅ Đã xuất {len(self.published_links)} dòng!\n\n📁 File: {filename}\n📂 Thu muc: {docs_folder}")
            
            # Auto open file with Excel
            try:
                os.startfile(file_path, 'open')
            except Exception as e:
                self.log(f"⚠️ Khong the mo file tu dong: {e}")
                # Try alternative method
                try:
                    subprocess.Popen(['start', file_path], shell=True)
                except:
                    pass
            
        except Exception as e:
            self.log(f"❌ Lỗi xuất Excel: {e}")
            messagebox.showerror("Lỗi", f"Không thể xuất file: {e}")

    def open_documents_folder(self):
        """Open Documents folder to view exported files"""
        try:
            import os
            docs_folder = os.path.expanduser("~\\Documents")
            os.makedirs(docs_folder, exist_ok=True)
            os.startfile(docs_folder)
            self.log(f"📂 Mở thư mục Documents: {docs_folder}")
        except Exception as e:
            self.log(f"❌ Lỗi mở thư mục: {e}")
            messagebox.showerror("Lỗi", f"Không thể mở thư mục Documents: {e}")

    def smart_export_history(self):
        """Export history in the same format as imported file"""
        # Check if user imported a file
        if hasattr(self, '_imported_file_format'):
            ext = self._imported_file_format
            format_name = {"csv": "CSV", "xlsx": "Excel", "docx": "Word", "odt": "ODT"}.get(ext[1:], "CSV")
            self.log(f"📤 Xuất theo định dạng đã import: {format_name}")
            
            if ext == ".csv":
                self.export_history_to_word()  # This exports CSV now
            elif ext == ".xlsx":
                self.export_history_to_excel()
            elif ext == ".docx":
                self.export_history_to_docx()
            elif ext == ".odt":
                self.export_history_to_odt()
        else:
            # Default to CSV if no import
            self.log("📤 Chưa import file, xuất mặc định CSV")
            self.export_history_to_word()  # This exports CSV now

    def export_history_to_docx(self):
        """Export to Word DOCX format"""
        try:
            if not getattr(self, "published_links", []):
                messagebox.showwarning("Canh bao", "Lich su trong!")
                return
            
            from docx import Document
            from docx.shared import Inches
            from datetime import datetime
            import os
            
            docs_folder = os.path.expanduser("~\\Documents")
            os.makedirs(docs_folder, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"DanhSachBaiViet_{timestamp}.docx"
            file_path = os.path.join(docs_folder, filename)
            
            doc = Document()
            doc.add_heading("Danh Sách Bài Viết", 0)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header: LINK FIRST, TITLE SECOND
            header_cells = table.rows[0].cells
            header_cells[0].text = "Link Bài Viết"
            header_cells[1].text = "Tiêu Đề"
            
            # Data: LINK FIRST, TITLE SECOND
            for item in self.published_links:
                if isinstance(item, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.get('link', 'N/A')
                    row_cells[1].text = item.get('title', 'N/A')
            
            doc.save(file_path)
            self.log(f"✅ Đã xuất Word: {filename}")
            messagebox.showinfo("Thành công", f"Đã xuất {len(self.published_links)} link!\n\nFile: {filename}")
            
            try:
                os.startfile(file_path)
            except:
                pass
                
        except ImportError:
            messagebox.showerror("Lỗi", "Cần cài đặt python-docx: pip install python-docx")
        except Exception as e:
            self.log(f"❌ Lỗi xuất Word: {e}")
            messagebox.showerror("Lỗi", f"Lỗi xuất Word: {e}")

    def export_history_to_odt(self):
        """Export to OpenDocument ODT format"""
        try:
            if not getattr(self, "published_links", []):
                messagebox.showwarning("Canh bao", "Lich su trong!")
                return
            
            from odf.opendocument import OpenDocumentText
            from odf.style import Style, TextProperties, ParagraphProperties, TableColumnProperties
            from odf.text import P, H
            from odf.table import Table, TableColumn, TableRow, TableCell
            from datetime import datetime
            import os
            
            docs_folder = os.path.expanduser("~\\Documents")
            os.makedirs(docs_folder, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"DanhSachBaiViet_{timestamp}.odt"
            file_path = os.path.join(docs_folder, filename)
            
            doc = OpenDocumentText()
            
            # Title
            h = H(outlinelevel=1, text="Danh Sách Bài Viết")
            doc.text.addElement(h)
            
            # Table
            table = Table()
            
            # Columns
            table.addElement(TableColumn())
            table.addElement(TableColumn())
            
            # Header row: LINK FIRST, TITLE SECOND
            tr = TableRow()
            tc = TableCell()
            tc.addElement(P(text="Link Bài Viết"))
            tr.addElement(tc)
            tc = TableCell()
            tc.addElement(P(text="Tiêu Đề"))
            tr.addElement(tc)
            table.addElement(tr)
            
            # Data rows: LINK FIRST, TITLE SECOND
            for item in self.published_links:
                if isinstance(item, dict):
                    tr = TableRow()
                    tc = TableCell()
                    tc.addElement(P(text=item.get('link', 'N/A')))
                    tr.addElement(tc)
                    tc = TableCell()
                    tc.addElement(P(text=item.get('title', 'N/A')))
                    tr.addElement(tc)
                    table.addElement(tr)
            
            doc.text.addElement(table)
            doc.save(file_path)
            
            self.log(f"✅ Đã xuất ODT: {filename}")
            messagebox.showinfo("Thành công", f"Đã xuất {len(self.published_links)} link!\n\nFile: {filename}")
            
            try:
                os.startfile(file_path)
            except:
                pass
                
        except ImportError:
            messagebox.showerror("Lỗi", "Cần cài đặt odfpy: pip install odfpy")
        except Exception as e:
            self.log(f"❌ Lỗi xuất ODT: {e}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Lỗi", f"Lỗi xuất ODT: {e}")

    def import_csv_for_export(self):
        """Import file with posted links (supports CSV, Excel, Word, ODT)"""
        try:
            import csv
            from pathlib import Path
            
            file_path = filedialog.askopenfilename(
                filetypes=[
                    ("CSV files", "*.csv"),
                    ("Excel files", "*.xlsx"),
                    ("Word files", "*.docx"),
                    ("OpenDocument", "*.odt"),
                    ("All files", "*.*")
                ],
                title="Chọn file có danh sách link đã đăng"
            )
            
            if not file_path:
                return
            
            ext = Path(file_path).suffix.lower()
            self.log(f"📥 Đang import {ext.upper()}: {os.path.basename(file_path)}")
            
            # Store the imported file format for export
            self._imported_file_format = ext
            
            imported_links = []
            skipped_count = 0
            
            # CSV
            if ext == ".csv":
                with open(file_path, 'r', encoding='utf-8-sig') as f:
                    reader = csv.reader(f)
                    header = next(reader, None)  # Skip header
                    
                    for row in reader:
                        if not row or len(row) < 2:
                            continue
                        
                        # NEW ORDER: Link first (col 0), Title second (col 1)
                        link = row[0].strip()
                        title = row[1].strip()
                        
                        # Skip empty rows or rows without valid links
                        if not title or not link:
                            skipped_count += 1
                            continue
                        
                        # Validate link (must contain http or domain)
                        if 'http' not in link.lower() and '.' not in link:
                            skipped_count += 1
                            self.log(f"⚠️ Bỏ qua dòng không hợp lệ: {link[:50]}")
                            continue
                        
                        imported_links.append({
                            'title': title,
                            'link': link,
                            'timestamp': '',
                            'status': 'imported'
                        })
            
            # Excel
            elif ext == ".xlsx":
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path)
                    ws = wb.active
                    
                    for row in ws.iter_rows(min_row=2, values_only=True):  # Skip header
                        if not row or len(row) < 2:
                            continue
                        
                        # NEW ORDER: Link first (col 0), Title second (col 1)
                        link = str(row[0]).strip() if row[0] else ""
                        title = str(row[1]).strip() if row[1] else ""
                        
                        # Skip empty or invalid
                        if not title or not link:
                            skipped_count += 1
                            continue
                        if 'http' not in link.lower() and '.' not in link:
                            skipped_count += 1
                            self.log(f"⚠️ Bỏ qua dòng không hợp lệ: {link[:50]}")
                            continue
                        
                        imported_links.append({
                            'title': title,
                            'link': link,
                            'timestamp': '',
                            'status': 'imported'
                        })
                except ImportError:
                    messagebox.showerror("Lỗi", "Cần cài đặt openpyxl: pip install openpyxl")
                    return
            
            # Word DOCX
            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(file_path)
                    
                    # Try to find table
                    if doc.tables:
                        table = doc.tables[0]
                        for i, row in enumerate(table.rows):
                            if i == 0:  # Skip header
                                continue
                            cells = row.cells
                            if len(cells) < 2:
                                continue
                            
                            # NEW ORDER: Link first (col 0), Title second (col 1)
                            link = cells[0].text.strip()
                            title = cells[1].text.strip()
                            
                            # Skip empty or invalid
                            if not title or not link:
                                continue
                            if 'http' not in link.lower() and '.' not in link:
                                continue
                            
                            imported_links.append({
                                'title': title,
                                'link': link,
                                'timestamp': '',
                                'status': 'imported'
                            })
                except ImportError:
                    messagebox.showerror("Lỗi", "Cần cài đặt python-docx: pip install python-docx")
                    return
            
            # OpenDocument ODT
            elif ext == ".odt":
                try:
                    from odf import text, teletype, table
                    from odf.opendocument import load
                    
                    doc = load(file_path)
                    
                    # Try to read from table first
                    tables = doc.getElementsByType(table.Table)
                    if tables:
                        tbl = tables[0]
                        rows = tbl.getElementsByType(table.TableRow)
                        for i, row in enumerate(rows):
                            if i == 0:  # Skip header
                                continue
                            cells = row.getElementsByType(table.TableCell)
                            if len(cells) < 2:
                                continue
                            
                            # NEW ORDER: Link first (col 0), Title second (col 1)
                            link = teletype.extractText(cells[0]).strip()
                            title = teletype.extractText(cells[1]).strip()
                            
                            # Skip empty or invalid
                            if not title or not link:
                                continue
                            if 'http' not in link.lower() and '.' not in link:
                                continue
                            
                            imported_links.append({
                                'title': title,
                                'link': link,
                                'timestamp': '',
                                'status': 'imported'
                            })
                    else:
                        # Fallback: read paragraphs
                        paragraphs = doc.getElementsByType(text.P)
                        lines = [teletype.extractText(p).strip() for p in paragraphs if teletype.extractText(p).strip()]
                        
                        for i in range(0, len(lines)-1, 2):
                            title = lines[i]
                            link = lines[i+1] if i+1 < len(lines) else ""
                            
                            # Skip empty or invalid
                            if not title or not link:
                                continue
                            if 'http' not in link.lower() and '.' not in link:
                                continue
                            
                            imported_links.append({
                                'title': title,
                                'link': link,
                                'timestamp': '',
                                'status': 'imported'
                            })
                except ImportError:
                    messagebox.showerror("Lỗi", "Cần cài đặt odfpy: pip install odfpy")
                    return
            
            else:
                messagebox.showwarning("Không hỗ trợ", f"Định dạng {ext} chưa được hỗ trợ")
                return
            
            if not imported_links:
                self.log("⚠️ File không có dữ liệu hợp lệ")
                self.log("💡 Kiểm tra: Cột 1 = Link (có http://), Cột 2 = Tiêu đề")
                messagebox.showwarning("Trống", "File không có dữ liệu hợp lệ!\n\nĐảm bảo:\n- Cột 1: Link bài viết (có http://)\n- Cột 2: Tiêu đề bài viết\n- Có ít nhất 1 dòng dữ liệu (ngoài header)")
                return
            
            # Add to published_links
            if not hasattr(self, 'published_links'):
                self.published_links = []
            
            self.published_links.extend(imported_links)
            
            # Update history display
            if hasattr(self, 'history_textbox'):
                self.history_textbox.configure(state="normal")
                for item in imported_links:
                    self.history_textbox.insert("end", f"📥 {item['link']}\n   📝 {item['title']}\n\n")
                self.history_textbox.configure(state="disabled")
            
            format_name = {"csv": "CSV", "xlsx": "Excel", "docx": "Word", "odt": "ODT"}.get(ext[1:], ext.upper())
            self.log(f"✅ Đã import {len(imported_links)} link từ {format_name}")
            if skipped_count > 0:
                self.log(f"⚠️ Đã bỏ qua {skipped_count} dòng không hợp lệ")
            self.log(f"💡 Khi xuất sẽ tự động dùng định dạng {format_name}")
            
            msg = f"Đã import {len(imported_links)} link từ {format_name}!"
            if skipped_count > 0:
                msg += f"\n\n⚠️ Đã bỏ qua {skipped_count} dòng không hợp lệ"
            msg += f"\n\nKhi xuất sẽ tự động dùng định dạng {format_name}."
            messagebox.showinfo("Thành công", msg)
            
        except Exception as e:
            self.log(f"❌ Lỗi import: {e}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Lỗi", f"Lỗi import file: {e}")

    def clear_history(self):
        """Clear all history"""
        self.published_links = []
        if hasattr(self, 'history_textbox'):
            self.history_textbox.configure(state="normal")
            self.history_textbox.delete("1.0", "end")
            self.history_textbox.configure(state="disabled")
        self.log("🗑️ Đã xóa toàn bộ lịch sử.")

    def browse_video_upload(self):
        filenames = filedialog.askopenfilenames(
            title="Chọn Video Upload",
            filetypes=[("Video Files", "*.mp4 *.mkv *.avi *.mov *.wmv"), ("All Files", "*.*")]
        )
        if filenames:
            paths_str = "; ".join(filenames)
            self.entry_upload_path.delete(0, "end")
            self.entry_upload_path.insert(0, paths_str)
            self.log(f"Đã chọn {len(filenames)} file video.")

    def browse_thumbnail(self):
        """Open file dialog to select a thumbnail image"""
        filename = filedialog.askopenfilename(
            title="Chọn Ảnh Thumbnail",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png *.webp"), ("All Files", "*.*")]
        )
        if filename:
            self.entry_image.delete(0, "end")
            self.entry_image.insert(0, filename)
            self.log(f"🖼️ Đã chọn ảnh thumbnail: {os.path.basename(filename)}")

    def browse_content_image(self):
        """Open file dialog to select a content image"""
        filename = filedialog.askopenfilename(
            title="Chọn Ảnh Content 1",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png *.webp"), ("All Files", "*.*")]
        )
        if filename:
            self.entry_content_image.delete(0, "end")
            self.entry_content_image.insert(0, filename)
            self.log(f"🖼️ Đã chọn ảnh content 1: {os.path.basename(filename)}")
    
    def browse_content_image2(self):
        """Open file dialog to select content image 2"""
        filename = filedialog.askopenfilename(
            title="Chọn Ảnh Content 2",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png *.webp"), ("All Files", "*.*")]
        )
        if filename:
            self.entry_content_image2.delete(0, "end")
            self.entry_content_image2.insert(0, filename)
            self.log(f"🖼️ Đã chọn ảnh content 2: {os.path.basename(filename)}")
    
    def browse_content_image3(self):
        """Open file dialog to select content image 3"""
        filename = filedialog.askopenfilename(
            title="Chọn Ảnh Content 3",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png *.webp"), ("All Files", "*.*")]
        )
        if filename:
            self.entry_content_image3.delete(0, "end")
            self.entry_content_image3.insert(0, filename)
            self.log(f"🖼️ Đã chọn ảnh content 3: {os.path.basename(filename)}")

    def paste_image_from_clipboard(self, event=None):
        """Paste image from clipboard and save to thumbnails folder (for thumbnail field)"""
        return self.paste_image_from_clipboard_to_field(event, self.entry_image, "thumb")
    
    def paste_image_from_clipboard_to_field(self, event, target_field, field_name):
        """Paste image from clipboard and save to thumbnails folder (generic method)"""
        try:
            # Try to get image from clipboard
            img = ImageGrab.grabclipboard()
            
            if img is None:
                self.log("⚠️ Clipboard không chứa ảnh. Hãy chụp màn hình (PrtScn) hoặc copy ảnh trước.")
                return "break"
            
            # Create thumbnails folder if not exists
            if not os.path.exists("thumbnails"):
                os.makedirs("thumbnails")
            
            # Generate filename with timestamp
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            filename = f"thumbnails/pasted_{field_name}_{timestamp}.png"
            
            # Save image
            img.save(filename, "PNG")
            
            # Update the target field
            target_field.delete(0, "end")
            target_field.insert(0, filename)
            
            self.log(f"✅ Đã paste ảnh từ clipboard: {filename}")
            
            # Prevent default paste behavior
            return "break"
            
        except Exception as e:
            self.log(f"❌ Lỗi khi paste ảnh: {e}")
            return "break"


    def add_uploaded_video_link(self, title, embed_code, thumb=None):
        if hasattr(self, 'txt_upload_list'):
            self.txt_upload_list.configure(state="normal")
            display_text = f"✅ [UPLOADED] {title}\n"
            display_text += f"🔗 Embed Code:\n{embed_code}\n"
            if thumb:
                display_text += f"🖼 Thumbnail: {thumb}\n"
            display_text += "-"*50 + "\n"
            
            self.txt_upload_list.insert("end", display_text)
            self.txt_upload_list.see("end")
            self.txt_upload_list.configure(state="disabled")

    def on_upload_click(self):
        files_str = self.entry_upload_path.get().strip()
        if not files_str:
            self.log("❌ Chưa chọn file video!")
            return

        files = files_str.split("; ")
        self.log(f"Bắt đầu upload {len(files)} video...")
        self.btn_upload.configure(state="disabled", text="⏳ Đang tải lên...")
        
        import threading
        from model.vimeo_helper import VimeoHelper
        import queue
        
        # Create a queue for results
        result_queue = queue.Queue()
        completed_count = [0]  # Use list to allow modification in nested function
        total_files = len(files)

        def _upload_single_video(file_path, file_index, account_index):
            """Upload a single video with its own browser instance"""
            helper = VimeoHelper()
            try:
                # Load all available accounts first (important for rotation)
                helper.load_all_available_accounts()
                
                # Check headless option
                use_headless = bool(self.chk_headless_upload.get())
                mode_text = "Headless - Ẩn" if use_headless else "Visible - Hiện"
                self.after(0, lambda idx=file_index: self.log(f"[Luồng {account_index+1}] --- Upload File {idx+1}/{total_files}: {os.path.basename(file_path)} ---"))
                
                helper.init_driver(headless=use_headless)
                
                # Log callback for helper
                def _log(m): 
                    self.after(0, lambda msg=m, idx=file_index: self.log(f"[File {idx+1}] {msg}"))

                # Upload
                success, msg, data, quota = helper.upload_video(file_path, log_callback=_log)
                
                result_queue.put({
                    'index': file_index,
                    'success': success,
                    'msg': msg,
                    'data': data,
                    'quota': quota,
                    'file_path': file_path
                })
                
            except Exception as e:
                self.after(0, lambda err=e, idx=file_index: self.log(f"[File {idx+1}] Lỗi Thread: {err}"))
                result_queue.put({
                    'index': file_index,
                    'success': False,
                    'msg': str(e),
                    'data': None,
                    'quota': False,
                    'file_path': file_path
                })
            finally:
                helper.close()

        def _process_results():
            """Process upload results from queue"""
            try:
                result = result_queue.get_nowait()
                
                file_path = result['file_path']
                success = result['success']
                msg = result['msg']
                data = result['data']
                quota = result['quota']
                idx = result['index']
                
                if success:
                    # Check if we have data
                    if data:
                        title = data.get("title", "Video")
                        embed = data.get("embed_code", "No Embed Code")
                        thumb = data.get("thumbnail", None)
                        video_link = data.get("video_link", "")
                    else:
                        # No data but success - create basic info from filename
                        title = os.path.basename(file_path)
                        embed = f"<!-- Video đang xử lý: {title} -->"
                        thumb = None
                        video_link = ""
                        self.log("⚠️ Video đang xử lý, chưa có embed code đầy đủ")
                    
                    self.log(f"✅ {msg}")
                    self.add_uploaded_video_link(title, embed, thumb)
                    
                    # Auto add to queue if option is checked
                    if self.chk_auto_add_queue.get():
                        self.add_video_to_queue(title, embed, video_link, thumb)
                else:
                    self.log(f"❌ Lỗi: {msg}")
                    if quota:
                        self.log("⚠️ Quota Full!")
                
                completed_count[0] += 1
                
                # Check if all done
                if completed_count[0] >= total_files:
                    self.log(f"=== Hoàn thành upload {completed_count[0]}/{total_files} video ===")
                    self.btn_upload.configure(state="normal", text="⬆️ Bắt đầu Upload")
                else:
                    # Continue checking for more results
                    self.after(500, _process_results)
                    
            except queue.Empty:
                # No result yet, check again later
                if completed_count[0] < total_files:
                    self.after(500, _process_results)
                else:
                    self.btn_upload.configure(state="normal", text="⬆️ Bắt đầu Upload")

        # Determine number of available accounts
        total_accounts = 1
        try:
            if os.path.exists("vimeo_accounts.txt"):
                with open("vimeo_accounts.txt", "r", encoding="utf-8") as f:
                    acc_lines = [l for l in f.readlines() if l.strip()]
                    if acc_lines:
                        total_accounts = len(acc_lines)
            elif os.path.exists("vimeo_cookies.txt"):
                with open("vimeo_cookies.txt", "r", encoding="utf-8") as f:
                    c_lines = [l for l in f.readlines() if l.strip()]
                    if c_lines:
                        total_accounts = len(c_lines)
        except:
            pass
            
        self.log(f"🔎 Tìm thấy {total_accounts} tài khoản Vimeo để xoay vòng.")

        # Start parallel uploads with limited workers (2-3 threads max)
        max_parallel = 3  # Giới hạn 3 luồng song song
        self.log(f"🚀 Chạy song song tối đa {max_parallel} luồng upload...")
        
        # Use ThreadPoolExecutor to limit concurrent uploads
        from concurrent.futures import ThreadPoolExecutor
        
        def submit_upload(idx, file_path):
            """Submit a single upload task"""
            file_path = file_path.strip()
            if not file_path:
                completed_count[0] += 1
                return
            
            # Assign account index (rotate through ALL accounts)
            account_index = idx % total_accounts
            
            # Execute upload
            _upload_single_video(file_path, idx, account_index)
        
        # Create thread pool and submit all tasks
        executor = ThreadPoolExecutor(max_workers=max_parallel)
        
        for idx, file_path in enumerate(files):
            executor.submit(submit_upload, idx, file_path)
        
        # Don't wait for completion here - let them run in background
        executor.shutdown(wait=False)
        
        # Start result processor
        self.after(1000, _process_results)

    def add_video_to_queue(self, title, embed_code, video_link, thumbnail_path=None):
        """Add uploaded video to post queue automatically with thumbnail (content will be added when posting)"""
        try:
            # Clean title (remove file extension if present)
            clean_title = title.replace('.mp4', '').replace('.avi', '').replace('.mov', '').replace('.mkv', '')
            clean_title = clean_title.replace('_', ' ').strip()
            
            # Create post data with thumbnail (NO content yet)
            post_data = {
                'title': clean_title,
                'video_url': embed_code if embed_code else video_link,
                'image_url': thumbnail_path if thumbnail_path else '',
                'content': '',
                'needs_body_content': True
            }
            
            # Thread-safe: schedule append + UI update on main thread
            def _add():
                self.post_queue.append(post_data)
                self.update_queue_display()
            self.after(0, _add)
            self.log(f"📝 Đã thêm vào hàng chờ: {clean_title}")
            self.log(f"   🔗 Video: {video_link if video_link else 'Processing...'}")
            if thumbnail_path:
                self.log(f"   🖼️ Thumbnail: {os.path.basename(thumbnail_path)}")
            
            # Show content pool status
            if hasattr(self, 'content_pool'):
                remaining = len(self.content_pool)
                if remaining > 0:
                    self.log(f"   💾 Có {remaining} nội dung sẵn sàng để ghép khi đăng bài")
                else:
                    self.log(f"   ℹ️ Chưa có nội dung body (chỉ có video embed)")
            
        except Exception as e:
            self.log(f"❌ Lỗi thêm vào hàng chờ: {e}")

    def log_vimeo(self, message):
        """Log specifically to Vimeo tab"""
        # Update Main Log
        self.log(f"[VIMEO] {message}")
        
        # Update Vimeo Tab Log
        if hasattr(self, 'txt_vimeo_log'):
            self.txt_vimeo_log.configure(state="normal")
            self.txt_vimeo_log.insert("end", f"[{time.strftime('%H:%M:%S')}] {message}\n")
            self.txt_vimeo_log.see("end")
            self.txt_vimeo_log.configure(state="disabled")

    def test_vimeo_cookie_login(self):
        """Test Vimeo cookie login functionality"""
        self.log_vimeo("🍪 Bắt đầu test cookie login...")
        self.btn_test_cookie.configure(state="disabled", text="⏳ Đang test...")
        
        import threading
        from model.vimeo_helper import VimeoHelper
        
        def _test_login():
            helper = VimeoHelper()
            try:
                # Check headless option
                use_headless = bool(self.chk_headless_test.get())
                helper.init_driver(headless=use_headless)
                
                # Log callback
                def _log(msg): 
                    self.after(0, lambda m=msg: self.log_vimeo(m))
                
                # Test cookie login
                success = helper.auto_login(log_callback=_log)
                
                if success:
                    self.after(0, lambda: self.log_vimeo("✅ Cookie login thành công!"))
                    
                    # Test upload page access
                    self.after(0, lambda: self.log_vimeo("🔍 Kiểm tra quyền truy cập upload..."))
                    helper.driver.get("https://vimeo.com/upload")
                    time.sleep(3)
                    
                    if "upload" in helper.driver.current_url.lower():
                        self.after(0, lambda: self.log_vimeo("✅ Có thể truy cập trang upload!"))
                    else:
                        self.after(0, lambda: self.log_vimeo("❌ Không thể truy cập trang upload."))
                        
                else:
                    self.after(0, lambda: self.log_vimeo("❌ Cookie login thất bại!"))
                    
            except Exception as e:
                self.after(0, lambda err=e: self.log_vimeo(f"❌ Lỗi test: {err}"))
            finally:
                helper.close()
                self.after(0, lambda: self.btn_test_cookie.configure(state="normal", text="🍪 Test Cookie Login"))
        
        threading.Thread(target=_test_login, daemon=True).start()

    def on_vimeo_reg_click(self):
        try:
            count = int(self.entry_vm_count.get())
        except:
            self.log("❌ Số lượng không hợp lệ (nhập số nguyên).")
            return

        self.log_vimeo(f"Bắt đầu quy trình tạo {count} tài khoản Vimeo...")
        self.btn_vm_reg.configure(state="disabled", text="⏳ Đang chạy...")
        
        import threading
        import random
        import string
        from model.vimeo_helper import VimeoHelper

        def _generate_identity():
            name = f"User{random.randint(10000, 99999)}"
            
            # Luôn dùng @gmail.com - guerrillamailblock bị Vimeo block
            rand_user = ''.join(random.choices(string.ascii_lowercase + string.digits, k=10))
            email = f"{rand_user}@gmail.com"
            
            p_upper = random.choice(string.ascii_uppercase)
            p_lower = random.choice(string.ascii_lowercase)
            p_digits = ''.join(random.choices(string.digits, k=3))
            pwd = f"{p_upper}{p_lower}@{p_digits}xyz"
            return name, email, pwd

        def _clear_brave_profile():
            """Xóa profile Brave cũ để tránh bị track"""
            import shutil, tempfile, glob
            try:
                tmp = tempfile.gettempdir()
                for d in glob.glob(os.path.join(tmp, f"chrome_profile_*")):
                    try:
                        shutil.rmtree(d, ignore_errors=True)
                    except: pass
                brave_cache = os.path.join(os.environ.get('LOCALAPPDATA',''), 
                                           r"BraveSoftware\Brave-Browser\User Data\Default\Cache")
                if os.path.exists(brave_cache):
                    shutil.rmtree(brave_cache, ignore_errors=True)
                print("[BRAVE] 🧹 Đã xóa profile/cache cũ")
            except Exception as e:
                print(f"[BRAVE] ⚠️ Xóa profile lỗi: {e}")

        # Load proxy list khi bắt đầu
        current_proxy = [None]  # dùng list để có thể modify trong nested function
        try:
            proxies = VimeoHelper.fetch_free_proxies()
            if proxies:
                self.after(0, lambda n=len(proxies): self.log_vimeo(f"🌐 Loaded {n} free proxies"))
        except: pass

        def _get_working_proxy():
            """Lấy proxy tiếp theo - không test, dùng thẳng"""
            p = VimeoHelper.get_next_proxy()
            return p

        def _run_batch():
            created = 0
            consecutive_failures = 0
            ip_blocked = False
            
            for i in range(count):
                self.after(0, lambda idx=i: self.log_vimeo(f"--- TÀI KHOẢN {idx+1}/{count} ---"))
                
                # 1. Gen Info
                name, email, pwd = _generate_identity()
                self.after(0, lambda e=email: self.log_vimeo(f"📧 {e}"))
                
                # 2. Run Helper
                helper = VimeoHelper()
                try:
                    # _cb chỉ hiện các message quan trọng (có ✅ ❌ ⚠️), bỏ log chi tiết
                    def _cb(msg):
                        if any(s in msg for s in ['✅', '❌', '⚠️', '🎣', 'Thành công', 'Thất bại', 'SAVE', 'API-REG']):
                            self.after(0, lambda m=msg: self.log_vimeo(m))

                    # ── Thử Pure API ──
                    api_ok, api_data = helper.register_pure_api(name, email, pwd, log_callback=_cb)

                    if api_ok:
                        helper._save_account_from_api(email, pwd, name, api_data, _cb)
                        success, msg = True, "SUCCESS_VIA_PURE_API"
                    else:
                        # Fallback: Selenium với Brave + proxy
                        use_headless = bool(self.chk_headless_vimeo.get())
                        proxy = current_proxy[0]
                        if proxy:
                            self.after(0, lambda p=proxy: self.log_vimeo(f"🌐 Dùng proxy: {p}"))
                        helper.init_driver(headless=use_headless, browser_type='brave', proxy=proxy)
                        success, msg = helper.fill_registration_form(name, email, pwd, log_callback=_cb)
                    
                    if success:
                        self.after(0, lambda: self.log_vimeo("✅ Tạo thành công!"))
                        created += 1
                        consecutive_failures = 0
                    else:
                        self.after(0, lambda m=msg: self.log_vimeo(f"❌ Thất bại: {m}"))
                        
                        # Khi bị block/fail → đổi proxy + xóa profile
                        if msg in ["IP_BLOCKED", "RATE_LIMITED", "ACCOUNT_LIMIT", "NETWORK_ERROR",
                                   "TIMEOUT_JOIN_PAGE", "WRONG_BUTTON_GOOGLE"]:
                            consecutive_failures += 1
                            _clear_brave_profile()
                            VimeoHelper._shared_vimeo_cookies = None  # reset shared cookies
                            
                            if msg in ["IP_BLOCKED", "RATE_LIMITED"]:
                                # Đổi proxy ngay
                                self.after(0, lambda: self.log_vimeo("� IP bị chặn → đổi proxy..."))
                                new_proxy = _get_working_proxy()
                                if new_proxy:
                                    current_proxy[0] = new_proxy
                                    self.after(0, lambda p=new_proxy: self.log_vimeo(f"✅ Proxy mới: {p}"))
                                else:
                                    self.after(0, lambda: self.log_vimeo("⚠️ Không tìm được proxy hoạt động"))
                                    if consecutive_failures >= 3:
                                        ip_blocked = True
                                        self.after(0, lambda: self.log_vimeo("🛑 DỪNG: Hết proxy!"))
                                        self.after(0, lambda: self.show_ip_block_warning())
                                        break
                        else:
                            consecutive_failures = 0
                        
                except Exception as e:
                    self.after(0, lambda err=e: self.log_vimeo(f"Lỗi: {err}"))
                finally:
                    helper.close()
            
            # Final summary
            if ip_blocked:
                self.after(0, lambda: self.log_vimeo(f"⚠️ KẾT THÚC SỚM: Tạo được {created}/{count} TK (IP bị chặn)"))
            else:
                self.after(0, lambda: self.log_vimeo(f"=== KẾT THÚC: Tạo được {created}/{count} TK ==="))
            
            self.after(0, lambda: self.btn_vm_reg.configure(state="normal", text="🚀 Bắt đầu Tạo"))

        threading.Thread(target=_run_batch, daemon=True).start()

    def show_ip_block_warning(self):
        """Show popup warning when IP is blocked"""
        try:
            messagebox.showwarning(
                "⚠️ IP Bị Chặn",
                "🚫 Vimeo đã chặn IP của bạn!\n\n"
                "📌 Nguyên nhân:\n"
                "• Tạo quá nhiều tài khoản từ cùng 1 IP\n"
                "• IP bị đánh dấu spam/bot\n"
                "• Rate limit (quá nhiều request)\n\n"
                "✅ Giải pháp:\n"
                "1. Đổi IP/VPN khác\n"
                "2. Đợi 30-60 phút rồi thử lại\n"
                "3. Sử dụng proxy xoay (rotating proxy)\n\n"
                "💡 Khuyến nghị: Đổi VPN/Proxy ngay!"
            )
        except Exception as e:
            print(f"[GUI] Error showing popup: {e}")
    
    def show_about_dialog(self):
        """Show About dialog with version info and changelog"""
        try:
            # Import version info
            try:
                from version import get_version_info, __app_name__, __author__, __build_date__
                version_info = get_version_info()
            except:
                version_info = {
                    "version": "2.0.2",
                    "app_name": "WprTool - WordPress Auto Posting",
                    "author": "NguyenDuyDuc",
                    "build_date": "2026-02-05",
                    "features": []
                }
            
            # Create popup window
            popup = ctk.CTkToplevel(self)
            popup.title("Thông Tin Phiên Bản")
            popup.geometry("500x400")
            popup.resizable(False, False)
            
            # Center window
            popup.update_idletasks()
            x = (popup.winfo_screenwidth() // 2) - (500 // 2)
            y = (popup.winfo_screenheight() // 2) - (400 // 2)
            popup.geometry(f"500x400+{x}+{y}")
            
            # Main container
            main_frame = ctk.CTkFrame(popup, fg_color=self.colors['bg_light'])
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Header with icon and title
            header = ctk.CTkFrame(main_frame, fg_color=self.colors['primary'], corner_radius=12)
            header.pack(fill="x", pady=(0, 20))
            
            ctk.CTkLabel(
                header,
                text="🚀",
                font=("Segoe UI", 48)
            ).pack(pady=(20, 5))
            
            ctk.CTkLabel(
                header,
                text=version_info['app_name'],
                font=("Segoe UI", 20, "bold"),
                text_color="white"
            ).pack(pady=(0, 5))
            
            ctk.CTkLabel(
                header,
                text=f"Version {version_info['version']}",
                font=("Segoe UI", 14),
                text_color="white"
            ).pack(pady=(0, 20))
            
            # Info section
            info_frame = ctk.CTkFrame(main_frame, fg_color=self.colors['bg_card'], corner_radius=12)
            info_frame.pack(fill="x", pady=(0, 15))
            
            info_items = [
                # ("👤 Tác giả", version_info['author']),
                # ("📅 Ngày phát hành", version_info['build_date']),
                ("Build", version_info['version']),
            ]
            
            for label, value in info_items:
                row = ctk.CTkFrame(info_frame, fg_color="transparent")
                row.pack(fill="x", padx=15, pady=8)
                
                ctk.CTkLabel(
                    row,
                    text=label,
                    font=("Segoe UI", 11, "bold"),
                    text_color=self.colors['text_secondary']
                ).pack(side="left")
                
                ctk.CTkLabel(
                    row,
                    text=value,
                    font=("Segoe UI", 11),
                    text_color=self.colors['text_primary']
                ).pack(side="right")
            
            # Footer with close button
            footer = ctk.CTkFrame(main_frame, fg_color="transparent")
            footer.pack(fill="x", pady=(20, 0))
            
            ctk.CTkButton(
                footer,
                text="Đóng",
                width=150,
                height=36,
                font=("Segoe UI", 12, "bold"),
                fg_color=self.colors['primary'],
                hover_color=self.colors['primary_hover'],
                corner_radius=10,
                command=popup.destroy
            ).pack(pady=5)
            
            # Make popup modal
            popup.transient(self)
            popup.grab_set()
            
        except Exception as e:
            print(f"[GUI] Error showing about dialog: {e}")
            import traceback
            traceback.print_exc()
    
    def copy_all_links_with_titles(self):
        """Copy all published posts with titles in format: Title - Link"""
        try:
            import pyperclip
            
            if not self.published_posts:
                messagebox.showinfo("Thông Báo", "Chưa có bài viết nào được đăng!")
                return
            
            # Format: Title - Link (one per line)
            formatted_text = ""
            for post in self.published_posts:
                title = post.get('title', 'Untitled')
                link = post.get('link', '')
                formatted_text += f"{title} - {link}\n"
            
            # Copy to clipboard
            pyperclip.copy(formatted_text.strip())
            
            # Show success message
            count = len(self.published_posts)
            messagebox.showinfo(
                "Thành Công", 
                f"✅ Đã copy {count} bài viết!\n\nĐịnh dạng:\nTiêu đề - Link"
            )
            
            print(f"[COPY] Copied {count} posts to clipboard")
            
        except Exception as e:
            print(f"[COPY] Error: {e}")
            messagebox.showerror("Lỗi", f"Không thể copy: {e}")
    
    def kill_chrome_processes(self):
        """Kill ONLY Chrome processes running from this application's folder"""
        try:
            import subprocess
            import os
            
            # Update status
            self.status_label.configure(text="Cleanup Chrome...", text_color=self.colors['warning'])
            self.update()
            
            # Project root
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            
            # Method 1: Use safe batch file
            bat_file = os.path.join(project_root, "kill_chrome.bat")
            
            log_msg = ""
            
            if os.path.exists(bat_file):
                try:
                    # Run batch file that filters by %~dp0
                    subprocess.run([bat_file], shell=True, capture_output=True, timeout=10)
                    log_msg = "Run kill_chrome.bat"
                except Exception as e:
                    log_msg = f"Bat error: {e}"
            else:
                # Fallback Safe PowerShell if bat missing
                try:
                    target_path_escaped = project_root.replace("\\", "\\\\").replace("'", "''")
                    ps_cmd = f"Get-WmiObject Win32_Process | Where-Object {{ ($_.Name -eq 'chrome.exe' -or $_.Name -eq 'chromedriver.exe') -and $_.ExecutablePath -like '*{target_path_escaped}*' }} | ForEach-Object {{ Stop-Process -Id $_.ProcessId -Force }}"
                    subprocess.run(["powershell", "-NoProfile", "-Command", ps_cmd], capture_output=True, timeout=10)
                    log_msg = "Run Safe PowerShell"
                except Exception as e:
                    log_msg = f"PS error: {e}"

            # Verify
            print(f"[KILL] {log_msg}")
            
            self.status_label.configure(
                text="✅ Đã đóng Chrome của Tool", 
                text_color=self.colors['success']
            )
            
            messagebox.showinfo(
                "Hoàn Thành",
                "Đã đóng các tiến trình Chrome/Driver của Tool.\n(Trình duyệt cá nhân của bạn không bị ảnh hưởng)"
            )
            
            # Reset status after 3s
            self.after(3000, lambda: self.status_label.configure(
                text="✅ Sẵn sàng", 
                text_color=self.colors['success']
            ))
            
        except Exception as e:
            print(f"[KILL] Critical Error: {e}")
            self.status_label.configure(text="❌ Lỗi", text_color=self.colors['danger'])
    
    def add_published_post(self, title, link):
        """Add a published post to the list for copying"""
        try:
            self.published_posts.append({
                'title': title,
                'link': link
            })
            print(f"[PUBLISHED] Added: {title} - {link}")
        except Exception as e:
            print(f"[PUBLISHED] Error adding post: {e}")

    def logout(self):
        """Đăng xuất và quay lại màn hình login"""
        # Confirm logout
        if messagebox.askyesno("Xác nhận", "Bạn có chắc muốn đăng xuất?"):
            # Destroy main screen
            if hasattr(self, 'main_frame'):
                self.main_frame.destroy()
            
            # Recreate login screen
            self.create_login_screen()
            
            self.log("👋 Đã đăng xuất")

# Mock Controller để test giao diện
class MockController:
    username = "AdminUser"
    def handle_login(self, s, u, p, headless):
        print(f"Login: {s}, {u}, Headless={headless}")
        # Giả lập login thành công sau 1s
        app.after(1000, app.login_success)
        
    def handle_post_request(self, data, is_batch=False):
        print(f"Posting: {data.title}")
        # Giả lập post thành công
        app.after(1500, lambda: app.on_post_finished(True, f"https://site.com/{data.title.replace(' ','-')}", is_batch))

if __name__ == "__main__":
    controller = MockController()
    app = GUIView(controller)
    app.mainloop()