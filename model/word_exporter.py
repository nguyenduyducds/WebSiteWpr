"""
Word Exporter - Xuất danh sách bài viết đã đăng ra file Word
"""

import os
from datetime import datetime


class WordExporter:
    """Export posted articles to Word document"""
    
    @staticmethod
    def export_posted_articles(articles, output_path=None, append=False):
        """
        Export list of posted articles to Word document (Simple format: Title + Link)
        
        Args:
            articles: List of dicts with keys: title, post_url
            output_path: Where to save the Word file
            append: If True, append to existing file instead of creating new
        
        Returns:
            Path to created Word file or None
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            
            if not output_path:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"DanhSachBaiViet_{timestamp}.docx"
            
            # Ensure .docx extension
            if not output_path.endswith('.docx'):
                output_path = output_path.replace('.odt', '.docx').replace('.doc', '.docx')
                if not output_path.endswith('.docx'):
                    output_path += '.docx'
            
            # Check if file exists and append mode is enabled
            if append and os.path.exists(output_path):
                print(f"[WORD] 📝 Appending to existing file: {output_path}")
                doc = Document(output_path)
                # Find the table (should be the last table)
                if doc.tables:
                    table = doc.tables[-1]
                else:
                    # Create table if doesn't exist
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid Accent 1'
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Tiêu Đề"
                    header_cells[1].text = "Link Bài Viết"
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            else:
                # Create new document
                print(f"[WORD] ✨ Creating new file: {output_path}")
                doc = Document()
                
                # Title
                title = doc.add_heading("📋 Danh Sách Bài Viết Đã Đăng", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Info
                info_para = doc.add_paragraph()
                from datetime import datetime
                info_para.add_run(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
                info_para.add_run(f"Tổng số bài: {len(articles)}")
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Spacer
                
                # Table with 2 columns: Title + Link
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Tiêu Đề"
                header_cells[1].text = "Link Bài Viết"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Add data rows
            for article in articles:
                row_cells = table.add_row().cells
                row_cells[0].text = article.get('title', 'N/A')
                row_cells[1].text = article.get('post_url', 'N/A')
            
            # Set column widths
            widths = [Inches(3.0), Inches(3.5)]
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            
            # Add footer only if new document
            if not append or not os.path.exists(output_path):
                doc.add_paragraph()
                footer = doc.add_paragraph()
                footer.add_run("✅ Tất cả bài viết đã được đăng thành công!").font.italic = True
                footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save - CRITICAL: Use .docx format
            doc.save(output_path)
            print(f"[WORD] ✅ Exported to: {output_path}")
            return output_path
            
        except ImportError as e:
            print(f"[WORD] ⚠️ python-docx not installed or import error: {e}")
            print("[WORD] Install: pip install python-docx")
            return None
        except Exception as e:
            print(f"[WORD] ❌ Error exporting: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    @staticmethod
    def export_scan_results(videos, output_path=None):
        """
        Export scan results to Word document (for import later)
        
        Args:
            videos: List of dicts with keys: url, title, embed_code, thumbnail, video_id
            output_path: Where to save the Word file
        
        Returns:
            Path to created Word file or None
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"scan_results_{timestamp}.docx"
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading("📊 Kết Quả Scan Video", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Ngày scan: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            info_para.add_run(f"Tổng số video: {len(videos)}")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacer
            
            # Table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "STT"
            header_cells[1].text = "Tiêu Đề"
            header_cells[2].text = "Link Video"
            header_cells[3].text = "Video ID"
            header_cells[4].text = "Trạng Thái"
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for idx, video in enumerate(videos, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = video.get('title', 'N/A')
                row_cells[2].text = video.get('url', 'N/A')
                row_cells[3].text = video.get('video_id', 'N/A')
                row_cells[4].text = video.get('status', 'pending')
            
            # Set column widths
            widths = [Inches(0.5), Inches(2.0), Inches(2.5), Inches(1.5), Inches(1.0)]
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.add_run("💡 Bạn có thể import lại file này để tiếp tục làm việc").font.italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save
            doc.save(output_path)
            print(f"[WORD] ✅ Exported to: {output_path}")
            return output_path
            
        except ImportError:
            print("[WORD] ⚠️ python-docx not installed. Install: pip install python-docx")
            return None
        except Exception as e:
            print(f"[WORD] ❌ Error exporting: {e}")
            return None


# Test
if __name__ == "__main__":
    exporter = WordExporter()
    
    # Test data
    test_articles = [
        {
            'title': 'Watch full rescues - Bài viết 1',
            'url': 'https://vimeo.com/123456789',
            'post_url': 'https://example.com/post/1',
            'thumbnail': '/path/to/thumb1.jpg',
            'status': 'published'
        },
        {
            'title': 'Watch full rescues - Bài viết 2',
            'url': 'https://vimeo.com/987654321',
            'post_url': 'https://example.com/post/2',
            'thumbnail': '/path/to/thumb2.jpg',
            'status': 'published'
        }
    ]
    
    # Export
    result = exporter.export_posted_articles(test_articles)
    if result:
        print(f"✅ Test export successful: {result}")
