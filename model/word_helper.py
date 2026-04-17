"""
Word Helper - Đọc link từ file Word (.docx)
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple, Optional, Dict


_URL_RE = re.compile(r"(https?://[^\s<>\"]+)", re.IGNORECASE)


class WordHelper:
    """Helper để đọc link từ file Word (.docx)."""

    @staticmethod
    def read_docx(file_path: str) -> Tuple[Optional[List[Dict[str, str]]], Optional[str]]:
        """
        Trả về (videos, error)
        videos: [{ 'url': str, 'title': str }]
        """
        try:
            from docx import Document
        except ImportError:
            return None, "python-docx not installed"

        try:
            p = Path(file_path)
            if p.suffix.lower() != ".docx":
                return None, "Chỉ hỗ trợ Word .docx (file .doc không hỗ trợ)"

            doc = Document(file_path)

            texts: List[str] = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text:
                    texts.append(para.text)

            # Tables (common when exported as a list)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or "").strip()
                        if t:
                            texts.append(t)

            raw = "\n".join(texts)
            urls = _URL_RE.findall(raw)

            # Deduplicate while preserving order
            seen = set()
            unique_urls: List[str] = []
            for u in urls:
                u = u.strip().rstrip(").,;]}>\"'")
                if not u or u in seen:
                    continue
                seen.add(u)
                unique_urls.append(u)

            videos = [{"url": u, "title": "Auto-generated"} for u in unique_urls]
            return videos, None
        except Exception as e:
            return None, f"Lỗi đọc Word: {e}"

